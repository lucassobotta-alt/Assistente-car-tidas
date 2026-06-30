import re
import json
import unicodedata
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# ── Funções de Áudio e Voz ────────────────────────────────────────────────────

def _normalizar(texto: str) -> str:
    return ''.join(
        c for c in unicodedata.normalize('NFD', texto.lower())
        if unicodedata.category(c) != 'Mn'
    )

def parse_comando_voz(texto: str) -> dict:
    t = _normalizar(texto)
    updates = {}
    nums = re.findall(r'\d+[,.]?\d*', t)
    num = float(nums[0].replace(',', '.')) if nums else None
    lado = None
    if any(w in t for w in ['direita', 'direito', 'dir']):
        lado = 'dir'
    elif any(w in t for w in ['esquerda', 'esquerdo', 'esq']):
        lado = 'esq'

    if t.startswith('nome ') or t.startswith('paciente '):
        updates['w_nome'] = texto.split(' ', 1)[1].strip().title()
        return updates
    if any(w in t for w in ['cmi', 'medio-intimal', 'medio intimal', 'intimal']):
        if lado and num is not None:
            updates[f'w_cmi_{lado}'] = min(max(num, 0.0), 5.0)
        return updates
    if 'suboclusao' in t or ('sub' in t and 'oclusao' in t):
        for ld in ([lado] if lado else ['dir', 'esq']):
            updates[f'w_estado_aci_{ld}'] = 'Suboclusão'
        return updates
    if 'oclusao' in t:
        for ld in ([lado] if lado else ['dir', 'esq']):
            updates[f'w_estado_aci_{ld}'] = 'Oclusão'
        return updates
    if 'pervia' in t:
        for ld in ([lado] if lado else ['dir', 'esq']):
            updates[f'w_estado_aci_{ld}'] = 'Pérvia (Calcular por Velocidade)'
        return updates
    if 'vertebral' in t or 'vert' in t:
        espectro = None
        if 'hipoplasia' in t:
            espectro = 'Hipoplasia'
        elif 'total' in t or 'retrogrado' in t:
            espectro = 'Roubo Total (Fluxo Retrógrado)'
        elif 'parcial' in t or 'alternante' in t:
            espectro = 'Roubo Parcial (Fluxo Alternante)'
        elif 'latente' in t:
            espectro = 'Roubo Latente'
        elif 'normal' in t or 'anterogrado' in t:
            espectro = 'Normal (Fluxo Anterógrado)'
        if espectro and lado:
            updates[f'w_espectro_vert_{lado}'] = espectro
        elif num is not None and lado:
            updates[f'w_vps_vert_{lado}'] = num
        return updates
    if any(w in t for w in ['vps', 'velocidade', 'pico', 'sistolico']):
        if any(w in t for w in ['interna', 'aci']) and lado and num is not None:
            updates[f'w_vps_aci_{lado}'] = num
        elif any(w in t for w in ['comum', 'acc']) and lado and num is not None:
            updates[f'w_vcc_{lado}'] = num
        elif any(w in t for w in ['vertebral', 'vert']) and lado and num is not None:
            updates[f'w_vps_vert_{lado}'] = num
    return updates


def render_voice_input():
    st.markdown("#### 🎤 Entrada por Voz / Texto")
    st.markdown(
        "Digite ou use a **digitação por voz do seu sistema** e clique em ✅ Aplicar.\n\n"
        "**💡 Ativar voz:** "
        "Mac → `Fn Fn` · Windows → `Win + H` · Chrome/Android → microfone no teclado\n\n"
        "**Exemplos de comandos:** "
        "*\"nome João Silva\"* · *\"VPS carótida interna direita 150\"* · "
        "*\"CMI direita 0 vírgula 8\"* · *\"oclusão esquerda\"* · *\"vertebral direita hipoplasia\"*"
    )
    col_inp, col_btn = st.columns([4, 1])
    with col_inp:
        comando = st.text_input(
            "Comando:",
            key="stt_input",
            label_visibility="collapsed",
            placeholder='Ex: "VPS carótida interna direita 150"'
        )
    with col_btn:
        aplicar = st.button("✅ Aplicar", use_container_width=True, key="stt_aplicar")

    if aplicar and comando.strip():
        updates = parse_comando_voz(comando.strip())
        if updates:
            campos = ', '.join(updates.keys())
            st.success(f"✅ Aplicado: {comando.strip()}")
            st.session_state.update(updates)
            st.session_state['stt_input'] = ''
            st.rerun()
        else:
            st.warning("Comando não reconhecido. Verifique os exemplos acima.")


def render_audio_player(texto: str, key: str = "tts"):
    texto_js = json.dumps(texto)
    components.html(f"""
    <div style="display:flex; gap:10px; align-items:center; flex-wrap:wrap; margin:4px 0;">
        <button id="btn_play_{key}" onclick="falarTexto_{key}()"
            style="padding:8px 18px; background:#1a56db; color:#fff; border:none;
                   border-radius:6px; cursor:pointer; font-size:14px;">
            ▶️ Ouvir Laudo
        </button>
        <button id="btn_pause_{key}" onclick="pausarTexto_{key}()" style="display:none;
            padding:8px 18px; background:#f59e0b; color:#fff; border:none;
            border-radius:6px; cursor:pointer; font-size:14px;">
            ⏸ Pausar
        </button>
        <button id="btn_stop_{key}" onclick="pararTexto_{key}()" style="display:none;
            padding:8px 18px; background:#e02424; color:#fff; border:none;
            border-radius:6px; cursor:pointer; font-size:14px;">
            ⏹ Parar
        </button>
        <span id="status_{key}" style="font-size:13px; color:#555;"></span>
    </div>
    <script>
    var _paused_{key} = false;
    function falarTexto_{key}() {{
        if (window.speechSynthesis.speaking && !_paused_{key}) return;
        if (_paused_{key}) {{
            window.speechSynthesis.resume(); _paused_{key} = false;
            document.getElementById('btn_pause_{key}').textContent = '⏸ Pausar';
            document.getElementById('status_{key}').textContent = '▶ Reproduzindo...'; return;
        }}
        window.speechSynthesis.cancel();
        var partes = {texto_js}.match(/[\\s\\S]{{1,200}}(?=[.!?]|$)/g) || [{texto_js}];
        var idx = 0;
        function next() {{
            if (idx >= partes.length) {{
                document.getElementById('btn_play_{key}').style.display = 'inline-block';
                document.getElementById('btn_pause_{key}').style.display = 'none';
                document.getElementById('btn_stop_{key}').style.display = 'none';
                document.getElementById('status_{key}').textContent = '✅ Concluído'; return;
            }}
            var u = new SpeechSynthesisUtterance(partes[idx]);
            u.lang = 'pt-BR'; u.rate = 0.95;
            u.onend = function() {{ idx++; next(); }};
            window.speechSynthesis.speak(u);
        }}
        document.getElementById('btn_play_{key}').style.display = 'none';
        document.getElementById('btn_pause_{key}').style.display = 'inline-block';
        document.getElementById('btn_stop_{key}').style.display = 'inline-block';
        document.getElementById('status_{key}').textContent = '▶ Reproduzindo...';
        next();
    }}
    function pausarTexto_{key}() {{
        if (window.speechSynthesis.speaking && !_paused_{key}) {{
            window.speechSynthesis.pause(); _paused_{key} = true;
            document.getElementById('btn_pause_{key}').textContent = '▶ Retomar';
            document.getElementById('status_{key}').textContent = '⏸ Pausado';
        }} else if (_paused_{key}) {{
            window.speechSynthesis.resume(); _paused_{key} = false;
            document.getElementById('btn_pause_{key}').textContent = '⏸ Pausar';
            document.getElementById('status_{key}').textContent = '▶ Reproduzindo...';
        }}
    }}
    function pararTexto_{key}() {{
        window.speechSynthesis.cancel(); _paused_{key} = false;
        document.getElementById('btn_play_{key}').style.display = 'inline-block';
        document.getElementById('btn_pause_{key}').style.display = 'none';
        document.getElementById('btn_stop_{key}').style.display = 'none';
        document.getElementById('status_{key}').textContent = '';
    }}
    </script>
    """, height=60)


# Inicialização segura do estado da sessão
if 'lista_placas' not in st.session_state:
    st.session_state.lista_placas = []
if 'lesoes_incipientes' not in st.session_state:
    st.session_state.lesoes_incipientes = []
if 'calcificacoes_isoladas' not in st.session_state:
    st.session_state.calcificacoes_isoladas = []
if 'lesoes_nao_ateromatosas' not in st.session_state:
    st.session_state.lesoes_nao_ateromatosas = []

# Aplicar sync de velocidade de tortuosidade antes de renderizar widgets
if '_sync_vps' in st.session_state:
    for _k, _v in st.session_state['_sync_vps'].items():
        st.session_state[_k] = _v
    del st.session_state['_sync_vps']

# --- FUNÇÃO AUXILIAR PARA ESTIMAR PLAQUE-RADS ---
def retirar_prefixo_numerico(opcao_texto):
    if ". " in opcao_texto:
        return opcao_texto.split(". ", 1)[1]
    return opcao_texto

def estimar_plaque_rads(opcao_texto):
    if opcao_texto.startswith("1."):
        return "Classificação: Plaque-RADS 1"
    elif opcao_texto.startswith("2."):
        return "Classificação: Plaque-RADS 2"
    elif opcao_texto.startswith("3."):
        return "Classificação: Plaque-RADS 3"
    elif opcao_texto.startswith("4."):
        return "Classificação: Plaque-RADS 4"
    elif opcao_texto.startswith("5."):
        return "Classificação: Plaque-RADS 5"
    elif opcao_texto.startswith("6."):
        return "Classificação: Plaque-RADS 5"
    return None

# --- CLASSIFICADORES HEMODINÂMICOS ADAPTATIVOS (SBC 2023 vs NASCET) ---
def obter_texto_hemo_continuo(estado, vps_aci, vcc, tem_placa=False, diretriz="Diretriz SBC 2023", incluir_vel=True):
    if estado == "Oclusão":
        return "Oclusão", "determinando oclusão completa do vaso, caracterizada por ausência total de fluxo ao estudo Doppler pulsado e mapeamento a cores."
    elif estado == "Suboclusão":
        return "Suboclusão", "determinando suboclusão do vaso, caracterizada por estreitamento luminal severo com padrão de fluxo filiforme ('trickle flow') ao estudo Doppler."

    relacao = round(vps_aci / max(vcc, 1), 2)
    vel_aci  = f" de {vps_aci} cm/s" if incluir_vel else ""
    vel_rel  = f" e relação artéria carótida interna / artéria carótida comum de {relacao}" if incluir_vel else ""
    vel_rel2 = f" e relação ACI/ACC de {relacao}" if incluir_vel else ""

    if diretriz == "Diretriz SBC 2023":
        if vps_aci < 140:
            if tem_placa:
                det = f", caracterizada por velocidade de pico sistólico na artéria carótida interna{vel_aci}" if incluir_vel else ""
                return "Estenose < 50%", f"determinando estenose leve (<50% pelos critérios da Diretriz SBC 2023){det}."
            else:
                return "Normal", "com fluxo bifásico anterógrado de baixa resistência, caracterizado por diástole sustentada e velocidades dentro da normalidade, compatível com irrigação de leito encefálico de baixa impedância. Não há sinais de estenose ou turbulência."

        if vps_aci > 400 or relacao > 5.0:
            det = f", caracterizada por acentuada elevação das velocidades de fluxo com VPS na artéria carótida interna{vel_aci}{vel_rel}" if incluir_vel else ""
            return "Estenose > 90%", f"determinando estenose acentuada (>90% pelos critérios da Diretriz SBC 2023){det}."
        elif 230 < vps_aci <= 400 or relacao > 4.0:
            det = f", caracterizada por VPS na artéria carótida interna{vel_aci}{vel_rel}" if incluir_vel else ""
            return "Estenose de 70-89%", f"determinando estenose hemodinamicamente significativa (70-89% pelos critérios da Diretriz SBC 2023){det}."
        elif 3.2 <= relacao <= 4.0:
            det = f", caracterizada por relação artéria carótida interna / artéria carótida comum de {relacao}{' e VPS na artéria carótida interna' + vel_aci if incluir_vel else ''}" if incluir_vel else ""
            return "Estenose de 60-69%", f"determinando estenose moderada (60-69% pelos critérios da Diretriz SBC 2023){det}."
        else:
            det = f", caracterizada por VPS na artéria carótida interna{vel_aci}{vel_rel}" if incluir_vel else ""
            return "Estenose de 50-59%", f"determinando estenose moderada (50-59% pelos critérios da Diretriz SBC 2023){det}."

    else:  # Critérios Clássicos do NASCET
        if vps_aci < 125:
            if tem_placa:
                det = f", com VPS na artéria carótida interna{vel_aci}" if incluir_vel else ""
                return "Estenose < 50%", f"determinando estenose leve (<50% pelos critérios do Consenso NASCET){det}."
            else:
                return "Normal", "apresentando padrão de velocidades normais ao estudo Doppler, sem critérios para estenose hemodinâmica pelo Consenso NASCET."

        if vps_aci >= 230 or relacao >= 4.0:
            det = f", caracterizada por VPS na artéria carótida interna{vel_aci}{vel_rel2}" if incluir_vel else ""
            return "Estenose ≥ 70%", f"determinando estenose severa (≥70% pelos critérios do Consenso NASCET){det}."
        else:
            det = f", caracterizada por VPS na artéria carótida interna{vel_aci}{vel_rel2}" if incluir_vel else ""
            return "Estenose de 50-69%", f"determinando estenose moderada (50-69% pelos critérios do Consenso NASCET){det}."

def texto_tortuosidade_report(na):
    seg = na['segmento']
    vaso = na['vaso'].lower()
    lado = na['lado'].lower()
    if na['hemo'] and na.get('vps_tort', 0) > 0:
        hemo_text = f"com alteração hemodinâmica caracterizada por aumento expressivo na velocidade de pico sistólico em até {na['vps_tort']:.0f} cm/s no ponto da maior curvatura"
    else:
        hemo_text = "sem alteração hemodinâmica"
    subtipo = na['subtipo']
    if subtipo == "Acotovelamento (Kinking)":
        return f"Tortuosidade de trajeto com acotovelamento (<i>Kinking</i>) no segmento cervical {seg} da {vaso} {lado}, {hemo_text}."
    elif subtipo == "Looping":
        return f"Tortuosidade de trajeto sem acotovelamento, configurando <i>loop</i>, no segmento cervical {seg} da {vaso} {lado}, {hemo_text}."
    else:
        return f"Tortuosidade de trajeto configurando <i>coil</i> no segmento cervical {seg} da {vaso} {lado}, {hemo_text}."

def impressao_tortuosidade_text(na):
    seg = na['segmento']
    vaso = na['vaso'].lower()
    lado = na['lado'].lower()
    hemo_imp = "com alteração hemodinâmica" if na['hemo'] else "sem alteração hemodinâmica"
    subtipo = na['subtipo']
    if subtipo == "Acotovelamento (Kinking)":
        return f"– Acotovelamento (Kinking) no segmento cervical {seg} da {vaso} {lado}, {hemo_imp}."
    elif subtipo == "Looping":
        return f"– Looping no segmento cervical {seg} da {vaso} {lado}, {hemo_imp}."
    else:
        return f"– Coil no segmento cervical {seg} da {vaso} {lado}, {hemo_imp}."

def avaliar_vertebral(espectro, vps_vert):
    if espectro == "Normal (Fluxo Anterógrado)":
        if vps_vert >= 100:
            return "Estenose de Vertebral (>50%)", f"Artéria vertebral apresentando fluxo anterógrado com acentuada elevação focal de velocidades (VPS de {vps_vert} cm/s) e turbulência local, compatível com estenose segmentar superior a 50%."
        else:
            return "Normal", "Artéria vertebral pérvia, com fluxo anterógrado de baixa resistência e diástole sustentada, compatível com adequada perfusão vertebrobasilar."
    elif espectro == "Hipoplasia": 
        return "Hipoplasia de Vertebral", f"Artéria vertebral apresentando fluxo anterógrado de baixa resistência, porém exibindo calibre reduzido e velocidades proporcionalmente baixas (VPS de {vps_vert} cm/s), compatível com variante anatomofuncional (hipoplasia)."
    elif espectro == "Roubo Latente": 
        return "Sinal de Roubo Latente da Subclávia", "Artéria vertebral apresentando fluxo anterógrado, porém com morfologia de onda alterada devido a uma desaceleração mesosistólica abrupta, sugerindo alteração hemodinâmica inicial por estenose da artéria subclávia proximal ipsilateral."
    elif espectro == "Roubo Parcial (Fluxo Alternante)": 
        return "Sinal de Roubo Parcial da Subclávia", "Artéria vertebral apresentando padrão de fluxo alternante, caracterizado por vetor sistólico retrógrado e vetor diastólico anterógrado, indicando inversão parcial do fluxo por estenose acentuada da artéria subclávia proximal ipsilateral."
    elif espectro == "Roubo Total (Fluxo Retrógrado)": 
        return "Sinal de Roubo Total da Subclávia", "Artéria vertebral apresentando inversão completa e contínua do seu vetor de fluxo, confirmando o fenômeno de roubo de subclávia secundário a oclusão da artéria subclávia proximal ipsilateral."
    return "Alterada", f"Artéria vertebral com alterações inespecíficas do padrão de fluxo. VPS: {vps_vert} cm/s."

# --- CONFIGURAÇÃO DA INTERFACE STREAMLIT ---
st.title("⚕️ Assistente de Laudos: Duplex Scan Arterial Carotídeo")

# ==========================================
#       PAINEL DE CONTROLE LATERAL
# ==========================================
with st.sidebar:
    st.markdown("## ⚙️ Painel de Controle Avançado")
    
    # 1. Parâmetros Científicos / Diretrizes
    st.markdown("### 📚 Critérios Científicos")
    diretriz_selecionada = "Diretriz SBC 2023"
    st.markdown(
        "**Diretrizes de referência:**<br>"
        "📄 <a href='https://abccardiol.org/article/atualizacao-da-recomendacao-para-avaliacao-da-doenca-das-arterias-carotidas-e-vertebrais-pela-ultrassonografia-vascular-dic-cbr-sbacv-2023/' target='_blank'>"
        "Ultrassonografia Carotídea e Vertebral – DIC/CBR/SBACV 2023</a><br>"
        "📄 <a href='https://abccardiol.org/wp-content/uploads/articles_xml/0066-782X-abc-122-09-e20250640/0066-782X-abc-122-09-e20250640.x66747.pdf' target='_blank'>"
        "Diretriz Brasileira de Dislipidemias e Prevenção da Aterosclerose 2025</a><br>"
        "📄 <a href='https://www.jacc.org/doi/10.1016/j.jcmg.2023.09.005' target='_blank'>"
        "Carotid Plaque-RADS – JACC 2023</a>",
        unsafe_allow_html=True
    )
    ano_dislipidemia = st.selectbox("Ano da Diretriz de Dislipidemia:", ["2025", "2023"], index=0)
    
    st.markdown("---")
    
    # 2. Configurações de Formatação Visual do Documento
    st.markdown("### 📝 Formatação Externa (.docx)")
    fonte_doc = st.selectbox("Família da Fonte:", ["Arial", "Calibri", "Times New Roman"], index=0)
    tamanho_fonte = st.slider("Tamanho do Texto (pt):", 10, 14, 11, step=1)
    espacamento_linhas = st.slider("Espaçamento entre Linhas:", 1.0, 1.5, 1.15, step=0.05)
    quebrar_pagina_diag = st.toggle("Separar Impressão Diagnóstica em Nova Página", value=False)

    modo_saida = st.radio(
        "Modo de saída do laudo:",
        ["Somente DOCX", "Somente Visualização", "Visualização + DOCX"],
        index=2
    )
    
    st.markdown("---")
    
    # 3. Identidade Visual e Assinatura Automatizada
    st.markdown("### ✍️ Identidade & Assinatura")
    nome_clinica = st.text_input("Cabeçalho / Nome da Clínica:", placeholder="Ex: Instituto de Diagnóstico por Imagem")
    nome_medico = st.text_input("Nome do Médico:", "")
    colcrm1, colcrm2 = st.columns([2,1])
    with colcrm1:
        crm_medico = st.text_input("CRM:")
    with colcrm2:
        crm_uf = st.selectbox("UF CRM", ["AC","AL","AP","AM","BA","CE","DF","ES","GO","MA","MT","MS","MG","PA","PB","PR","PE","PI","RJ","RN","RS","RO","RR","SC","SP","SE","TO"], index=25)
    rqe_medico = st.text_input("RQE:")
    incluir_observacoes = st.toggle("Incluir observações complementares", value=True)
    incluir_velocidades = st.toggle("Incluir velocidades (VPS / relação ACI-ACC) no laudo", value=True)

    st.markdown("---")
    
    if st.button("🔄 Resetar Todos os Parâmetros", use_container_width=True, type="secondary"):
        st.session_state.lista_placas = []
        st.session_state.lesoes_incipientes = []
        st.session_state.calcificacoes_isoladas = []
        st.session_state.lesoes_nao_ateromatosas = []
        # Redefine todos os widgets hemodinâmicos para seus valores padrão
        st.session_state.w_nome = ""
        st.session_state.w_tecnica = "1. Exame sem limitações técnicas"
        st.session_state.w_cmi_dir = 0.4
        st.session_state.w_estado_aci_dir = "Pérvia (Calcular por Velocidade)"
        st.session_state.w_vps_aci_dir = 0.0
        st.session_state.w_vcc_dir = 0.0
        st.session_state.w_ace_dir = "Com padrão espectral de alta resistência, compatível com perfusão de leitos musculares extracranianos."
        st.session_state.w_espectro_vert_dir = "Normal (Fluxo Anterógrado)"
        st.session_state.w_vps_vert_dir = 0.0
        st.session_state.w_cmi_esq = 0.4
        st.session_state.w_estado_aci_esq = "Pérvia (Calcular por Velocidade)"
        st.session_state.w_vps_aci_esq = 0.0
        st.session_state.w_vcc_esq = 0.0
        st.session_state.w_ace_esq = "Com padrão espectral de alta resistência, compatível com perfusão de leitos musculares extracranianos."
        st.session_state.w_espectro_vert_esq = "Normal (Fluxo Anterógrado)"
        st.session_state.w_vps_vert_esq = 0.0
        st.toast("🔄 Todos os dados clínicos foram limpos!")
        st.rerun()

# Opções técnicas fixas
opcoes_tecnicas = {
    "1. Exame sem limitações técnicas": "Exame realizado em decúbito dorsal, utilizando transdutor linear de alta frequência, com avaliação bidimensional, mapeamento de fluxo a cores e Doppler pulsado, sem limitações técnicas.",
    "2. Exame com limitação por condições anatômicas desfavoráveis": "Exame realizado em decúbito dorsal, utilizando transdutor linear de alta frequência, com avaliação bidimensional, mapeamento de fluxo a cores e Doppler pulsado. Devido a condições anatômicas desfavoráveis para insonação dos vasos cervicais, foi necessária avaliação complementar com transdutor convexo, o que pode reduzir a sensibilidade para identificação de placas ateroscleróticas de pequenas dimensões.",
    "3. Exame realizado à beira do leito (UTI)": "Exame realizado à beira do leito em unidade de terapia intensiva, utilizando transdutor linear de alta frequência, com limitações técnicas inerentes às condições do exame.",
    "4. Exame à beira do leito (UTI) com curativos cervicais": "Exame realizado à beira do leito em unidade de terapia intensiva, utilizando transdutor linear de alta frequência, com limitações técnicas inerentes às condições do exame e à presença de curativos cervicais sobre acessos jugulares."
}

col_id1, col_id2 = st.columns([2, 2])
with col_id1:
    nome = st.text_input("Nome do Paciente", "", key="w_nome")
with col_id2:
    opcao_selecionada = st.selectbox("Condições Técnicas do Exame:", list(opcoes_tecnicas.keys()), key="w_tecnica")
    texto_tecnica_final = opcoes_tecnicas[opcao_selecionada]

st.markdown("---")
render_voice_input()
st.markdown("---")
st.markdown("### 📊 Parâmetros Hemodinâmicos")
col_hemo_dir, col_hemo_esq = st.columns(2)

with col_hemo_dir:
    st.header("LADO DIREITO")
    cmi_dir = st.number_input("CMI Artéria Carótida Comum Direita (mm)", min_value=0.0, max_value=5.0, value=0.4, step=0.1, key="w_cmi_dir")
    estado_aci_dir = st.selectbox("Estado da Artéria Carótida Interna Direita", ["Pérvia (Calcular por Velocidade)", "Suboclusão", "Oclusão"], key="w_estado_aci_dir")
    vps_aci_dir = st.number_input("VPS Artéria Carótida Interna Direita (cm/s)", min_value=0.0, value=0.0, step=5.0, key="w_vps_aci_dir")
    vcc_dir = st.number_input("VPS Artéria Carótida Comum Direita (cm/s)", min_value=0.0, value=0.0, step=5.0, key="w_vcc_dir")
    ace_dir = st.selectbox("Artéria Carótida Externa Direita", ["Com padrão espectral de alta resistência, compatível com perfusão de leitos musculares extracranianos.", "Alterada / Estenose hemodinâmica"], key="w_ace_dir")
    espectro_vert_dir = st.selectbox("Espectro Artéria Vertebral Direita", ["Normal (Fluxo Anterógrado)", "Hipoplasia", "Roubo Latente", "Roubo Parcial (Fluxo Alternante)", "Roubo Total (Fluxo Retrógrado)"], key="w_espectro_vert_dir")
    vps_vert_dir = st.number_input("VPS Artéria Vertebral Direita (cm/s)", min_value=0.0, value=0.0, step=5.0, key="w_vps_vert_dir")

with col_hemo_esq:
    st.header("LADO ESQUERDO")
    cmi_esq = st.number_input("CMI Artéria Carótida Comum Esquerda (mm)", min_value=0.0, max_value=5.0, value=0.4, step=0.1, key="w_cmi_esq")
    estado_aci_esq = st.selectbox("Estado da Artéria Carótida Interna Esquerda", ["Pérvia (Calcular por Velocidade)", "Suboclusão", "Oclusão"], key="w_estado_aci_esq")
    vps_aci_esq = st.number_input("VPS Artéria Carótida Interna Esquerda (cm/s)", min_value=0.0, value=0.0, step=5.0, key="w_vps_aci_esq")
    vcc_esq = st.number_input("VPS Artéria Carótida Comum Esquerda (cm/s)", min_value=0.0, value=0.0, step=5.0, key="w_vcc_esq")
    ace_esq = st.selectbox("Artéria Carótida Externa Esquerda", ["Com padrão espectral de alta resistência, compatível com perfusão de leitos musculares extracranianos.", "Alterada / Estenose hemodinâmica"], key="w_ace_esq")
    espectro_vert_esq = st.selectbox("Espectro Artéria Vertebral Esquerda", ["Normal (Fluxo Anterógrado)", "Hipoplasia", "Roubo Latente", "Roubo Parcial (Fluxo Alternante)", "Roubo Total (Fluxo Retrógrado)"], key="w_espectro_vert_esq")
    vps_vert_esq = st.number_input("VPS Artéria Vertebral Esquerda (cm/s)", min_value=0.0, value=0.0, step=5.0, key="w_vps_vert_esq")

st.markdown("---")

with st.expander("2. Lesões Não Ateromatosas (Tortuosidades e Vasculite)"):
    categoria_na = st.radio("Tipo de lesão:", ["Tortuosidade", "Vasculite"], horizontal=True, key="cat_na")

    if categoria_na == "Tortuosidade":
        col_t1, col_t2, col_t3 = st.columns(3)
        with col_t1:
            subtipo_tort = st.selectbox("Tipo de tortuosidade:", ["Acotovelamento (Kinking)", "Looping", "Coil"], key="subtipo_tort")
            lado_tort = st.selectbox("Lado:", ["Direito", "Esquerdo", "Bilateral"], key="lado_tort")
        with col_t2:
            vaso_tort = st.selectbox("Artéria:", ["Artéria carótida interna", "Artéria carótida comum"], key="vaso_tort")
            segmento_tort = st.selectbox("Segmento cervical:", ["proximal", "médio", "distal"], key="seg_tort")
        with col_t3:
            hemo_tort = st.toggle("Há alteração hemodinâmica?", value=False, key="hemo_tort")
            vps_tort = 0.0
            if hemo_tort:
                vps_tort = st.number_input("VPS no ponto de maior curvatura (cm/s):", min_value=0.0, value=0.0, step=5.0, key="vps_tort_val")

        if st.button("💾 Registrar Tortuosidade"):
            lados_add = ["Direito", "Esquerdo"] if lado_tort == "Bilateral" else [lado_tort]
            for ld in lados_add:
                item_na = {"categoria": "tortuosidade", "subtipo": subtipo_tort, "vaso": vaso_tort, "lado": ld, "segmento": segmento_tort, "hemo": hemo_tort, "vps_tort": vps_tort}
                if item_na not in st.session_state.lesoes_nao_ateromatosas:
                    st.session_state.lesoes_nao_ateromatosas.append(item_na)
                if hemo_tort and vps_tort > 0:
                    suffix_ld = "dir" if ld == "Direito" else "esq"
                    if '_sync_vps' not in st.session_state:
                        st.session_state['_sync_vps'] = {}
                    if "interna" in vaso_tort.lower():
                        st.session_state['_sync_vps'][f"w_vps_aci_{suffix_ld}"] = vps_tort
                    elif "comum" in vaso_tort.lower():
                        st.session_state['_sync_vps'][f"w_vcc_{suffix_ld}"] = vps_tort
            st.toast("✅ Tortuosidade registrada com sucesso!")
            st.rerun()

    else:
        col_v1, col_v2, col_v3 = st.columns(3)
        with col_v1:
            vaso_na = st.selectbox("Vaso:", ["Artéria carótida interna", "Artéria carótida comum", "Bulbo carotídeo", "Artéria vertebral"], key="vaso_vasc")
            lado_na = st.selectbox("Lado:", ["Direito", "Esquerdo", "Bilateral"], key="lado_vasc")
        with col_v2:
            tipo_na = st.selectbox("Tipo de vasculite:", [
                "Vasculite / Arterite (Espessamento parietal concêntrico e homogêneo - Sinal do Halo)",
                "Vasculite / Arterite (Espessamento difuso irregular de padrão inflamatório)"
            ], key="tipo_vasc")
        with col_v3:
            mensuracao_na = st.number_input("Espessura parietal (mm):", min_value=0.0, max_value=10.0, value=2.1, step=0.1, key="med_vasc")
            hemo_na = st.toggle("Alteração hemodinâmica local?", value=False, key="hemo_vasc")

        if st.button("💾 Registrar Vasculite"):
            lados_add = ["Direito", "Esquerdo"] if lado_na == "Bilateral" else [lado_na]
            for ld in lados_add:
                item_na = {"categoria": "vasculite", "vaso": vaso_na, "lado": ld, "tipo": tipo_na, "medida": mensuracao_na, "hemo": hemo_na}
                if item_na not in st.session_state.lesoes_nao_ateromatosas:
                    st.session_state.lesoes_nao_ateromatosas.append(item_na)
            st.toast("✅ Vasculite registrada com sucesso!")

    if st.session_state.lesoes_nao_ateromatosas:
        for idx, na in enumerate(st.session_state.lesoes_nao_ateromatosas):
            if na.get('categoria') == 'tortuosidade':
                hemo_lbl = f"com hemo ({na.get('vps_tort', 0):.0f} cm/s)" if na['hemo'] else "sem alteração hemo"
                st.write(f"• `{idx+1:02d}` **{na['vaso']} {na['lado']}**: {na['subtipo']} — {na['segmento']} — {hemo_lbl}")
            else:
                suf_h = "com repercussão" if na.get('hemo') else "sem repercussão"
                st.write(f"• `{idx+1:02d}` **{na.get('vaso','')} {na.get('lado','')}**: {na.get('tipo','')} ({na.get('medida','')} mm) — {suf_h}")
        if st.button("❌ Limpar Lista Não Ateromatosa"):
            st.session_state.lesoes_nao_ateromatosas = []
            st.rerun()

st.markdown("---")

with st.expander("3. Lesões Estruturais Incipientes (Alterações Precoces ≤ 1.5 mm)"):
    col_inc1, col_inc2, col_inc3 = st.columns(3)
    with col_inc1:
        vaso_inc = st.selectbox("Vaso (Incipiente):", [
            "Artéria carótida comum direita", "Bulbo carotídeo direito", "Artéria carótida interna direita",
            "Artéria carótida comum esquerda", "Bulbo carotídeo esquerdo", "Artéria carótida interna esquerda"
        ], key="vaso_inc")
    with col_inc2:
        local_inc = st.selectbox("Localização / Segmento (Incipiente):", ["terço proximal", "terço médio", "terço distal", "segmento total"], key="local_inc")
    with col_inc3:
        espessura_inc = st.number_input("Espessura Máxima (mm) [Max: 1.5mm]:", min_value=0.5, max_value=1.5, value=1.2, step=0.1)

    if st.button("💾 Registrar Lesão Incipiente"):
        nova_incipiente = {"vaso": vaso_inc, "localizacao": local_inc, "espessura": espessura_inc}
        if nova_incipiente not in st.session_state.lesoes_incipientes:
            st.session_state.lesoes_incipientes.append(nova_incipiente)
        st.toast("✅ Lesão estrutural incipiente registrada!")

    if st.session_state.lesoes_incipientes:
        for idx, inc in enumerate(st.session_state.lesoes_incipientes):
            st.write(f"• `{idx+1:02d}` **{inc['vaso']}** ({inc['localizacao']}) — Espessura: {inc['espessura']} mm.")
        if st.button("❌ Limpar Lista Incipientes"):
            st.session_state.lesoes_incipientes = []
            st.rerun()

st.markdown("---")

with st.expander("4. Mapeamento de Placas Ateroscleróticas (Consolidadas ≥ 2.0 mm)"):
    c_vaso, c_loc = st.columns(2)
    with c_vaso:
        vaso_selecionado = st.selectbox("Selecione a Artéria:", [
            "Artéria carótida comum direita", "Bulbo carotídeo direito", "Artéria carótida interna direita",
            "Artéria carótida comum esquerda", "Bulbo carotídeo esquerdo", "Artéria carótida interna esquerda"
        ], key="vaso_placa")
    with c_loc:
        localizacao_selecionada = st.selectbox("Localização / Segmento:", ["terço proximal", "terço médio", "terço distal", "segmento total/bifurcação"], key="local_placa")

    culpada_hemo = st.toggle("Esta lesão determina estenose significativa?", value=False, key="culpada_placa")

    c1, c2, c3 = st.columns(3)
    with c1:
        composicao = st.selectbox("Composição / Tipo da Placa (Plaque-RADS):", [
            "1. Placa calcificada",
            "2. Placa uniformemente ecogênica (fibrosa estável)",
            "3. Placa predominantemente ecogênica (fibrosa)",
            "4. Placa com componente anecogênico delimitado por halo hiperecogênico",
            "5. Placa predominantemente anecogênica",
            "6. Placa uniformemente anecogênica sem cápsula fibrosa definida"
        ])
        usar_plaque_rads = st.toggle("Deseja estimar e incluir a classificação Plaque-RADS?", value=False)
        espessura = st.number_input("Espessura Máxima (mm):", min_value=2.0, max_value=20.0, value=2.5, step=0.1)
    with c2:
        superficie = st.selectbox("Superfície:", ["Regular", "Irregular", "Ulcerada"])
        profundidade_ulcura = st.text_input("Nicho Ulceroso:", placeholder="Ex: medindo 1.2 mm") if superficie == "Ulcerada" else ""
    with c3:
        st.markdown("**Achados Adicionais:**")
        tem_trombo = st.checkbox("Trombo adjacente")
        tem_fluxo = st.checkbox("Fluxo intraplaca")
        tem_calc_intra = st.checkbox("Focos de calcificação de permeio")

    if st.button("💾 Registrar Placa Ateromatosa"):
        achados_texto_lista = []
        if tem_trombo: achados_texto_lista.append("trombo adjacente")
        if tem_fluxo: achados_texto_lista.append("fluxo intraplaca")
        if tem_calc_intra: achados_texto_lista.append("calcificação intraplaca")
        pr_estimado = estimar_plaque_rads(composicao) if usar_plaque_rads else None
        composicao_limpa = retirar_prefixo_numerico(composicao)
        nova_placa = {
            "vaso": vaso_selecionado, "localizacao": localizacao_selecionada, "culpada_hemo": culpada_hemo,
            "composicao_texto": composicao_limpa, "superficie_texto": superficie, "profundidade_ulcura": profundidade_ulcura,
            "espessura": espessura, "achados_adicionais": achados_texto_lista,
            "plaque_rads": pr_estimado
        }
        if nova_placa not in st.session_state.lista_placas:
            st.session_state.lista_placas.append(nova_placa)
        st.toast("✅ Placa registrada!")

    if st.session_state.lista_placas:
        for idx, p in enumerate(st.session_state.lista_placas):
            pr_tag = f" | {p['plaque_rads']}" if p['plaque_rads'] else ""
            st.write(f"`Item {idx+1:02d}` **{p['vaso']}** ({p['localizacao']}) — {p['composicao_texto']} | {p['espessura']} mm ({p['superficie_texto'].lower()}){pr_tag}.")
        _placas_alerta = [p for p in st.session_state.lista_placas
                         if p['composicao_texto'] != "Placa calcificada" and p['espessura'] >= 3.0]
        for _pa in _placas_alerta:
            st.warning(
                f"⚠️ **Placa de risco elevado — {_pa['vaso']} ({_pa['localizacao']}):** "
                f"placa não completamente calcificada com espessura de {_pa['espessura']} mm (≥ 3 mm). "
                f"Característica associada a maior risco de instabilidade plaqueária."
            )
        if st.button("❌ Limpar Lista de Placas"):
            st.session_state.lista_placas = []
            st.rerun()

st.markdown("---")

with st.expander("5. Calcificações Parietais Isoladas (Sem Formação de Placa)"):
    tem_calc_parietal = st.checkbox("Incluir registro de Calcificações Parietais Isoladas?")
    if tem_calc_parietal:
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            lado_calc = st.selectbox("Lado da Calcificação:", ["Direito", "Esquerdo", "Bilateral"])
        with col_c2:
            topografia_calc = st.selectbox("Localização Anatômica:", ["bulbo carotídeo", "artéria carótida comum", "artéria carótida interna"])
        
        if st.button("💾 Adicionar Calcificação Isolada"):
            lados_calc_add = ["Direito", "Esquerdo"] if lado_calc == "Bilateral" else [lado_calc]
            for ld in lados_calc_add:
                item_calc = {"lado": ld, "topografia": topografia_calc}
                if item_calc not in st.session_state.calcificacoes_isoladas:
                    st.session_state.calcificacoes_isoladas.append(item_calc)
            st.toast("✅ Calcificação registrada!")

        if st.session_state.calcificacoes_isoladas:
            for idx, c in enumerate(st.session_state.calcificacoes_isoladas):
                st.write(f"• `{idx+1:02d}` Calcificação no Lado **{c['lado']}** — {c['topografia'].title()}")
            if st.button("❌ Limpar Lista de Calcificações"):
                st.session_state.calcificacoes_isoladas = []
                st.rerun()

st.markdown("---")

with st.expander("6. Template Personalizado do Laudo (opcional)"):
    st.markdown(
        "Cole abaixo um modelo de laudo usando os **marcadores** listados. "
        "Se o campo estiver vazio, o formato padrão será utilizado.\n\n"
        "**Marcadores disponíveis:**\n"
        "`{{nome}}` `{{tecnica}}` `{{cmi_dir}}` `{{cmi_esq}}`\n\n"
        "`{{acc_dir}}` `{{bulbo_dir}}` `{{aci_dir}}` `{{ace_dir}}` `{{vertebral_dir}}`\n\n"
        "`{{acc_esq}}` `{{bulbo_esq}}` `{{aci_esq}}` `{{ace_esq}}` `{{vertebral_esq}}`\n\n"
        "`{{impressao}}` `{{observacoes}}` `{{medico}}` `{{crm}}` `{{rqe}}`"
    )
    template_usuario = st.text_area(
        "Template personalizado:",
        height=250,
        placeholder="Exemplo:\nDUPLEX SCAN DAS ARTÉRIAS CARÓTIDAS E VERTEBRAIS\nPaciente: {{nome}}\n\nLADO DIREITO\n{{acc_dir}}\n{{bulbo_dir}}\n{{aci_dir}}\n\nLADO ESQUERDO\n{{acc_esq}}\n{{bulbo_esq}}\n{{aci_esq}}\n\nIMPRESSÃO DIAGNÓSTICA\n{{impressao}}\n\n{{observacoes}}\n\n{{medico}}\n{{crm}}\n{{rqe}}"
    )

st.markdown("---")
gerar_laudo = st.button("🚀 Gerar Laudo Clínico Completo", use_container_width=True)

# --- MOTOR DE GERAÇÃO TEXTUAL ADAPTATIVO & FORMATADOR ---
if gerar_laudo:
    tem_placa_aci_dir = any("interna direita" in p['vaso'].lower() for p in st.session_state.lista_placas)
    tem_placa_aci_esq = any("interna esquerda" in p['vaso'].lower() for p in st.session_state.lista_placas)

    status_aci_dir_limpo, sufixo_hemo_aci_dir = obter_texto_hemo_continuo(estado_aci_dir, vps_aci_dir, vcc_dir, tem_placa_aci_dir, diretriz_selecionada, incluir_velocidades)
    status_aci_esq_limpo, sufixo_hemo_aci_esq = obter_texto_hemo_continuo(estado_aci_esq, vps_aci_esq, vcc_esq, tem_placa_aci_esq, diretriz_selecionada, incluir_velocidades)

    # Override hemodinâmico quando tortuosidade é a causa da aceleração focal
    tort_hemo_aci_dir = next((na for na in st.session_state.lesoes_nao_ateromatosas
                              if na.get('categoria') == 'tortuosidade' and na['hemo']
                              and na['lado'] == 'Direito' and 'interna' in na['vaso'].lower()), None)
    tort_hemo_aci_esq = next((na for na in st.session_state.lesoes_nao_ateromatosas
                              if na.get('categoria') == 'tortuosidade' and na['hemo']
                              and na['lado'] == 'Esquerdo' and 'interna' in na['vaso'].lower()), None)
    if tort_hemo_aci_dir:
        vps_t = tort_hemo_aci_dir.get('vps_tort', 0)
        sufixo_hemo_aci_dir = f"apresentando elevação focal da velocidade de pico sistólico (VPS de {vps_t:.0f} cm/s) no ponto de maior curvatura, secundária à tortuosidade de trajeto, sem evidência de processo estenótico ateromatoso associado."
        status_aci_dir_limpo = "Tortuosidade com Repercussão Hemodinâmica"
    elif vps_aci_dir == 0.0 and estado_aci_dir == "Pérvia (Calcular por Velocidade)":
        status_aci_dir_limpo = "Velocidade Não Informada"
    if tort_hemo_aci_esq:
        vps_t = tort_hemo_aci_esq.get('vps_tort', 0)
        sufixo_hemo_aci_esq = f"apresentando elevação focal da velocidade de pico sistólico (VPS de {vps_t:.0f} cm/s) no ponto de maior curvatura, secundária à tortuosidade de trajeto, sem evidência de processo estenótico ateromatoso associado."
        status_aci_esq_limpo = "Tortuosidade com Repercussão Hemodinâmica"
    elif vps_aci_esq == 0.0 and estado_aci_esq == "Pérvia (Calcular por Velocidade)":
        status_aci_esq_limpo = "Velocidade Não Informada"

    # Texto de fluxo normal para usar quando há placa mas sem velocidade informada
    _, _txt_fluxo_normal = obter_texto_hemo_continuo("Pérvia (Calcular por Velocidade)", 0, 0, False, diretriz_selecionada, incluir_velocidades)
    txt_fluxo_normal_aci = _txt_fluxo_normal.lstrip("com ")

    _, texto_vert_dir = avaliar_vertebral(espectro_vert_dir, vps_vert_dir)
    _, texto_vert_esq = avaliar_vertebral(espectro_vert_esq, vps_vert_esq)

    doc = Document()
    
    # Configuração Dinâmica da Fonte Padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = fonte_doc
    font.size = Pt(tamanho_fonte)
    
    # Funções Auxiliares de Injeção de Texto Adaptativo
    def adicionar_titulo(texto):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(texto.upper())
        run.bold = True
        run.font.name = fonte_doc
        run.font.size = Pt(tamanho_fonte + 1)

    def adicionar_subtitulo(texto):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(texto)
        run.bold = True
        run.font.name = fonte_doc
        run.font.size = Pt(tamanho_fonte)

    def adicionar_texto_esquerda(texto, bold_prefix=None, force_page_break=False, italico=False, fonte_menor=False):
        p = doc.add_paragraph()
        if force_page_break:
            p.insert_paragraph_before().add_run().add_break()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = espacamento_linhas
        tamanho_atual = Pt(tamanho_fonte - 2) if fonte_menor else Pt(tamanho_fonte)

        if bold_prefix:
            run_pre = p.add_run(bold_prefix)
            run_pre.bold = True
            run_pre.italic = italico
            run_pre.font.name = fonte_doc
            run_pre.font.size = tamanho_atual

        parts = re.split(r'(<i>.*?</i>)', texto)
        for part in parts:
            if part.startswith('<i>') and part.endswith('</i>'):
                run = p.add_run(part[3:-4])
                run.italic = True
            else:
                run = p.add_run(part)
                run.italic = italico
            run.font.name = fonte_doc
            run.font.size = tamanho_atual

    # --- CONSTRUÇÃO DOS PARÁGRAFOS DO RELATÓRIO ---

    txt_comum_dir = f"Artéria carótida comum direita pérvia, com diâmetro e trajeto conservados, apresentando fluxo bifásico anterógrado de baixa resistência. Espessura do complexo médio-intimal: {cmi_dir} mm."
    for na in [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Direito" and "comum" in x['vaso'].lower()]:
        if na.get('categoria') == 'tortuosidade':
            txt_comum_dir += " " + texto_tortuosidade_report(na)
        else:
            suf_h = "com alteração hemodinâmica local" if na.get('hemo') else "sem repercussão hemodinâmica"
            txt_comum_dir += f" Identifica-se espessamento parietal sugestivo de processo inflamatório, medindo {na.get('medida', 0)} mm, {suf_h}."
    for inc in [x for x in st.session_state.lesoes_incipientes if "carótida comum direita" in x['vaso'].lower()]:
        txt_comum_dir += f" Identifica-se espessamento focal da camada médio-intimal no {inc['localizacao']}, medindo {inc['espessura']} mm."
    for p in [x for x in st.session_state.lista_placas if "carótida comum direita" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        comp_dir = p['composicao_texto'].lower().removeprefix("placa ")
        txt_comum_dir += f" Identifica-se na parede uma placa de ateroma {comp_dir}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Direito" and "comum" in c['topografia']:
            txt_comum_dir += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."

    txt_bulbo_dir = "Bulbo carotídeo direito pérvio, com diâmetro e trajeto conservados."
    tem_achado_bulbo_dir = False
    for na in [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Direito" and "bulbo" in x['vaso'].lower()]:
        if na.get('categoria') == 'tortuosidade':
            txt_bulbo_dir += " " + texto_tortuosidade_report(na)
        else:
            suf_h = "com alteração hemodinâmica local" if na.get('hemo') else "sem repercussão hemodinâmica"
            txt_bulbo_dir += f" Identifica-se espessamento parietal sugestivo de processo inflamatório, medindo {na.get('medida', 0)} mm, {suf_h}."
        tem_achado_bulbo_dir = True
    for inc in [x for x in st.session_state.lesoes_incipientes if "bulbo carotídeo direito" in x['vaso'].lower()]:
        txt_bulbo_dir += f" Identifica-se espessamento focal da camada médio-intimal no {inc['localizacao']}, medindo {inc['espessura']} mm."
        tem_achado_bulbo_dir = True
    for p in [x for x in st.session_state.lista_placas if "bulbo carotídeo direito" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        comp = p['composicao_texto'].lower().removeprefix("placa ")
        txt_bulbo_dir += f" Apresenta na parede uma placa de ateroma {comp}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
        tem_achado_bulbo_dir = True
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Direito" and "bulbo" in c['topografia']:
            txt_bulbo_dir += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."
            tem_achado_bulbo_dir = True
    if not tem_achado_bulbo_dir:
        txt_bulbo_dir += " Sem evidências de placas ou alterações estruturais."

    placas_aci_dir = [p for p in st.session_state.lista_placas if "interna direita" in p['vaso'].lower()]
    nas_aci_dir = [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Direito" and "interna" in x['vaso'].lower()]
    calcs_aci_dir = [c for c in st.session_state.calcificacoes_isoladas if c['lado'] == "Direito" and "interna" in c['topografia']]
    incs_aci_dir = [x for x in st.session_state.lesoes_incipientes if "interna direita" in x['vaso'].lower()]

    if estado_aci_dir in ["Oclusão", "Suboclusão"]:
        if placas_aci_dir:
            p = placas_aci_dir[0]
            suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
            comp_aci_dir = p['composicao_texto'].lower().removeprefix("placa ")
            txt_aci_dir = f"Artéria carótida interna direita apresentando na parede uma placa de ateroma {comp_aci_dir}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, {sufixo_hemo_aci_dir}"
        else:
            txt_aci_dir = f"Artéria carótida interna direita {sufixo_hemo_aci_dir}"
    elif placas_aci_dir:
        p = placas_aci_dir[0]
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        comp_aci_dir = p['composicao_texto'].lower().removeprefix("placa ")
        if vps_aci_dir == 0.0 and not tort_hemo_aci_dir:
            txt_aci_dir = f"Artéria carótida interna direita pérvia. Apresenta na parede uma placa de ateroma {comp_aci_dir}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, sem repercussão hemodinâmica. Mantém {txt_fluxo_normal_aci}"
        else:
            txt_aci_dir = f"Artéria carótida interna direita pérvia, apresentando na parede uma placa de ateroma {comp_aci_dir}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, {sufixo_hemo_aci_dir}"
    else:
        txt_aci_dir = f"Artéria carótida interna direita pérvia, {sufixo_hemo_aci_dir}"
    for inc in incs_aci_dir:
        txt_aci_dir += f" Identifica-se espessamento focal da camada médio-intimal no {inc['localizacao']}, medindo {inc['espessura']} mm."
    for na in nas_aci_dir:
        if na.get('categoria') == 'tortuosidade':
            txt_aci_dir += " " + texto_tortuosidade_report(na)
        else:
            suf_h = "com alteração hemodinâmica local" if na.get('hemo') else "sem repercussão hemodinâmica"
            txt_aci_dir += f" Identifica-se espessamento parietal sugestivo de processo inflamatório, medindo {na.get('medida', 0)} mm, {suf_h}."
    for c in calcs_aci_dir:
        txt_aci_dir += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."

    # --- LADO ESQUERDO ---
    
    txt_comum_esq = f"Artéria carótida comum esquerda pérvia, com diâmetro e trajeto conservados, apresentando fluxo bifásico anterógrado de baixa resistência. Espessura do complexo médio-intimal: {cmi_esq} mm."
    for na in [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Esquerdo" and "comum" in x['vaso'].lower()]:
        if na.get('categoria') == 'tortuosidade':
            txt_comum_esq += " " + texto_tortuosidade_report(na)
        else:
            suf_h = "com alteração hemodinâmica local" if na.get('hemo') else "sem repercussão hemodinâmica"
            txt_comum_esq += f" Identifica-se espessamento parietal sugestivo de processo inflamatório, medindo {na.get('medida', 0)} mm, {suf_h}."
    for inc in [x for x in st.session_state.lesoes_incipientes if "carótida comum esquerda" in x['vaso'].lower()]:
        txt_comum_esq += f" Identifica-se espessamento focal da camada médio-intimal no {inc['localizacao']}, medindo {inc['espessura']} mm."
    for p in [x for x in st.session_state.lista_placas if "carótida comum esquerda" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        comp = p['composicao_texto'].lower().removeprefix("placa ")
        txt_comum_esq += f" Identifica-se na parede uma placa de ateroma {comp}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Esquerdo" and "comum" in c['topografia']:
            txt_comum_esq += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."

    txt_bulbo_esq = "Bulbo carotídeo esquerdo pérvio, com diâmetro e trajeto conservados."
    tem_achado_bulbo_esq = False
    for na in [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Esquerdo" and "bulbo" in x['vaso'].lower()]:
        if na.get('categoria') == 'tortuosidade':
            txt_bulbo_esq += " " + texto_tortuosidade_report(na)
        else:
            suf_h = "com alteração hemodinâmica local" if na.get('hemo') else "sem repercussão hemodinâmica"
            txt_bulbo_esq += f" Identifica-se espessamento parietal sugestivo de processo inflamatório, medindo {na.get('medida', 0)} mm, {suf_h}."
        tem_achado_bulbo_esq = True
    for inc in [x for x in st.session_state.lesoes_incipientes if "bulbo carotídeo esquerdo" in x['vaso'].lower()]:
        txt_bulbo_esq += f" Identifica-se espessamento focal da camada médio-intimal no {inc['localizacao']}, medindo {inc['espessura']} mm."
        tem_achado_bulbo_esq = True
    for p in [x for x in st.session_state.lista_placas if "bulbo carotídeo esquerdo" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        comp = p['composicao_texto'].lower().removeprefix("placa ")
        txt_bulbo_esq += f" Apresenta na parede uma placa de ateroma {comp}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
        tem_achado_bulbo_esq = True
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Esquerdo" and "bulbo" in c['topografia']:
            txt_bulbo_esq += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."
            tem_achado_bulbo_esq = True
    if not tem_achado_bulbo_esq:
        txt_bulbo_esq += " Sem evidências de placas ou alterações estruturais."

    placas_aci_esq = [p for p in st.session_state.lista_placas if "interna esquerda" in p['vaso'].lower()]
    nas_aci_esq = [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Esquerdo" and "interna" in x['vaso'].lower()]
    calcs_aci_esq = [c for c in st.session_state.calcificacoes_isoladas if c['lado'] == "Esquerdo" and "interna" in c['topografia']]
    incs_aci_esq = [x for x in st.session_state.lesoes_incipientes if "interna esquerda" in x['vaso'].lower()]

    if estado_aci_esq in ["Oclusão", "Suboclusão"]:
        if placas_aci_esq:
            p = placas_aci_esq[0]
            suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
            comp_aci_esq = p['composicao_texto'].lower().removeprefix("placa ")
            txt_aci_esq = f"Artéria carótida interna esquerda apresentando na parede uma placa de ateroma {comp_aci_esq}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, {sufixo_hemo_aci_esq}"
        else:
            txt_aci_esq = f"Artéria carótida interna esquerda {sufixo_hemo_aci_esq}"
    elif placas_aci_esq:
        p = placas_aci_esq[0]
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        comp_aci_esq = p['composicao_texto'].lower().removeprefix("placa ")
        if vps_aci_esq == 0.0 and not tort_hemo_aci_esq:
            txt_aci_esq = f"Artéria carótida interna esquerda pérvia. Apresenta na parede uma placa de ateroma {comp_aci_esq}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, sem repercussão hemodinâmica. Mantém {txt_fluxo_normal_aci}"
        else:
            txt_aci_esq = f"Artéria carótida interna esquerda pérvia, apresentando na parede uma placa de ateroma {comp_aci_esq}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, {sufixo_hemo_aci_esq}"
    else:
        txt_aci_esq = f"Artéria carótida interna esquerda pérvia, {sufixo_hemo_aci_esq}"
    for inc in incs_aci_esq:
        txt_aci_esq += f" Identifica-se espessamento focal da camada médio-intimal no {inc['localizacao']}, medindo {inc['espessura']} mm."
    for na in nas_aci_esq:
        if na.get('categoria') == 'tortuosidade':
            txt_aci_esq += " " + texto_tortuosidade_report(na)
        else:
            suf_h = "com alteração hemodinâmica local" if na.get('hemo') else "sem repercussão hemodinâmica"
            txt_aci_esq += f" Identifica-se espessamento parietal sugestivo de processo inflamatório, medindo {na.get('medida', 0)} mm, {suf_h}."
    for c in calcs_aci_esq:
        txt_aci_esq += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."

    # --- COLETA DA IMPRESSÃO DIAGNÓSTICA ---
    impressao_linhas = []
    tem_achado = False
    cmi_alterado = (cmi_dir > 0.9 or cmi_esq > 0.9)
    tem_placa = len(st.session_state.lista_placas) > 0
    maior_que_plaque_rads_2 = any(p['espessura'] > 2.0 for p in st.session_state.lista_placas)

    if cmi_dir > 0.9 and cmi_esq > 0.9:
        impressao_linhas.append("– Espessamento do complexo médio-intimal bilateralmente.")
        tem_achado = True
    elif cmi_dir > 0.9:
        impressao_linhas.append("– Espessamento do complexo médio-intimal à direita.")
        tem_achado = True
    elif cmi_esq > 0.9:
        impressao_linhas.append("– Espessamento do complexo médio-intimal à esquerda.")
        tem_achado = True

    for na in st.session_state.lesoes_nao_ateromatosas:
        tem_achado = True
        if na.get('categoria') == 'tortuosidade':
            impressao_linhas.append(impressao_tortuosidade_text(na))
        else:
            impressao_linhas.append(f"– Alterações sugestivas de processo inflamatório (vasculite) na {na.get('vaso','').lower()} {na.get('lado','').lower()}.")

    for inc in st.session_state.lesoes_incipientes:
        impressao_linhas.append(f"– Alteração ateromatosa incipiente na {inc['vaso'].lower()} ({inc['localizacao']}), medindo {inc['espessura']} mm.")
        tem_achado = True

    for p in st.session_state.lista_placas:
        tem_achado = True
        v_nome = p['vaso'].lower()
        if "interna" in v_nome:
            status_hemo = status_aci_dir_limpo if "direita" in v_nome else status_aci_esq_limpo
        else:
            status_hemo = "Velocidade Não Informada" if not p['culpada_hemo'] else "Estenose de 50-59%"

        pr_str = p.get('plaque_rads') or ""
        try:
            pr_num = int(pr_str.split("Plaque-RADS")[1].strip()) if "Plaque-RADS" in pr_str else 0
        except (ValueError, IndexError):
            pr_num = 0
        risk_note = f", com achados sugestivos de risco de instabilidade plaqueária ({pr_str})" if pr_num >= 4 else ""

        if "Oclusão" in status_hemo:
            impressao_linhas.append(f"– Placa de ateroma determinando oclusão completa da {v_nome}{risk_note}.")
        elif "Suboclusão" in status_hemo:
            impressao_linhas.append(f"– Placa de ateroma determinando suboclusão da {v_nome}{risk_note}.")
        elif "Normal" in status_hemo or "< 50%" in status_hemo or "Velocidade Não Informada" in status_hemo or "Tortuosidade" in status_hemo:
            impressao_linhas.append(f"– Placa de ateroma discreta na {v_nome}{risk_note}.")
        else:
            pct = status_hemo.replace('Estenose de ', '').replace('Estenose ', '')
            impressao_linhas.append(f"– Placa de ateroma determinando estenose de {pct} na {v_nome}{risk_note}.")

    has_bulbo_dir = any(c['lado'] == "Direito" and c['topografia'] == "bulbo carotídeo" for c in st.session_state.calcificacoes_isoladas)
    has_bulbo_esq = any(c['lado'] == "Esquerdo" and c['topografia'] == "bulbo carotídeo" for c in st.session_state.calcificacoes_isoladas)
    if has_bulbo_dir and has_bulbo_esq:
        impressao_linhas.append("– Calcificações parietais isoladas sem repercussão hemodinâmica nos bulbos carotídeos bilateralmente.")
        tem_achado = True
        for c in [x for x in st.session_state.calcificacoes_isoladas if x['topografia'] != "bulbo carotídeo"]:
            prep = "na" if c['topografia'].startswith("artéria") else "no"
            impressao_linhas.append(f"– Calcificações parietais isoladas sem repercussão hemodinâmica {prep} {c['topografia']} {c['lado'].lower()}.")
            tem_achado = True
    else:
        for c in st.session_state.calcificacoes_isoladas:
            prep = "na" if c['topografia'].startswith("artéria") else "no"
            impressao_linhas.append(f"– Calcificações parietais isoladas sem repercussão hemodinâmica {prep} {c['topografia']} {c['lado'].lower()}.")
            tem_achado = True

    if not tem_achado:
        impressao_linhas.append("– Artérias carótidas e vertebrais pérvias, com trajetos e padrões de fluxo normais, dentro dos limites da normalidade.")

    # --- COLETA DAS OBSERVAÇÕES ---
    obs_ativas = []
    if incluir_observacoes:
        if cmi_alterado:
            obs_ativas.append(
                "\"O espessamento do complexo médio-intimal carotídeo é considerado marcador de aterosclerose subclínica "
                "e associa-se a aumento do risco de eventos cardiovasculares, devendo sua interpretação ser integrada ao "
                "contexto clínico e aos demais fatores de risco do paciente.\" Referências: Mannheim Carotid Intima-Media "
                "Thickness Consensus (2004–2006); ESC/EAS Guidelines for the Management of Dyslipidaemias (2021)."
            )
        if st.session_state.lesoes_incipientes:
            obs_ativas.append(
                "\"A presença de lesão ateromatosa incipiente, embora sem repercussão hemodinâmica significativa, "
                "constitui marcador de aterosclerose subclínica e possui relevância na estratificação do risco "
                "cardiovascular global, devendo ser considerada em conjunto com os demais fatores de risco e "
                "achados clínicos do paciente.\""
            )
        if tem_placa:
            obs_ativas.append(
                f"\"A presença de placa aterosclerótica carotídea, independentemente do grau de estenose, caracteriza "
                f"aterosclerose subclínica e constitui fator agravante de risco cardiovascular, devendo ser considerada "
                f"na estratificação global do risco cardiovascular, conforme a Diretriz Brasileira de Dislipidemias e "
                f"Prevenção da Aterosclerose – {ano_dislipidemia}.\""
            )
        if maior_que_plaque_rads_2 or any(p['plaque_rads'] is not None for p in st.session_state.lista_placas):
            obs_ativas.append(
                "\"A classificação Plaque-RADS padroniza a caracterização ultrassonográfica das placas carotídeas, "
                "incorporando aspectos morfológicos relacionados à vulnerabilidade da placa e fornecendo informação "
                "complementar ao grau de estenose na estratificação do risco de eventos cerebrovasculares.\" Referências: "
                "Plaque-RADS™ Consensus Statement (2023); recomendações da Society of Radiologists in Ultrasound para "
                "avaliação ultrassonográfica da doença carotídea."
            )
        _placas_alerta_obs = [p for p in st.session_state.lista_placas
                              if p['composicao_texto'] != "Placa calcificada" and p['espessura'] >= 3.0]
        if _placas_alerta_obs:
            obs_ativas.append(
                "\"Identifica-se placa aterosclerótica não completamente calcificada com espessura ≥ 3 mm, "
                "característica associada a maior risco de instabilidade e de eventos cerebrovasculares isquêmicos. "
                "Recomenda-se otimização do tratamento das dislipidemias e dos demais fatores de risco cardiovascular, "
                "além de acompanhamento ultrassonográfico periódico.\""
            )

    # Monta texto de obs para marcador
    obs_texto_linhas = []
    for i, obs in enumerate(obs_ativas):
        prefix = "– Observação: " if len(obs_ativas) == 1 else f"– Observação {i+1}: "
        obs_texto_linhas.append(prefix + obs)

    # --- DICIONÁRIO DE MARCADORES ---
    ace_dir_frase = f"Artéria carótida externa direita {(ace_dir if ace_dir.endswith('.') else ace_dir + '.').lower()}"
    ace_esq_frase = f"Artéria carótida externa esquerda {(ace_esq if ace_esq.endswith('.') else ace_esq + '.').lower()}"
    marcadores = {
        "{{nome}}": nome.strip(),
        "{{tecnica}}": texto_tecnica_final,
        "{{cmi_dir}}": str(cmi_dir),
        "{{cmi_esq}}": str(cmi_esq),
        "{{acc_dir}}": txt_comum_dir,
        "{{bulbo_dir}}": txt_bulbo_dir,
        "{{aci_dir}}": txt_aci_dir,
        "{{ace_dir}}": ace_dir_frase,
        "{{vertebral_dir}}": texto_vert_dir,
        "{{acc_esq}}": txt_comum_esq,
        "{{bulbo_esq}}": txt_bulbo_esq,
        "{{aci_esq}}": txt_aci_esq,
        "{{ace_esq}}": ace_esq_frase,
        "{{vertebral_esq}}": texto_vert_esq,
        "{{impressao}}": "\n".join(impressao_linhas),
        "{{observacoes}}": "\n".join(obs_texto_linhas),
        "{{medico}}": nome_medico,
        "{{crm}}": f"CRM-{crm_uf} {crm_medico}" if crm_medico else "",
        "{{rqe}}": f"RQE {rqe_medico}" if rqe_medico else "",
    }

    # --- GERAÇÃO DO DOCUMENTO ---
    if template_usuario.strip():
        # MODO TEMPLATE: substituição de marcadores
        if nome_clinica:
            p_cl = doc.add_paragraph()
            p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cl = p_cl.add_run(nome_clinica.upper())
            r_cl.bold = True
            r_cl.font.size = Pt(tamanho_fonte + 2)
        for linha_tmpl in template_usuario.split("\n"):
            texto_sub = linha_tmpl
            for marcador, valor in marcadores.items():
                texto_sub = texto_sub.replace(marcador, valor)
            for sub_linha in texto_sub.split("\n"):
                if sub_linha.strip():
                    adicionar_texto_esquerda(sub_linha)
    else:
        # MODO PADRÃO: layout estruturado existente
        if nome_clinica:
            p_cl = doc.add_paragraph()
            p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cl = p_cl.add_run(nome_clinica.upper())
            r_cl.bold = True
            r_cl.font.size = Pt(tamanho_fonte + 2)
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

        adicionar_titulo('DUPLEX SCAN DAS ARTÉRIAS CARÓTIDAS E VERTEBRAIS')
        if nome.strip():
            adicionar_texto_esquerda(f"Paciente: {nome.strip()}")
        adicionar_texto_esquerda(texto_tecnica_final, bold_prefix="Técnica: ")
        adicionar_subtitulo('RELATÓRIO')
        adicionar_subtitulo('LADO DIREITO')
        adicionar_texto_esquerda(txt_comum_dir)
        adicionar_texto_esquerda(txt_bulbo_dir)
        adicionar_texto_esquerda(txt_aci_dir)
        adicionar_texto_esquerda(ace_dir_frase)
        adicionar_texto_esquerda(texto_vert_dir)
        adicionar_subtitulo('LADO ESQUERDO')
        adicionar_texto_esquerda(txt_comum_esq)
        adicionar_texto_esquerda(txt_bulbo_esq)
        adicionar_texto_esquerda(txt_aci_esq)
        adicionar_texto_esquerda(ace_esq_frase)
        adicionar_texto_esquerda(texto_vert_esq)

        if quebrar_pagina_diag:
            doc.add_page_break()
        adicionar_subtitulo('IMPRESSÃO DIAGNÓSTICA')
        for linha in impressao_linhas:
            adicionar_texto_esquerda(linha)

        total_obs = len(obs_ativas)
        if total_obs == 1:
            adicionar_texto_esquerda(obs_ativas[0], bold_prefix="– Observação: ", italico=True, fonte_menor=True)
        elif total_obs > 1:
            for i, obs in enumerate(obs_ativas):
                adicionar_texto_esquerda(obs, bold_prefix=f"– Observação {i+1}: ", italico=True, fonte_menor=True)

    # Bloco de Assinatura (comum a ambos os modos)
    if nome_medico or crm_medico:
        doc.add_paragraph().paragraph_format.space_before = Pt(36)
        p_assinatura = doc.add_paragraph()
        p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if nome_medico:
            run_n = p_assinatura.add_run(f"{nome_medico}\n")
            run_n.bold = True
        if crm_medico:
            p_assinatura.add_run(f"CRM-{crm_uf} {crm_medico}\n")
        if rqe_medico:
            p_assinatura.add_run(f"RQE {rqe_medico}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("Laudo integrado gerado com sucesso!")

    # Visualização do laudo na própria página
    texto_visualizacao = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    if modo_saida in ["Somente Visualização", "Visualização + DOCX"]:
        st.markdown("## 👁️ Visualização do Laudo")
        st.text_area("Laudo Gerado", value=texto_visualizacao, height=700)
    st.markdown("### 🔊 Leitura em Áudio do Laudo")
    render_audio_player(texto_visualizacao, key="arterial")

    # Download DOCX
    if modo_saida in ["Somente DOCX", "Visualização + DOCX"]:
        st.download_button(
            label="📥 Baixar Laudo Formatado (.docx)",
            data=buffer,
            file_name="Laudo_Vascular_Completo.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
