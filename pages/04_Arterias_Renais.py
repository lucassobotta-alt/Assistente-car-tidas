# pages/03_Arterias_Renais.py
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

st.title("🫘 Assistente de Laudos: Duplex Scan de Aorta e Artérias Renais")

# --- SIDEBAR ---
with st.sidebar:
    st.markdown("## ⚙️ Painel de Controle")

    st.markdown("### 📚 Referências Científicas")
    st.markdown(
        "📄 <a href='https://www.arquivosonline.com.br/2019/113/4/' target='_blank'>"
        "SBC – Vascular Ultrasound Statement (Arq Bras Cardiol, 2019)</a><br>"
        "📄 <a href='https://www.escardio.org/Guidelines/Clinical-Practice-Guidelines/Peripheral-Arterial-and-Aortic-Diseases' target='_blank'>"
        "ESC 2024 Guidelines – Peripheral Arterial and Aortic Diseases</a>",
        unsafe_allow_html=True
    )

    st.markdown("---")
    st.markdown("### 📝 Formatação Externa (.docx)")
    fonte_doc = st.selectbox("Família da Fonte:", ["Arial", "Calibri", "Times New Roman"], key="ar_fonte")
    tamanho_fonte = st.slider("Tamanho do Texto (pt):", 10, 14, 11, key="ar_tam")
    espacamento_linhas = st.slider("Espaçamento entre Linhas:", 1.0, 1.5, 1.15, step=0.05, key="ar_esp")
    quebrar_pagina_diag = st.toggle("Separar Impressão em Nova Página", value=False, key="ar_qpd")
    modo_saida = st.radio(
        "Modo de saída:",
        ["Somente DOCX", "Somente Visualização", "Visualização + DOCX"],
        index=2, key="ar_saida"
    )

    st.markdown("---")
    st.markdown("### ✍️ Identidade & Assinatura")
    nome_clinica = st.text_input("Cabeçalho / Nome da Clínica:", placeholder="Ex: Instituto de Diagnóstico Vascular", key="ar_clinica")
    nome_medico = st.text_input("Nome do Médico:", "Lucas Santos Guimarães", key="ar_medico")
    colcrm1, colcrm2 = st.columns([2, 1])
    with colcrm1:
        crm_medico = st.text_input("CRM:", "4061", key="ar_crm")
    with colcrm2:
        crm_uf = st.selectbox("UF", ["AC","AL","AP","AM","BA","CE","DF","ES","GO","MA","MT","MS","MG","PA","PB","PR","PE","PI","RJ","RN","RS","RO","RR","SC","SP","SE","TO"], index=25, key="ar_uf")
    rqe_medico = st.text_input("RQE:", "", key="ar_rqe")
    incluir_assinatura = st.toggle("Incluir assinatura / carimbo no laudo", value=True, key="ar_assin")

    st.markdown("---")
    st.markdown("### 🔬 Avaliação Complementar")
    sugerir_complementar = st.toggle("Sugerir avaliação complementar no laudo", value=False, key="ar_sug_comp")
    metodos_complementares = []
    if sugerir_complementar:
        metodos_complementares = st.multiselect(
            "Método(s) sugerido(s):",
            ["Angiotomografia (AngioTC)", "Angioressonância Magnética (AngioRM)", "Arteriografia digital"],
            default=["Angiotomografia (AngioTC)"],
            key="ar_sug_metodos"
        )

    st.markdown("---")
    _AR_SIDEBAR_KEYS = {
        "ar_fonte", "ar_tam", "ar_esp", "ar_qpd", "ar_saida",
        "ar_clinica", "ar_medico", "ar_crm", "ar_uf", "ar_rqe",
        "ar_assin", "ar_rodape_tog", "ar_rodape_url",
        "ar_sug_comp", "ar_sug_metodos",
    }
    if st.button("🔄 Resetar Todos os Parâmetros", use_container_width=True, type="secondary"):
        for _k in [k for k in st.session_state if k not in _AR_SIDEBAR_KEYS]:
            del st.session_state[_k]
        st.toast("🔄 Todos os dados clínicos foram limpos!")
        st.rerun()

    st.markdown("---")
    st.markdown("### 🔗 Rodapé do Documento")
    incluir_rodape_link = st.toggle("Incluir nota de rodapé com link do sistema", value=False, key="ar_rodape_tog")
    rodape_url = ""
    if incluir_rodape_link:
        rodape_url = st.text_input("URL do sistema:", placeholder="Ex: https://seu-app.streamlit.app", key="ar_rodape_url")

# --- PACIENTE ---
nome_paciente = st.text_input("Nome do Paciente:", "", key="ar_paciente")

modo_template = st.radio(
    "Template de avaliação:",
    ["Duplex Scan de Aorta e Artérias Renais", "Avaliação de Rim Transplantado"],
    horizontal=True,
    key="ar_modo_template"
)

# =========================================================
# MODO: RIM TRANSPLANTADO
# =========================================================
if modo_template == "Avaliação de Rim Transplantado":

    _ECOTEXTURA_OPTS_TX = [
        "normais",
        "ecotextura homogênea com diferenciação corticomedular preservada",
        "ecotextura aumentada com diferenciação corticomedular reduzida",
        "ecotextura heterogênea com diferenciação corticomedular reduzida",
        "ecotextura heterogênea com diferenciação corticomedular abolida",
    ]

    st.markdown("---")
    st.markdown("### 1. Localização do Rim Transplantado")
    localizacao_tx = st.radio(
        "Rim transplantado localizado em:",
        ["Fossa ilíaca direita", "Fossa ilíaca esquerda"],
        horizontal=True,
        key="tx_loc"
    )

    st.markdown("---")
    st.markdown("### 2. Avaliação Morfológica do Rim Transplantado")
    col_tx1, col_tx2 = st.columns(2)
    with col_tx1:
        comp_tx = st.text_input("Comprimento (cm):", "11.0", key="tx_comp")
    with col_tx2:
        cortical_tx = st.text_input("Espessura cortical (cm):", "1.2", key="tx_cortical")

    ecotextura_tx = st.selectbox(
        "Ecotextura e diferenciação corticomedular:",
        _ECOTEXTURA_OPTS_TX,
        key="tx_eco"
    )

    c_tx1, c_tx2 = st.columns(2)
    with c_tx1:
        calculos_tx = st.toggle("Cálculos identificados?", key="tx_calc")
        calc_desc_tx = ""
        if calculos_tx:
            calc_desc_tx = st.text_input("Descrição:", "cálculo de até ___ mm no polo inferior", key="tx_calc_desc")
    with c_tx2:
        hidronefrose_tx = st.toggle("Hidronefrose?", key="tx_hidro")
        hidro_desc_tx = ""
        if hidronefrose_tx:
            hidro_desc_tx = st.selectbox(
                "Grau:",
                ["leve (grau I)", "moderada (grau II)", "acentuada (grau III)"],
                key="tx_hidro_grau"
            )

    try:
        _comp_tx_f = float(comp_tx.replace(",", "."))
        if _comp_tx_f < 9.0:
            st.warning(f"⚠️ Rim transplantado com comprimento {comp_tx} cm < 9,0 cm — considerar redução volumétrica.")
    except ValueError:
        pass

    st.markdown("**Coleções Perinéfricas**")
    colecao_tx = st.toggle("Coleção perinéfrica identificada?", key="tx_colecao")
    colecao_tipo_tx = ""
    colecao_desc_tx = ""
    if colecao_tx:
        colecao_tipo_tx = st.selectbox(
            "Tipo de coleção:",
            ["hematoma", "seroma", "linfocele", "abscesso", "urinoma"],
            key="tx_col_tipo"
        )
        colecao_desc_tx = st.text_input(
            "Dimensões e localização:",
            "coleção de ___ × ___ cm em topografia ___",
            key="tx_col_desc"
        )

    st.markdown("---")
    st.markdown("### 3. Anastomose Arterial")
    tipo_anast_tx = st.selectbox(
        "Tipo de anastomose arterial:",
        [
            "término-lateral com artéria ilíaca externa",
            "término-lateral com artéria ilíaca comum",
            "término-terminal com artéria ilíaca interna",
        ],
        key="tx_anast_tipo"
    )

    col_a1, col_a2, col_a3 = st.columns(3)
    with col_a1:
        vps_anast_tx = st.text_input("VPS na anastomose (cm/s):", "120", key="tx_anast_vps")
    with col_a2:
        vdf_anast_tx = st.text_input("VDF na anastomose (cm/s):", "40", key="tx_anast_vdf")
    with col_a3:
        ir_anast_tx = st.text_input("IR na anastomose:", "0.65", key="tx_anast_ir")

    estenose_anast_tx = st.toggle("Turbulência / estenose na anastomose?", key="tx_anast_esten")
    estenose_anast_desc_tx = ""
    if estenose_anast_tx:
        estenose_anast_desc_tx = st.text_input(
            "Descrição da alteração:",
            "aceleração focal com turbulência pós-estenótica na anastomose",
            key="tx_anast_esten_desc"
        )

    st.markdown("---")
    st.markdown("### 4. Artéria Principal do Rim Transplantado")
    col_b1, col_b2, col_b3, col_b4 = st.columns(4)
    with col_b1:
        vps_art_tx = st.text_input("VPS (cm/s):", "120", key="tx_art_vps")
    with col_b2:
        vdf_art_tx = st.text_input("VDF (cm/s):", "40", key="tx_art_vdf")
    with col_b3:
        ir_art_tx = st.text_input("IR:", "0.65", key="tx_art_ir")
    with col_b4:
        ta_art_tx = st.text_input("Tempo de aceleração (ms):", "50", key="tx_art_ta")

    st.markdown("---")
    st.markdown("### 5. Fluxo Intrarrenal")
    _FLUXO_INTRA_TX_OPTS = [
        "baixa resistência",
        "alta resistência",
        "padrão tardus-parvus",
        "ausente (oclusão)",
    ]
    col_fi1, col_fi2, col_fi3, col_fi4 = st.columns(4)
    with col_fi1:
        fluxo_intra_tx = st.selectbox("Padrão:", _FLUXO_INTRA_TX_OPTS, key="tx_fi_pad")
    with col_fi2:
        vps_intra_tx = st.text_input("VPS intrarrenal (cm/s):", "25", key="tx_fi_vps")
    with col_fi3:
        ir_intra_tx = st.text_input("IR intrarrenal:", "0.65", key="tx_fi_ir")
    with col_fi4:
        ta_intra_tx = st.text_input("Tempo de aceleração (ms):", "70", key="tx_fi_ta")

    def _safe_float_tx(s):
        try:
            return float(str(s).replace(",", ".").strip())
        except (ValueError, TypeError):
            return None

    _ir_intra_f_tx = _safe_float_tx(ir_intra_tx)
    if _ir_intra_f_tx is not None:
        if _ir_intra_f_tx < 0.70:
            st.success(f"🟢 IR intrarrenal {ir_intra_tx} — dentro da normalidade (< 0,70).")
        elif _ir_intra_f_tx <= 0.80:
            st.warning(f"🟡 IR intrarrenal {ir_intra_tx} — limítrofe (0,70–0,80). Correlacionar clinicamente.")
        else:
            st.error(f"🔴 IR intrarrenal {ir_intra_tx} — elevado (> 0,80). Considerar rejeição aguda, necrose tubular aguda ou obstrução.")

    st.markdown("---")
    st.markdown("### 6. Anastomose Venosa")
    veia_anast_tx = st.selectbox(
        "Anastomose venosa:",
        [
            "pérvia, com fluxo venoso de padrão habitual",
            "com fluxo venoso de padrão pulsátil (elevação de pressão venosa)",
            "não identificada ao método (limitação técnica)",
        ],
        key="tx_veia"
    )

    st.markdown("---")
    st.markdown("### 7. Observação Técnica (Opcional)")
    incluir_obs_tx = st.toggle("Incluir nota sobre limitações técnicas?", value=False, key="tx_obs_tog")
    obs_tx = ""
    if incluir_obs_tx:
        obs_tx = st.text_area("Texto da observação técnica:", key="tx_obs_txt", height=80)

    st.markdown("---")

    def construir_laudo_transplante():
        doc = Document()
        doc.styles['Normal'].font.name = fonte_doc
        doc.styles['Normal'].font.size = Pt(tamanho_fonte)

        def add_p(text, bold_pre=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=4, bullet=False, italic=False):
            p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
            p.alignment = align
            p.paragraph_format.line_spacing = espacamento_linhas
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            if bold_pre:
                r_p = p.add_run(bold_pre)
                r_p.bold = True
            r = p.add_run(text)
            if italic:
                r.italic = True

        if nome_clinica.strip():
            p_cl = doc.add_paragraph()
            p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cl = p_cl.add_run(nome_clinica.upper())
            r_cl.bold = True
            r_cl.font.name = fonte_doc
            r_cl.font.size = Pt(tamanho_fonte + 2)
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        add_p("DE RIM TRANSPLANTADO", bold_pre="DUPLEX SCAN ", space_after=12)
        if nome_paciente.strip():
            add_p(f" {nome_paciente}", bold_pre="Paciente:")

        add_p("TÉCNICA", space_before=12, space_after=6)
        add_p(
            "Exame realizado com transdutor convexo e linear multifrequencial, utilizando recursos de modo B, "
            "Doppler colorido e Doppler espectral, com avaliação morfológica e hemodinâmica do rim transplantado, "
            "anastomoses arterial e venosa, fluxo intrarrenal e estruturas perinéfricas.",
            space_after=12
        )

        add_p("RELATÓRIO", space_before=6, space_after=8)

        add_p("MORFOLOGIA DO RIM TRANSPLANTADO", space_after=4)
        _calc_txt_tx = (
            f"Identificado(s) {calc_desc_tx}."
            if calculos_tx and calc_desc_tx
            else "Não foram observados cálculos com mais de 5 mm."
        )
        _hidro_txt_tx = (
            f"Identificada hidronefrose {hidro_desc_tx}."
            if hidronefrose_tx and hidro_desc_tx
            else "Ausência de hidronefrose."
        )
        _eco_tx = ecotextura_tx if ecotextura_tx != "normais" else "normais"
        add_p(
            f"Rim transplantado localizado em {localizacao_tx.lower()}, medindo {comp_tx} cm, com espessura "
            f"cortical de {cortical_tx} cm, apresentando ecotextura e diferenciação corticomedular {_eco_tx}. "
            f"{_calc_txt_tx} {_hidro_txt_tx}"
        )
        if colecao_tx and colecao_tipo_tx:
            add_p(f"Identificada {colecao_tipo_tx} perinéfrica: {colecao_desc_tx}.")
        else:
            add_p("Não se identificam coleções líquidas perinéfricas significativas.")

        add_p("ANASTOMOSE ARTERIAL", space_before=10, space_after=4)
        add_p(
            f"Anastomose arterial {tipo_anast_tx}, com VPS de {vps_anast_tx} cm/s, "
            f"VDF de {vdf_anast_tx} cm/s e IR de {ir_anast_tx}."
        )
        if estenose_anast_tx and estenose_anast_desc_tx:
            add_p(f"Observa-se {estenose_anast_desc_tx}.")
        else:
            add_p("Não se observam acelerações focais significativas ou turbulência pós-estenótica na anastomose.")

        add_p("ARTÉRIA PRINCIPAL DO RIM TRANSPLANTADO", space_before=10, space_after=4)
        _ta_art_tx_f = _safe_float_tx(ta_art_tx)
        if _ta_art_tx_f is not None and _ta_art_tx_f > 100:
            _tp_txt = f", com tempo de aceleração de {ta_art_tx} ms (prolongado)"
        else:
            _tp_txt = f", com tempo de aceleração de {ta_art_tx} ms"
        add_p(
            f"Artéria principal do rim transplantado pérvia, com VPS de {vps_art_tx} cm/s, "
            f"VDF de {vdf_art_tx} cm/s, IR de {ir_art_tx}{_tp_txt}."
        )

        add_p("FLUXO INTRARRENAL", space_before=10, space_after=4)
        if fluxo_intra_tx == "ausente (oclusão)":
            add_p("Fluxo intrarrenal ausente ao Doppler colorido e espectral, sugerindo oclusão arterial.")
        else:
            add_p(
                f"Fluxo intrarrenal com padrão de {fluxo_intra_tx}, VPS de {vps_intra_tx} cm/s, "
                f"IR de {ir_intra_tx} e tempo de aceleração de {ta_intra_tx} ms."
            )

        add_p("ANASTOMOSE VENOSA", space_before=10, space_after=4)
        add_p(f"Anastomose venosa {veia_anast_tx}.")

        if incluir_obs_tx and obs_tx.strip():
            add_p("OBSERVAÇÃO TÉCNICA", space_before=10, space_after=4)
            add_p(obs_tx.strip())

        if quebrar_pagina_diag:
            doc.add_page_break()
        add_p("⸻", space_after=12)
        add_p("IMPRESSÃO DIAGNÓSTICA", space_after=6)

        conclusoes_tx = []

        if hidronefrose_tx and hidro_desc_tx:
            conclusoes_tx.append(f"Hidronefrose {hidro_desc_tx} no rim transplantado.")
        if colecao_tx and colecao_tipo_tx:
            conclusoes_tx.append(f"{colecao_tipo_tx.capitalize()} perinéfrica: {colecao_desc_tx}.")
        if estenose_anast_tx:
            conclusoes_tx.append(
                "Sinais compatíveis com estenose da anastomose arterial. Correlacionar com avaliação complementar (AngioTC, AngioRM ou arteriografia)."
            )

        _ir_f_conc = _safe_float_tx(ir_intra_tx)
        if fluxo_intra_tx == "ausente (oclusão)":
            conclusoes_tx.append("Sinais de oclusão arterial do rim transplantado — trombose arterial não excluída.")
        elif fluxo_intra_tx == "padrão tardus-parvus":
            conclusoes_tx.append(
                "Padrão intrarrenal tardus-parvus — sinais indiretos de estenose proximal da artéria do enxerto renal."
            )
        elif _ir_f_conc is not None:
            if _ir_f_conc > 0.80:
                conclusoes_tx.append(
                    f"IR intrarrenal elevado ({ir_intra_tx}). Não se exclui rejeição aguda, necrose tubular aguda ou obstrução urinária. Correlacionar com quadro clínico e laboratorial."
                )
            elif _ir_f_conc > 0.70:
                conclusoes_tx.append(
                    f"IR intrarrenal limítrofe ({ir_intra_tx}). Correlacionar clinicamente e monitorar evolutivamente."
                )

        if "pulsátil" in veia_anast_tx:
            conclusoes_tx.append(
                "Fluxo venoso pulsátil na anastomose venosa — sugestivo de elevação da pressão venosa central. Correlacionar clinicamente."
            )

        if not conclusoes_tx:
            add_p(
                "Rim transplantado com avaliação morfológica e hemodinâmica dentro dos limites da normalidade ao método. "
                "Ausência de sinais de estenose arterial significativa, trombose ou rejeição ao estudo Doppler.",
                bullet=True
            )
        else:
            for c in conclusoes_tx:
                add_p(c, bullet=True)

        if sugerir_complementar and metodos_complementares:
            if len(metodos_complementares) == 1:
                _met_txt = metodos_complementares[0]
            elif len(metodos_complementares) == 2:
                _met_txt = f"{metodos_complementares[0]} e {metodos_complementares[1]}"
            else:
                _met_txt = ", ".join(metodos_complementares[:-1]) + f" e {metodos_complementares[-1]}"
            add_p(
                f"OBSERVAÇÃO: Os achados ultrassonográficos do presente exame indicam correlação com método de imagem "
                f"complementar. Sugere-se a realização de {_met_txt} para melhor caracterização anatômica e hemodinâmica "
                "do enxerto renal e planejamento terapêutico adequado.",
                space_before=10, italic=True
            )

        if incluir_assinatura and (nome_medico or crm_medico):
            doc.add_paragraph().paragraph_format.space_before = Pt(25)
            p_assin = doc.add_paragraph()
            p_assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_assin.paragraph_format.line_spacing = espacamento_linhas
            if nome_medico:
                r = p_assin.add_run(f"{nome_medico}\n")
                r.bold = True
            if crm_medico:
                p_assin.add_run(f"CRM-{crm_uf} {crm_medico}\n")
            if rqe_medico.strip():
                p_assin.add_run(f"RQE {rqe_medico}")

        if incluir_rodape_link and rodape_url.strip():
            p_rod = doc.add_paragraph()
            p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_rod.paragraph_format.space_before = Pt(30)
            r_rod = p_rod.add_run(f"Laudo gerado com suporte do sistema: {rodape_url.strip()}")
            r_rod.italic = True
            r_rod.font.size = Pt(max(tamanho_fonte - 2, 8))

        return doc

    if st.button("🚀 Gerar Laudo de Rim Transplantado", use_container_width=True):
        doc = construir_laudo_transplante()
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.success("Laudo gerado com sucesso!")
        if modo_saida in ["Somente Visualização", "Visualização + DOCX"]:
            texto_viz = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            st.markdown("## 👁️ Visualização do Laudo")
            st.text_area("Laudo Gerado", value=texto_viz, height=600)
        if modo_saida in ["Somente DOCX", "Visualização + DOCX"]:
            st.download_button(
                "📥 Baixar Laudo de Rim Transplantado (.docx)",
                buf,
                "Laudo_Rim_Transplantado.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    st.stop()

st.markdown("---")

# =========================================================
# 1. AORTA ABDOMINAL
# =========================================================
st.markdown("### 1. Aorta Abdominal")

col_ao1, col_ao2, col_ao3 = st.columns(3)
with col_ao1:
    aorta_calibre = st.text_input("Calibre máximo (cm):", "1.8", key="ao_calibre")
with col_ao2:
    aorta_vps = st.text_input("VPS aórtica (cm/s):", "70", key="ao_vps")
with col_ao3:
    aorta_placas = st.selectbox("Placas ateromatosas parietais:", ["Ausentes", "Presentes"], key="ao_placas")

aorta_aneurisma = st.toggle("Dilatação aneurismática / dissecção?", value=False, key="ao_aneu")
aorta_aneurisma_desc = ""
if aorta_aneurisma:
    aorta_aneurisma_desc = st.text_area(
        "Descrição da alteração (aneurisma/dissecção):",
        "Dilatação aneurismática da aorta abdominal infrarrenal, com diâmetro máximo de ___ cm.",
        key="ao_aneu_desc"
    )

# Alerta VPS aórtica fora do range de referência da RRA
try:
    _vps_ao_float = float(aorta_vps.replace(",", "."))
    if _vps_ao_float < 40 or _vps_ao_float > 100:
        st.warning(
            f"⚠️ **VPS aórtica fora do intervalo de referência ({aorta_vps} cm/s):** "
            "valores < 40 ou > 100 cm/s podem comprometer a validade da Razão Renal-Aórtica (RRA). "
            "Interprete a RRA com cautela neste exame."
        )
except ValueError:
    _vps_ao_float = None

st.markdown("---")

# =========================================================
# 2. AVALIAÇÃO MORFOLÓGICA RENAL
# =========================================================
st.markdown("### 2. Avaliação Morfológica Renal")

_ECOTEXTURA_OPTS = [
    "normais",
    "ecotextura homogênea com diferenciação corticomedular preservada",
    "ecotextura aumentada com diferenciação corticomedular reduzida",
    "ecotextura heterogênea com diferenciação corticomedular reduzida",
    "ecotextura heterogênea com diferenciação corticomedular abolida",
]

morfo_rins = {}
tabs_rins = st.tabs(["🔴 Rim Direito", "🔵 Rim Esquerdo"])

for ri, lado_rim in enumerate(["Direito", "Esquerdo"]):
    with tabs_rins[ri]:
        c1, c2 = st.columns(2)
        with c1:
            comp_rim = st.text_input(f"Comprimento do rim (cm):", "10.5", key=f"rim_comp_{lado_rim}")
        with c2:
            cortical = st.text_input(f"Espessura cortical (cm):", "1.0", key=f"rim_cort_{lado_rim}")

        ecotextura = st.selectbox(
            "Ecotextura e diferenciação corticomedular:",
            _ECOTEXTURA_OPTS,
            key=f"rim_eco_{lado_rim}"
        )

        c3, c4 = st.columns(2)
        with c3:
            calculos = st.toggle("Cálculos identificados?", key=f"rim_calc_{lado_rim}")
            calc_desc = ""
            if calculos:
                calc_desc = st.text_input("Descrição (tamanho, localização):", "cálculo de até ___ mm no polo inferior", key=f"rim_calc_desc_{lado_rim}")
        with c4:
            hidronefrose = st.toggle("Hidronefrose?", key=f"rim_hidro_{lado_rim}")
            hidro_desc = ""
            if hidronefrose:
                hidro_desc = st.selectbox(
                    "Grau de hidronefrose:",
                    ["leve (grau I)", "moderada (grau II)", "acentuada (grau III)"],
                    key=f"rim_hidro_grau_{lado_rim}"
                )

        morfo_rins[lado_rim] = {
            "comp": comp_rim,
            "cortical": cortical,
            "ecotextura": ecotextura,
            "calculos": calculos,
            "calc_desc": calc_desc,
            "hidronefrose": hidronefrose,
            "hidro_desc": hidro_desc,
        }

        # Alerta de rim pequeno
        try:
            _comp_float = float(comp_rim.replace(",", "."))
            if _comp_float < 9.0:
                st.warning(f"⚠️ Rim {lado_rim.lower()} com comprimento {comp_rim} cm < 9,0 cm — considerar redução volumétrica renal.")
            if _comp_float > 13.0:
                st.info(f"ℹ️ Rim {lado_rim.lower()} com comprimento {comp_rim} cm > 13,0 cm — considerar aumento renal.")
        except ValueError:
            pass

st.markdown("---")

# =========================================================
# 3. ARTÉRIAS RENAIS
# =========================================================
st.markdown("### 3. Artérias Renais")

_ID_ARTERY_OPTS = [
    "desde sua origem",
    "somente no segmento proximal",
    "somente no segmento distal",
    "não identificada (limitação técnica)",
]
_CALIBRE_TRAJ_OPTS = [
    "habituais",
    "calibre reduzido",
    "calibre aumentado",
    "trajeto irregular / tortuoso",
]
_FLUXO_INTRA_OPTS = [
    "baixa resistência",
    "alta resistência",
    "padrão tardus-parvus",
    "ausente (oclusão)",
]

def _safe_float(s):
    try:
        return float(str(s).replace(",", ".").strip())
    except (ValueError, TypeError):
        return None

def classificar_estenose(vps, vdf, rra, ta):
    """Retorna (grau_texto, cor) baseado nos critérios SBC/ESC."""
    if vps is None:
        return "Não calculado", "gray"
    if vps == 0:
        return "Oclusão (ausência de fluxo)", "red"
    if vps < 200:
        return "Normal (sem estenose significativa)", "green"
    # VPS ≥ 200
    if rra is not None and rra >= 3.5:
        if vdf is not None and vdf >= 150:
            if ta is not None and ta >= 70:
                return "Estenose ≥ 80% (grave)", "red"
            return "Estenose ≥ 60–80%", "orange"
        return "Estenose ≥ 60%", "orange"
    # VPS ≥ 200 mas RRA < 3.5
    return "Estenose < 60% (hemodinamicamente não significativa)", "blue"

dados_arterias = {}
tabs_art = st.tabs(["🔴 Artéria Renal Direita", "🔵 Artéria Renal Esquerda"])

for ai, lado_art in enumerate(["Direita", "Esquerda"]):
    with tabs_art[ai]:
        id_art = st.selectbox(
            "Identificação da artéria renal:",
            _ID_ARTERY_OPTS,
            key=f"art_id_{lado_art}"
        )

        nao_identificada = id_art == "não identificada (limitação técnica)"
        limitacao_desc = ""
        if nao_identificada:
            limitacao_desc = st.text_input(
                "Motivo da limitação técnica:",
                "meteorismo intestinal",
                key=f"art_lim_{lado_art}"
            )
            dados_arterias[lado_art] = {
                "id": id_art,
                "limitacao_desc": limitacao_desc,
                "nao_identificada": True,
            }
            st.info("ℹ️ A ausência de visualização isoladamente não deve ser interpretada como oclusão.")
        else:
            calibre_traj = st.selectbox(
                "Calibre e trajeto:",
                _CALIBRE_TRAJ_OPTS,
                key=f"art_cal_{lado_art}"
            )

            col_v1, col_v2, col_v3 = st.columns(3)
            with col_v1:
                vps_art = st.text_input("VPS (cm/s):", "120", key=f"art_vps_{lado_art}")
            with col_v2:
                vdf_art = st.text_input("VDF (cm/s):", "40", key=f"art_vdf_{lado_art}")
            with col_v3:
                # RRA auto-calculada — sempre recomputada a cada render
                _vps_f = _safe_float(vps_art)
                _rra_auto = ""
                if _vps_f is not None and _vps_ao_float and _vps_ao_float > 0:
                    _rra_auto = f"{_vps_f / _vps_ao_float:.2f}"
                _key_rra = f"art_rra_{lado_art}"
                st.session_state[_key_rra] = _rra_auto if _rra_auto else "—"
                rra_art = st.text_input(
                    "Razão renal-aórtica (RRA):",
                    key=_key_rra,
                    disabled=True,
                    help="Calculado automaticamente como VPS renal / VPS aórtica."
                )

            turbulencia = st.toggle(
                "Aceleração focal / turbulência pós-estenótica?",
                key=f"art_turb_{lado_art}"
            )
            turb_desc = ""
            if turbulencia:
                turb_desc = st.text_input(
                    "Localização da aceleração/turbulência:",
                    "na origem",
                    key=f"art_turb_desc_{lado_art}"
                )

            st.markdown("**Fluxo Intrarrenal**")
            col_i1, col_i2, col_i3, col_i4 = st.columns(4)
            with col_i1:
                fluxo_intra_padrao = st.selectbox("Padrão:", _FLUXO_INTRA_OPTS, key=f"art_fi_pad_{lado_art}")
            with col_i2:
                vps_intra = st.text_input("VPS intrarrenal (cm/s):", "30", key=f"art_fi_vps_{lado_art}")
            with col_i3:
                ir_intra = st.text_input("IR:", "0.65", key=f"art_fi_ir_{lado_art}")
            with col_i4:
                ta_intra = st.text_input("Tempo de aceleração (ms):", "50", key=f"art_fi_ta_{lado_art}")

            # --- Classificação automática ---
            _vps_f = _safe_float(vps_art)
            _vdf_f = _safe_float(vdf_art)
            _rra_f = _safe_float(rra_art)
            _ta_f = _safe_float(ta_intra)

            _grau, _cor = classificar_estenose(_vps_f, _vdf_f, _rra_f, _ta_f)

            # Sinais indiretos: artéria principal normal mas intrarrenais alteradas
            _ta_val = _ta_f or 0
            _sinais_indiretos = (
                _cor == "green" and
                (fluxo_intra_padrao == "padrão tardus-parvus" or _ta_val > 100)
            )

            _cor_map = {"green": "🟢", "blue": "🔵", "orange": "🟠", "red": "🔴", "gray": "⚪"}
            st.markdown(
                f"**{_cor_map.get(_cor, '⚪')} Classificação automática — Artéria Renal {lado_art}:** {_grau}"
            )

            if _cor in ("orange", "red"):
                st.warning(
                    f"⚠️ **Critérios compatíveis com estenose hemodinamicamente significativa — Artéria Renal {lado_art}.**  \n"
                    "Considere avaliação complementar (AngioTC, AngioRM ou arteriografia)."
                )

            if _sinais_indiretos:
                st.warning(
                    f"⚠️ **Discordância hemodinâmica — Artéria Renal {lado_art}:** "
                    "o padrão intrarrenal sugere sinais **indiretos** de estenose hemodinamicamente "
                    "significativa (tardus-parvus e/ou tempo de aceleração > 100 ms), embora as "
                    "velocidades diretas estejam dentro da normalidade.  \n"
                    "**Verifique:** a artéria renal foi completamente visualizada ao longo de todo "
                    "o seu trajeto? Houve limitação técnica que possa ter comprometido a avaliação "
                    "direta da velocidade no ponto de estenose?"
                )

            dados_arterias[lado_art] = {
                "id": id_art,
                "calibre_traj": calibre_traj,
                "vps": vps_art,
                "vdf": vdf_art,
                "rra": rra_art,
                "turbulencia": turbulencia,
                "turb_desc": turb_desc,
                "fluxo_intra_padrao": fluxo_intra_padrao,
                "vps_intra": vps_intra,
                "ir_intra": ir_intra,
                "ta_intra": ta_intra,
                "grau": _grau,
                "cor": _cor,
                "sinais_indiretos": _sinais_indiretos,
                "nao_identificada": False,
            }

st.markdown("---")

# =========================================================
# 4. OBSERVAÇÃO TÉCNICA OPCIONAL
# =========================================================
st.markdown("### 4. Observação Técnica")
incluir_obs_tecnica = st.toggle("Incluir nota sobre limitações técnicas do exame?", value=False, key="ar_obs_tog")
obs_tecnica_texto = ""
if incluir_obs_tecnica:
    _LIMITACOES_OPTS = [
        "meteorismo intestinal",
        "obesidade",
        "respiração inadequada",
        "calcificações vasculares extensas",
        "dificuldade de demonstração das origens das artérias renais",
    ]
    limitacoes_sel = st.multiselect(
        "Fatores limitantes:",
        _LIMITACOES_OPTS,
        key="ar_lim_multi"
    )
    if limitacoes_sel:
        obs_tecnica_texto = "A avaliação foi limitada por " + ", ".join(limitacoes_sel) + ". Os segmentos com avaliação prejudicada estão indicados no relatório."

st.markdown("---")

# =========================================================
# 5. GERAÇÃO DO LAUDO
# =========================================================
def construir_laudo_renais():
    doc = Document()
    doc.styles['Normal'].font.name = fonte_doc
    doc.styles['Normal'].font.size = Pt(tamanho_fonte)

    def add_p(text, bold_pre=None, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=4, bullet=False, italic=False):
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.alignment = align
        p.paragraph_format.line_spacing = espacamento_linhas
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if bold_pre:
            r_p = p.add_run(bold_pre)
            r_p.bold = True
        r = p.add_run(text)
        if italic:
            r.italic = True

    if nome_clinica.strip():
        p_cl = doc.add_paragraph()
        p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cl = p_cl.add_run(nome_clinica.upper())
        r_cl.bold = True
        r_cl.font.name = fonte_doc
        r_cl.font.size = Pt(tamanho_fonte + 2)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_p("DE AORTA E ARTÉRIAS RENAIS", bold_pre="DUPLEX SCAN ", space_after=12)
    if nome_paciente.strip():
        add_p(f" {nome_paciente}", bold_pre="Paciente:")
    add_p("TÉCNICA", space_before=12, space_after=6)
    add_p(
        "Exame realizado com transdutor convexo multifrequencial, utilizando recursos de modo B, "
        "Doppler colorido e Doppler espectral, com avaliação morfológica e hemodinâmica da aorta "
        "abdominal, das artérias renais e dos ramos intrarrenais.",
        space_after=12
    )

    add_p("RELATÓRIO", space_before=6, space_after=8)

    # — AORTA —
    add_p("AORTA ABDOMINAL", space_after=4)
    _placas_txt = "com placas ateromatosas parietais" if aorta_placas == "Presentes" else "sem placas ateromatosas parietais"
    add_p(f"Aorta abdominal com calibre de até {aorta_calibre} cm, {_placas_txt}.")
    add_p(f"VPS na aorta abdominal de {aorta_vps} cm/s.")
    if aorta_aneurisma and aorta_aneurisma_desc.strip():
        add_p(aorta_aneurisma_desc.strip())
    else:
        add_p("Não se observam dilatações aneurismáticas, dissecções ou outras alterações morfológicas significativas da aorta abdominal ao método.")

    # — MORFOLOGIA RENAL —
    add_p("AVALIAÇÃO MORFOLÓGICA RENAL", space_before=10, space_after=4)
    for lado_rim, mrm in morfo_rins.items():
        _calc_txt = (
            f"Identificado(s) {mrm['calc_desc']}."
            if mrm["calculos"] and mrm["calc_desc"]
            else "Não foram observados cálculos com mais de 5 mm."
        )
        _hidro_txt = (
            f"Identificada hidronefrose {mrm['hidro_desc']}."
            if mrm["hidronefrose"] and mrm["hidro_desc"]
            else "Ausência de hidronefrose."
        )
        _eco = mrm["ecotextura"]
        _eco_txt = _eco if _eco != "normais" else "normais"
        add_p(
            f"Rim {lado_rim.lower()} medindo {mrm['comp']} cm, com espessura cortical de "
            f"{mrm['cortical']} cm, apresentando ecotextura e diferenciação corticomedular "
            f"{_eco_txt}. {_calc_txt} {_hidro_txt}"
        )

    # — ARTÉRIAS RENAIS —
    conclusoes = []
    for lado_art, da in dados_arterias.items():
        add_p(f"ARTÉRIA RENAL {lado_art.upper()}", space_before=10, space_after=4)

        if da.get("nao_identificada"):
            add_p(
                f"Artéria renal {lado_art.lower()} não identificada ao método, "
                f"possivelmente em razão de {da.get('limitacao_desc', 'limitação técnica')}. "
                "A ausência de visualização isoladamente não deve ser interpretada como oclusão."
            )
            conclusoes.append(f"Artéria renal {lado_art.lower()} não avaliada por limitação técnica ({da.get('limitacao_desc','')}).")
        else:
            _id_txt = da["id"]
            _cal_txt = da["calibre_traj"]
            add_p(
                f"Artéria renal {lado_art.lower()} identificada {_id_txt}, "
                f"apresentando calibre e trajeto {_cal_txt}."
            )
            add_p(f"VPS: {da['vps']} cm/s. VDF: {da['vdf']} cm/s. Razão renal-aórtica (RRA): {da['rra']}.")

            if da["turbulencia"]:
                add_p(
                    f"Observada aceleração focal do fluxo {da['turb_desc']}, com turbulência pós-estenótica ao mapeamento colorido."
                )
            else:
                add_p("Não se observam acelerações focais significativas do fluxo, alterações espectrais ou turbulência pós-estenótica.")

            _fi_pad = da["fluxo_intra_padrao"]
            if _fi_pad == "ausente (oclusão)":
                add_p(
                    f"Fluxo intrarrenal ausente ao Doppler colorido e espectral, "
                    "sugerindo oclusão arterial."
                )
                conclusoes.append(f"Sinais de oclusão da artéria renal {lado_art.lower()}.")
            else:
                add_p(
                    f"Fluxo intrarrenal com padrão de {_fi_pad}, apresentando VPS de "
                    f"{da['vps_intra']} cm/s, IR de {da['ir_intra']} e tempo de aceleração de "
                    f"{da['ta_intra']} ms."
                )
                # Conclusão baseada na classificação
                _grau = da["grau"]
                _cor = da["cor"]
                if _cor == "green":
                    if da.get("sinais_indiretos"):
                        _motivo = []
                        if da["fluxo_intra_padrao"] == "padrão tardus-parvus":
                            _motivo.append("padrão tardus-parvus")
                        _ta_val_r = _safe_float(da["ta_intra"]) or 0
                        if _ta_val_r > 100:
                            _motivo.append(f"tempo de aceleração de {da['ta_intra']} ms")
                        _motivo_txt = " e ".join(_motivo)
                        conclusoes.append(
                            f"Artéria renal {lado_art.lower()}: presença de sinais indiretos de "
                            f"estenose hemodinamicamente significativa ({_motivo_txt}), "
                            "na vigência de velocidades diretas na artéria renal dentro da normalidade. "
                            "Não se exclui estenose proximal de difícil demonstração ultrassonográfica."
                        )
                elif _cor == "blue":
                    conclusoes.append(
                        f"Artéria renal {lado_art.lower()}: sinais compatíveis com estenose hemodinamicamente não significativa (< 60%)."
                    )
                elif _cor == "orange":
                    conclusoes.append(
                        f"Artéria renal {lado_art.lower()}: sinais compatíveis com estenose hemodinamicamente significativa ({_grau})."
                    )
                elif _cor == "red":
                    conclusoes.append(
                        f"Artéria renal {lado_art.lower()}: sinais de estenose grave — {_grau}."
                    )

    # — OBSERVAÇÃO TÉCNICA —
    if obs_tecnica_texto:
        add_p("OBSERVAÇÃO TÉCNICA", space_before=10, space_after=4)
        add_p(obs_tecnica_texto)

    # — IMPRESSÃO DIAGNÓSTICA —
    if quebrar_pagina_diag:
        doc.add_page_break()
    add_p("⸻", space_after=12)
    add_p("IMPRESSÃO DIAGNÓSTICA", space_after=6)

    # Aorta
    if not aorta_aneurisma:
        add_p("Aorta abdominal sem alterações hemodinamicamente significativas ao estudo Doppler.", bullet=True)
    else:
        add_p(f"Aorta abdominal com alteração morfológica: {aorta_aneurisma_desc.strip()}", bullet=True)

    # Artérias renais
    _todas_normais = all(
        (not da.get("nao_identificada")) and da.get("cor") == "green" and not da.get("sinais_indiretos")
        for da in dados_arterias.values()
    )
    if _todas_normais:
        add_p("Ausência de sinais diretos ou indiretos de estenose hemodinamicamente significativa nas artérias renais.", bullet=True)
    else:
        for c in conclusoes:
            add_p(c, bullet=True)
        if not conclusoes:
            add_p("Ausência de sinais diretos ou indiretos de estenose hemodinamicamente significativa nas artérias renais.", bullet=True)

    # Observação — avaliação complementar
    if sugerir_complementar and metodos_complementares:
        if len(metodos_complementares) == 1:
            _met_txt = metodos_complementares[0]
        elif len(metodos_complementares) == 2:
            _met_txt = f"{metodos_complementares[0]} e {metodos_complementares[1]}"
        else:
            _met_txt = ", ".join(metodos_complementares[:-1]) + f" e {metodos_complementares[-1]}"
        add_p(
            f"OBSERVAÇÃO: Os achados ultrassonográficos do presente exame indicam correlação com "
            f"método de imagem complementar. Sugere-se a realização de {_met_txt} para melhor "
            f"caracterização anatômica e hemodinâmica das artérias renais e planejamento terapêutico adequado.",
            space_before=10, italic=True
        )

    # Assinatura
    if incluir_assinatura and (nome_medico or crm_medico):
        doc.add_paragraph().paragraph_format.space_before = Pt(25)
        p_assin = doc.add_paragraph()
        p_assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_assin.paragraph_format.line_spacing = espacamento_linhas
        if nome_medico:
            r = p_assin.add_run(f"{nome_medico}\n")
            r.bold = True
        if crm_medico:
            p_assin.add_run(f"CRM-{crm_uf} {crm_medico}\n")
        if rqe_medico.strip():
            p_assin.add_run(f"RQE {rqe_medico}")

    if incluir_rodape_link and rodape_url.strip():
        p_rod = doc.add_paragraph()
        p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rod.paragraph_format.space_before = Pt(30)
        r_rod = p_rod.add_run(f"Laudo gerado com suporte do sistema: {rodape_url.strip()}")
        r_rod.italic = True
        r_rod.font.size = Pt(max(tamanho_fonte - 2, 8))

    return doc


if st.button("🚀 Gerar Laudo de Artérias Renais", use_container_width=True):
    doc = construir_laudo_renais()
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.success("Laudo gerado com sucesso!")
    if modo_saida in ["Somente Visualização", "Visualização + DOCX"]:
        texto_viz = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        st.markdown("## 👁️ Visualização do Laudo")
        st.text_area("Laudo Gerado", value=texto_viz, height=600)
    if modo_saida in ["Somente DOCX", "Visualização + DOCX"]:
        st.download_button(
            "📥 Baixar Laudo de Artérias Renais (.docx)",
            buf,
            "Laudo_Arterias_Renais.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
