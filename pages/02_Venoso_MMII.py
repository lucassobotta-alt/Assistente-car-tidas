# venoso.py
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- CONFIGURAÇÃO DA PÁGINA ---
st.title("🌀 Assistente de Laudos: Duplex Scan Venoso de MMII")

# --- SIDEBAR: CONFIGURAÇÕES E DADOS DO MÉDICO ---
with st.sidebar:
    st.markdown("## ⚙️ Painel de Controle Avançado")

    st.markdown("### 📚 Referências Científicas")
    st.markdown(
        "📄 <a href='https://www.ejves.com/article/S1078-5884(21)00979-5/fulltext' target='_blank'>"
        "ESVS 2022 Clinical Practice Guidelines – Chronic Venous Disease</a><br>"
        "📄 <a href='https://www.minervamedica.it/en/journals/international-angiology/article.php?cod=R34Y2023N01A0045' target='_blank'>"
        "The First Latin American Consensus on Superficial and Perforating Venous Mapping</a><br>"
        "📄 <a href='https://cbr.org.br/wp-content/uploads/2020/05/Consenso-para-a-Sociedade-Bras.-de-Angiologia-e-Cirurgia-Vascular_2020.pdf' target='_blank'>"
        "Consenso Brasileiro de Angiologia e Cirurgia Vascular – SBACV 2020</a>",
        unsafe_allow_html=True
    )

    st.markdown("---")
    st.markdown("### 📝 Formatação Externa (.docx)")
    fonte_doc = st.selectbox("Família da Fonte:", ["Arial", "Calibri", "Times New Roman"])
    tamanho_fonte = st.slider("Tamanho do Texto (pt):", 10, 14, 11)
    espacamento_linhas = st.slider("Espaçamento entre Linhas:", 1.0, 1.5, 1.15, step=0.05)
    quebrar_pagina_diag = st.toggle("Separar Impressão Diagnóstica em Nova Página", value=False)
    modo_saida = st.radio(
        "Modo de saída do laudo:",
        ["Somente DOCX", "Somente Visualização", "Visualização + DOCX"],
        index=2
    )

    st.markdown("---")
    st.markdown("### ✍️ Identidade & Assinatura")
    incluir_identidade = st.toggle("Incluir identidade no laudo", value=False)
    nome_clinica = st.text_input("Cabeçalho / Nome da Clínica:", placeholder="Ex: Instituto de Diagnóstico Vascular")
    nome_medico = st.text_input("Nome do Médico:", "Lucas Santos Guimarães")
    colcrm1, colcrm2 = st.columns([2, 1])
    with colcrm1:
        crm_medico = st.text_input("CRM:", "4061")
    with colcrm2:
        crm_uf = st.selectbox("UF", ["AC","AL","AP","AM","BA","CE","DF","ES","GO","MA","MT","MS","MG","PA","PB","PR","PE","PI","RJ","RN","RS","RO","RR","SC","SP","SE","TO"], index=25)
    rqe_medico = st.text_input("RQE:", "")

    st.markdown("---")
    if st.button("🔄 Resetar Todos os Parâmetros", use_container_width=True, type="secondary"):
        _new_rc = st.session_state.get("_reset_count", 0) + 1
        for _k in list(st.session_state.keys()):
            del st.session_state[_k]
        st.session_state["_reset_count"] = _new_rc
        st.session_state["lista_perfurantes"] = {}
        st.session_state["segmentos_vsm_reg"] = {}
        st.session_state["vsp_reg"] = {}
        st.session_state["lista_varic_reg"] = {}
        st.session_state["_vsm_seg_fc"] = {}
        st.session_state["_vsp_fc"] = {}
        st.session_state["lista_varext_reg"] = {}
        st.toast("🔄 Todos os dados foram limpos!")
        st.rerun()

# --- IDENTIFICAÇÃO DO PACIENTE ---
nome_paciente = st.text_input("Nome do Paciente:", "")

TECNICAS = {
    "Técnica adequada (ortostase)": "Exame realizado com transdutor linear e paciente em ortostase, sem limitação técnica e com amostra adequada.",
    "Garroteamento (decúbito)": "Exame realizado com transdutor linear e paciente em decúbito, com garroteamento dos membros inferiores. Amostra adequada.",
    "Curativo na perna": "Exame realizado com transdutor linear e paciente em ortostase, com limitação técnica decorrente de curativo na perna, limitando a amostra adequada no segmento afetado.",
}
tecnica_selecionada = st.selectbox("Técnica do Exame:", list(TECNICAS.keys()))
tecnica_texto = TECNICAS[tecnica_selecionada]

formato_exame = st.selectbox("Tipo de Exame:", ["Unilateral", "Bilateral (Laudos Separados)", "Bilateral (Laudo Único)"])

membros_para_processar = [st.selectbox("Selecione o Lado Avaliado:", ["DIREITO", "ESQUERDO"])] if formato_exame == "Unilateral" else ["DIREITO", "ESQUERDO"]

st.markdown("---")

# Inicialização do estado de sessão para estruturas dinâmicas
if "_reset_count" not in st.session_state:
    st.session_state["_reset_count"] = 0
if "lista_perfurantes" not in st.session_state:
    st.session_state["lista_perfurantes"] = {}
if "segmentos_vsm_reg" not in st.session_state:
    st.session_state["segmentos_vsm_reg"] = {}
if "vsp_reg" not in st.session_state:
    st.session_state["vsp_reg"] = {}
if "lista_varic_reg" not in st.session_state:
    st.session_state["lista_varic_reg"] = {}
if "_vsm_seg_fc" not in st.session_state:
    st.session_state["_vsm_seg_fc"] = {}
if "_vsp_fc" not in st.session_state:
    st.session_state["_vsp_fc"] = {}
if "lista_varext_reg" not in st.session_state:
    st.session_state["lista_varext_reg"] = {}

_rc = st.session_state["_reset_count"]

dados_membros = {}

if len(membros_para_processar) > 1:
    abas = st.tabs(["🔴 Membro Inferior Direito (MID)", "🔵 Membro Inferior Esquerdo (MIE)"])
else:
    abas = [st.container()]

for idx, m_nome in enumerate(membros_para_processar):
    if m_nome not in st.session_state["lista_perfurantes"]:
        st.session_state["lista_perfurantes"][m_nome] = []
    if m_nome not in st.session_state["segmentos_vsm_reg"]:
        st.session_state["segmentos_vsm_reg"][m_nome] = []
    if m_nome not in st.session_state["vsp_reg"]:
        st.session_state["vsp_reg"][m_nome] = []
    if m_nome not in st.session_state["lista_varic_reg"]:
        st.session_state["lista_varic_reg"][m_nome] = []
    if m_nome not in st.session_state["_vsm_seg_fc"]:
        st.session_state["_vsm_seg_fc"][m_nome] = 0
    if m_nome not in st.session_state["_vsp_fc"]:
        st.session_state["_vsp_fc"][m_nome] = 0
    if m_nome not in st.session_state["lista_varext_reg"]:
        st.session_state["lista_varext_reg"][m_nome] = []
    _sfc = st.session_state["_vsm_seg_fc"][m_nome]
    _vfc = st.session_state["_vsp_fc"][m_nome]

    with abas[idx]:
        st.markdown(f"### 📋 Parâmetros Clínicos - Membro {m_nome}")
        
        # 1. SISTEMA VENOSO PROFUNDO (SVP)
        st.markdown("#### 1. Sistema Venoso Profundo (SVP)")
        svp_status = st.radio(f"Status do SVP ({m_nome}):", ["Normal", "Anormal"], horizontal=True, key=f"svp_stat_{m_nome}_{_rc}")
        svp_res = {"status": svp_status}
        if svp_status == "Anormal":
            c_svp1, c_svp2 = st.columns(2)
            with c_svp1: svp_res["tipo"] = st.selectbox("Tipo de Alteração:", ["Refluxo", "Trombose Venosa Profunda (TVP)"], key=f"svp_tipo_{m_nome}_{_rc}")
            with c_svp2: svp_res["veias"] = st.multiselect("Veias Acometidas:", ["Veia Femoral Comum (VFC)", "Veia Femoral (VF)", "Veia Femoral Profunda (VFP)", "Veia Poplítea (V POP)", "Veias Gastrocnêmias", "Veias Soleares", "Veias Tibiais Posteriores (VTP)", "Veias Fibulares"], key=f"svp_veias_{m_nome}_{_rc}")
        
        st.markdown("---")
        
        # 2. SISTEMA VENOSO SUPERFICIAL (SVS) - VEIA SAFENA MAGNA (VSM)
        st.markdown("#### 2. Sistema Venoso Superficial (SVS)")
        st.markdown("##### 2.1 Veia Safena Magna (VSM)")
        
        vsm_status_geral = st.selectbox("Status Geral da VSM:", ["Pérvia / Presente", "Pérvia, com segmento hipoplásico", "Ausente (Safenectomia Total)", "Ausente (Safenectomia Segmentar)"], key=f"vsm_status_geral_{m_nome}_{_rc}")
        vsm_dados_mapeamento = {"status_geral": vsm_status_geral}

        if "Ausente" in vsm_status_geral:
            pos_op_opt = st.selectbox("Pós-Operatório / Recidiva:", ["Não se aplica", "6.3 Sinais de neovascularização adjacentes"], key=f"pos_op_{m_nome}_{_rc}")
        else:
            pos_op_opt = "Não se aplica"

        if "Pérvia" in vsm_status_geral:
            st.markdown("**Avaliação da Junção Safenofemoral (JSF):**")
            jsf_status = st.radio("Status da JSF:", ["Competente", "Incompetente"], horizontal=True, key=f"jsf_status_{m_nome}_{_rc}")
            vsm_dados_mapeamento["jsf_status"] = jsf_status
            
            jsf_detalhes_input = {}
            
            if jsf_status == "Incompetente":
                jsf_valvulas = st.selectbox(
                    "Padrão de Incompetência Valvar da JSF:", 
                    ["Válvulas Terminal e Pré-terminal incompetentes", "Apenas a Válvula Pré-terminal incompetente", "Apenas a Válvula Terminal incompetente"], 
                    key=f"jsf_valvulas_{m_nome}_{_rc}"
                )
                vsm_dados_mapeamento["jsf_valvulas"] = jsf_valvulas
                
                st.markdown("<div style='background-color: #f0f7ff; padding: 12px; border-left: 4px solid #0056b3; border-radius: 4px; margin-top: 10px; margin-bottom: 10px;'><strong>📐 Extensão e Deságue do Refluxo Juncional</strong></div>", unsafe_allow_html=True)
                
                if "Terminal e Pré-terminal" in jsf_valvulas:
                    c1, c2, c3 = st.columns(3)
                    with c1: jsf_detalhes_input["extensao_refluxo"] = st.selectbox("O refluxo estende-se até:", ["toda sua extensão", "o terço proximal da coxa", "o terço médio da coxa", "o terço distal da coxa", "a altura do joelho", "o terço proximal da perna", "o terço médio da perna", "o terço distal da perna"], key=f"jsf_ext_{m_nome}_{_rc}")
                    with c2: jsf_detalhes_input["desague_tipo"] = st.selectbox("Drenagem / Deságue para:", ["tributárias epifasciais varicosas", "veia perfurante incompetente", "veia comunicante", "veias colaterais de face medial"], key=f"jsf_des_tipo_{m_nome}_{_rc}")
                    with c3: jsf_detalhes_input["ponto_ref_dist"] = st.selectbox("Ponto de Referência do Deságue:", ["Interlinha do Joelho", "Junção Safenofemoral (JSF)", "Face Plantar"], key=f"jsf_ref_dist_{m_nome}_{_rc}")
                    
                    c4, c5 = st.columns(2)
                    with c4:
                        if jsf_detalhes_input["ponto_ref_dist"] == "Interlinha do Joelho":
                            jsf_detalhes_input["posicao_joelho_dist"] = st.radio("Posição em relação ao joelho:", ["acima", "abaixo"], horizontal=True, key=f"jsf_pos_joelho_{m_nome}_{_rc}")
                        else: jsf_detalhes_input["posicao_joelho_dist"] = ""
                    with c5: jsf_detalhes_input["dist_fim"] = st.text_input("Distância do ponto de referência (cm):", "", key=f"jsf_dist_fim_{m_nome}_{_rc}")

                elif "Apenas a Válvula Pré-terminal" in jsf_valvulas:
                    c1, c2, c3 = st.columns(3)
                    with c1: jsf_detalhes_input["cm_ponto_j"] = st.text_input("Incompetente até cerca de (cm) da JSF/Ponto J:", "", key=f"jsf_det_j_{m_nome}_{_rc}")
                    with c2: jsf_detalhes_input["tipo_drenagem"] = st.selectbox("Drenagem para:", ["tributárias epifasciais varicosas", "veia perfurante incompetente", "veia comunicante"], key=f"jsf_det_dren_{m_nome}_{_rc}")
                    with c3: jsf_detalhes_input["terco_coxa"] = st.selectbox("Em qual terço anatômico:", ["terço proximal da coxa", "terço médio da coxa", "terço distal da coxa"], key=f"jsf_det_terco_{m_nome}_{_rc}")
                
                elif "Apenas a Válvula Terminal" in jsf_valvulas:
                    jsf_detalhes_input["destino_terminal"] = st.selectbox(
                        "Destino do refluxo da Válvula Terminal:",
                        ["Fluxo para tributárias não safênicas (Mantém VSM competente)", "Escape para a Veia Safena Acessória Anterior (VSAA)"],
                        key=f"jsf_dest_term_{m_nome}_{_rc}"
                    )
                    if "Acessória Anterior" in jsf_detalhes_input["destino_terminal"]:
                        c1, c2 = st.columns(2)
                        with c1: jsf_detalhes_input["vsaa_desague"] = st.text_input("Deságue da VSAA em veias epifasciais da face:", "", key=f"jsf_vsaa_des_{m_nome}_{_rc}")
                        with c2: jsf_detalhes_input["vsaa_extensao"] = st.text_input("Extensão dos trajetos varicosos até a face:", "", key=f"jsf_vsaa_ext_{m_nome}_{_rc}")
                    else:
                        jsf_detalhes_input["tributarias_tipo"] = st.text_input("Fluindo para tributárias não safênicas na:", "", key=f"jsf_det_term_trib_{m_nome}_{_rc}")

                vsm_dados_mapeamento["jsf_detalhes_input"] = jsf_detalhes_input

            incluir_segmento = st.toggle("Incluir Segmento Incompetente", key=f"incluir_seg_{m_nome}_{_rc}_{_sfc}")

            seg_sel = []
            seg_origem = "Insuficiência valvar isolada"
            seg_desague = "Tributária epifascial varicosa"
            seg_prox_ref = "Junção Safenofemoral (JSF)"
            seg_prox_pos = ""
            seg_prox_cm = ""
            seg_dist_ref = "Interlinha do Joelho"
            seg_dist_pos = "abaixo"
            seg_dist_cm = ""

            if incluir_segmento:
                # Linha 1: Origem
                seg_origem = st.selectbox(
                    "Origem do refluxo (escape proximal):",
                    ["Insuficiência valvar isolada", "Tributárias pélvicas", "Varizes ganglionares",
                     "Tributária epifascial incompetente", "Veia de Giacomini incompetente",
                     "Veia perfurante incompetente"],
                    key=f"seg_origem_{m_nome}_{_rc}_{_sfc}"
                )

                # Linha 2: Início — Referência e Distância
                _refs = ["Junção Safenofemoral (JSF)", "Interlinha do Joelho", "Face Plantar"]
                c_prox1, c_prox2, c_prox3 = st.columns(3)
                with c_prox1:
                    seg_prox_ref = st.selectbox("Início — Referência:", _refs, key=f"seg_prox_ref_{m_nome}_{_rc}_{_sfc}")
                with c_prox2:
                    if seg_prox_ref == "Interlinha do Joelho":
                        seg_prox_pos = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"seg_prox_pos_{m_nome}_{_rc}_{_sfc}")
                    else:
                        seg_prox_pos = ""
                        st.empty()
                with c_prox3:
                    seg_prox_cm = st.text_input("Distância (cm):", "", key=f"seg_prox_cm_{m_nome}_{_rc}_{_sfc}")

                # Linha 3: Ponto distal
                seg_desague = st.selectbox(
                    "Ponto distal (deságue do refluxo):",
                    ["Tributária epifascial varicosa", "Tributária varicosa troncular",
                     "Veia perfurante de drenagem", "Região maleolar"],
                    key=f"seg_desague_{m_nome}_{_rc}_{_sfc}"
                )

                # Linha 4: Fim — Referência e Distância (oculto se região maleolar)
                if seg_desague != "Região maleolar":
                    c_dist1, c_dist2, c_dist3 = st.columns(3)
                    with c_dist1:
                        seg_dist_ref = st.selectbox("Fim — Referência:", _refs, key=f"seg_dist_ref_{m_nome}_{_rc}_{_sfc}")
                    with c_dist2:
                        if seg_dist_ref == "Interlinha do Joelho":
                            seg_dist_pos = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"seg_dist_pos_{m_nome}_{_rc}_{_sfc}")
                        else:
                            seg_dist_pos = ""
                            st.empty()
                    with c_dist3:
                        seg_dist_cm = st.text_input("Distância (cm):", "", key=f"seg_dist_cm_{m_nome}_{_rc}_{_sfc}")
                else:
                    seg_dist_ref = ""
                    seg_dist_pos = ""
                    seg_dist_cm = ""

                if st.button("💾 Registrar Segmento", key=f"reg_seg_vsm_{m_nome}_{_rc}"):
                    st.session_state["segmentos_vsm_reg"][m_nome].append({
                        "segmentos": seg_sel,
                        "origem": seg_origem, "desague": seg_desague,
                        "prox_ref": seg_prox_ref, "prox_pos": seg_prox_pos, "prox_cm": seg_prox_cm,
                        "dist_ref": seg_dist_ref, "dist_pos": seg_dist_pos, "dist_cm": seg_dist_cm,
                    })
                    st.session_state["_vsm_seg_fc"][m_nome] += 1
                    st.rerun()

            if st.session_state["segmentos_vsm_reg"][m_nome]:
                for i, reg in enumerate(st.session_state["segmentos_vsm_reg"][m_nome]):
                    def _fmt_ref(ref, pos, cm):
                        if ref == "Interlinha do Joelho":
                            return f"{cm} cm {pos} da interlinha do joelho"
                        return f"{cm} cm da {ref}"
                    segs_label = ", ".join(reg["segmentos"]) if reg["segmentos"] else "insuficiência valvar"
                    desc = (f"`{i+1:02d}` **{segs_label}** | origem: {reg.get('origem','')} | deságue: {reg.get('desague','')} — "
                            f"de {_fmt_ref(reg['prox_ref'], reg['prox_pos'], reg['prox_cm'])} "
                            f"até {_fmt_ref(reg['dist_ref'], reg['dist_pos'], reg['dist_cm'])}")
                    st.markdown(desc)
                if st.button("❌ Limpar Segmentos Registrados", key=f"clear_seg_vsm_{m_nome}_{_rc}"):
                    st.session_state["segmentos_vsm_reg"][m_nome] = []
                    st.rerun()

            vsm_dados_mapeamento["tronco_refluxo"] = bool(st.session_state["segmentos_vsm_reg"][m_nome])
            vsm_dados_mapeamento["segmentos_lista"] = st.session_state["segmentos_vsm_reg"][m_nome]

        if "hipoplásico" in vsm_status_geral:
            st.markdown("<sub style='color: #444;'>Detalhamento do segmento hipoplásico:</sub>", unsafe_allow_html=True)
            _SEGS_VSM = [
                "Terço proximal da coxa", "Terço médio da coxa", "Terço distal da coxa",
                "Terço proximal da perna", "Terço médio da perna", "Terço distal da perna"
            ]
            vsm_dados_mapeamento["hipoplasico_segs"] = st.multiselect(
                "Segmento(s) hipoplásico(s):", _SEGS_VSM,
                key=f"vsm_hipo_segs_{m_nome}_{_rc}"
            )
            vsm_dados_mapeamento["hipoplasico_tributaria"] = st.toggle(
                "Tributária superficial acompanhando o segmento hipoplásico",
                key=f"vsm_hipo_trib_{m_nome}_{_rc}"
            )
            if vsm_dados_mapeamento["hipoplasico_tributaria"]:
                vsm_dados_mapeamento["hipoplasico_trib_calibre"] = st.text_input(
                    "Maior calibre da tributária (mm):", "",
                    key=f"vsm_hipo_trib_cal_{m_nome}_{_rc}"
                )

        # Mensurações da Veia Safena Magna
        if "Pérvia" in vsm_status_geral:
            st.markdown("**Mensurações da Veia Safena Magna (Diâmetros em mm):**")
            cm1, cm2, cm3, cm4 = st.columns(4)
            with cm1: jsf_mm = st.text_input("Junção safenofemoral (mm):", "", key=f"jsf_mm_{m_nome}_{_rc}")
            with cm2: vsm_prox_coxa = st.text_input("Terço proximal da coxa (mm):", "", key=f"prox_c_{m_nome}_{_rc}")
            with cm3: vsm_med_coxa = st.text_input("Terço médio da coxa (mm):", "", key=f"med_c_{m_nome}_{_rc}")
            with cm4: vsm_dist_coxa = st.text_input("Terço distal da coxa (mm):", "", key=f"dist_c_{m_nome}_{_rc}")
            cm5, cm6, cm7 = st.columns(3)
            with cm5: vsm_prox_perna = st.text_input("Terço proximal da perna (mm):", "", key=f"prox_p_{m_nome}_{_rc}")
            with cm6: vsm_med_perna = st.text_input("Terço médio da perna (mm):", "", key=f"med_p_{m_nome}_{_rc}")
            with cm7: vsm_dist_perna = st.text_input("Terço distal da perna (mm):", "", key=f"dist_p_{m_nome}_{_rc}")
            # Alertas de diâmetro da VSM (ESVS 2022)
            _alertas_vsm = []
            try:
                if float(jsf_mm) > 8.0:
                    _alertas_vsm.append(f"JSF: {jsf_mm} mm > 8,0 mm — dilatação acentuada da junção safenofemoral")
            except ValueError: pass
            for _lbl, _val in [("Terço proximal da coxa", vsm_prox_coxa), ("Terço médio da coxa", vsm_med_coxa),
                                ("Terço distal da coxa", vsm_dist_coxa), ("Terço proximal da perna", vsm_prox_perna),
                                ("Terço médio da perna", vsm_med_perna), ("Terço distal da perna", vsm_dist_perna)]:
                try:
                    if float(_val) > 6.0:
                        _alertas_vsm.append(f"{_lbl}: {_val} mm > 6,0 mm — tronco dilatado (limiar ESVS 2022 para indicação terapêutica)")
                except ValueError: pass
            if _alertas_vsm:
                st.warning("⚠️ **Diâmetros fora do limiar de referência (ESVS 2022):**\n\n" + "\n\n".join(f"• {a}" for a in _alertas_vsm))
        else: jsf_mm = vsm_prox_coxa = vsm_med_coxa = vsm_dist_coxa = vsm_prox_perna = vsm_med_perna = vsm_dist_perna = ""

        st.markdown("---")

        # 2.2 VEIA SAFENA PARVA (VSP)
        st.markdown("##### 2.2 Veia Safena Parva (VSP)")
        vsp_template = st.selectbox(
            "Avaliação da Junção Safenopoplítea / Veia Safena Parva:",
            [
                "Junção safenopoplítea ausente com extensão cranial da veia safena parva",
                "Junção safenopoplítea competente",
                "Junção safenopoplítea incompetente com escape do refluxo para a veia safena parva",
                "Junção safenopoplítea incompetente com refluxo drenado de forma ascendente para veia de Giacomini",
            ],
            key=f"vsp_temp_choise_{m_nome}_{_rc}"
        )
        vsp_dados_form = {"template": vsp_template}

        if "escape do refluxo para a veia safena parva" in vsp_template:
            vsp_dados_form["vsp_incomp_tipo"] = st.radio(
                "Padrão de incompetência da veia safena parva:",
                ["Incompetência completa", "Incompetência segmentar"],
                horizontal=True,
                key=f"vsp_incomp_tipo_{m_nome}_{_rc}_{_vfc}"
            )
            if vsp_dados_form["vsp_incomp_tipo"] == "Incompetência segmentar":
                st.markdown("<sub style='color: #444;'>Detalhamento do segmento incompetente:</sub>", unsafe_allow_html=True)
                _refs_vsp = ["Prega Poplítea", "Face Plantar"]

                vsp_dados_form["vsp_escape"] = st.selectbox(
                    "Escape proximal (origem do refluxo):",
                    ["Veia perfurante incompetente", "Tributária incompetente", "Veia varicosa troncular",
                     "Varizes do arco posterior originadas da VSM", "Varizes alimentadas por refluxo de veia ciática"],
                    key=f"vsp_escape_{m_nome}_{_rc}_{_vfc}"
                )

                cvsp_p1, cvsp_p2, cvsp_p3 = st.columns(3)
                with cvsp_p1:
                    vsp_dados_form["vsp_prox_ref"] = st.selectbox("Início — Referência:", _refs_vsp, key=f"vsp_prox_ref_{m_nome}_{_rc}_{_vfc}")
                with cvsp_p2:
                    if vsp_dados_form["vsp_prox_ref"] == "Prega Poplítea":
                        vsp_dados_form["vsp_prox_pos"] = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"vsp_prox_pos_{m_nome}_{_rc}_{_vfc}")
                    else:
                        vsp_dados_form["vsp_prox_pos"] = ""
                        st.empty()
                with cvsp_p3:
                    vsp_dados_form["vsp_prox_cm"] = st.text_input("Distância (cm):", "", key=f"vsp_prox_cm_{m_nome}_{_rc}_{_vfc}")

                vsp_dados_form["vsp_desague_tipo"] = st.selectbox(
                    "Deságue do refluxo:",
                    ["Varizes de face posterior", "Veia perfurante de drenagem", "Continuidade para a região maleolar"],
                    key=f"vsp_desague_tipo_{m_nome}_{_rc}_{_vfc}"
                )

                if vsp_dados_form["vsp_desague_tipo"] != "Continuidade para a região maleolar":
                    cvsp_d1, cvsp_d2, cvsp_d3 = st.columns(3)
                    with cvsp_d1:
                        vsp_dados_form["vsp_dist_ref"] = st.selectbox("Fim — Referência:", _refs_vsp, key=f"vsp_dist_ref_{m_nome}_{_rc}_{_vfc}")
                    with cvsp_d2:
                        if vsp_dados_form["vsp_dist_ref"] == "Prega Poplítea":
                            vsp_dados_form["vsp_dist_pos"] = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"vsp_dist_pos_{m_nome}_{_rc}_{_vfc}")
                        else:
                            vsp_dados_form["vsp_dist_pos"] = ""
                            st.empty()
                    with cvsp_d3:
                        vsp_dados_form["vsp_dist_cm"] = st.text_input("Distância (cm):", "", key=f"vsp_dist_cm_{m_nome}_{_rc}_{_vfc}")

        elif "extensão cranial" in vsp_template or vsp_template == "Junção safenopoplítea competente":
            vsp_dados_form["tem_seg_vsp"] = st.toggle(
                "Incluir segmento com refluxo na veia safena parva?",
                key=f"vsp_tem_seg_{m_nome}_{_rc}_{_vfc}"
            )
            if vsp_dados_form.get("tem_seg_vsp"):
                st.markdown("<sub style='color: #444;'>Detalhamento do segmento com refluxo:</sub>", unsafe_allow_html=True)
                _refs_vsp = ["Prega Poplítea", "Face Plantar"]
                vsp_dados_form["vsp_escape"] = st.selectbox(
                    "Escape proximal (origem do refluxo):",
                    ["Veia perfurante incompetente", "Tributária incompetente", "Veia varicosa troncular",
                     "Varizes do arco posterior originadas da VSM", "Varizes alimentadas por refluxo de veia ciática"],
                    key=f"vsp_escape_{m_nome}_{_rc}_{_vfc}"
                )
                cvsp_p1, cvsp_p2, cvsp_p3 = st.columns(3)
                with cvsp_p1:
                    vsp_dados_form["vsp_prox_ref"] = st.selectbox("Início — Referência:", _refs_vsp, key=f"vsp_prox_ref_{m_nome}_{_rc}_{_vfc}")
                with cvsp_p2:
                    if vsp_dados_form["vsp_prox_ref"] == "Prega Poplítea":
                        vsp_dados_form["vsp_prox_pos"] = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"vsp_prox_pos_{m_nome}_{_rc}_{_vfc}")
                    else:
                        vsp_dados_form["vsp_prox_pos"] = ""
                        st.empty()
                with cvsp_p3:
                    vsp_dados_form["vsp_prox_cm"] = st.text_input("Distância (cm):", "", key=f"vsp_prox_cm_{m_nome}_{_rc}_{_vfc}")
                vsp_dados_form["vsp_desague_tipo"] = st.selectbox(
                    "Deságue do refluxo:",
                    ["Varizes de face posterior", "Veia perfurante de drenagem", "Continuidade para a região maleolar"],
                    key=f"vsp_desague_tipo_{m_nome}_{_rc}_{_vfc}"
                )
                if vsp_dados_form["vsp_desague_tipo"] != "Continuidade para a região maleolar":
                    cvsp_d1, cvsp_d2, cvsp_d3 = st.columns(3)
                    with cvsp_d1:
                        vsp_dados_form["vsp_dist_ref"] = st.selectbox("Fim — Referência:", _refs_vsp, key=f"vsp_dist_ref_{m_nome}_{_rc}_{_vfc}")
                    with cvsp_d2:
                        if vsp_dados_form["vsp_dist_ref"] == "Prega Poplítea":
                            vsp_dados_form["vsp_dist_pos"] = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"vsp_dist_pos_{m_nome}_{_rc}_{_vfc}")
                        else:
                            vsp_dados_form["vsp_dist_pos"] = ""
                            st.empty()
                    with cvsp_d3:
                        vsp_dados_form["vsp_dist_cm"] = st.text_input("Distância (cm):", "", key=f"vsp_dist_cm_{m_nome}_{_rc}_{_vfc}")

        if "Ausente (Safenectomia" not in vsp_template:
            st.markdown("**Mensurações da Veia Safena Parva (Diâmetros em mm):**")
            cp1, cp2, cp3 = st.columns(3)
            label_jsp_dinamico = "Extensão cranial (mm):" if "extensão cranial" in vsp_template else "Junção safenopoplítea (mm):"
            with cp1: vsp_dados_form["jsp_mm"] = st.text_input(label_jsp_dinamico, "", key=f"jsp_mm_{m_nome}_{_rc}_{_vfc}", disabled="extensão cranial" in vsp_template)
            with cp2: vsp_dados_form["vsp_crossa"] = st.text_input("Crossa da safena parva (mm):", "", key=f"crossa_{m_nome}_{_rc}_{_vfc}")
            with cp3: vsp_dados_form["vsp_med_perna_diam"] = st.text_input("Terço médio da perna (mm):", "", key=f"med_per_{m_nome}_{_rc}_{_vfc}")

        if st.button("💾 Registrar Achado da VSP", key=f"reg_vsp_{m_nome}_{_rc}"):
            st.session_state["vsp_reg"][m_nome] = vsp_dados_form
            st.session_state["_vsp_fc"][m_nome] += 1
            st.rerun()

        vsp_dados_input = st.session_state["vsp_reg"][m_nome] if st.session_state["vsp_reg"][m_nome] else vsp_dados_form

        if st.session_state["vsp_reg"][m_nome]:
            reg_vsp = st.session_state["vsp_reg"][m_nome]
            lbl_vsp = reg_vsp["template"].split("(")[0].strip() if "(" in reg_vsp["template"] else reg_vsp["template"][:60]
            st.write(f"✅ **VSP registrada:** {lbl_vsp}")
            if st.button("❌ Limpar Registro VSP", key=f"clear_vsp_{m_nome}_{_rc}"):
                st.session_state["vsp_reg"][m_nome] = []
                st.rerun()

        jsp_mm = vsp_dados_input.get("jsp_mm", "")
        vsp_crossa = vsp_dados_input.get("vsp_crossa", "")
        vsp_med_perna_diam = vsp_dados_input.get("vsp_med_perna_diam", "")

        st.markdown("---")

        # --- 2.3 VEIAS PERFURANTES INCOMPETENTES ---
        st.markdown("#### 2.3 Veias Perfurantes Incompetentes")

        _perf_saved = st.session_state["lista_perfurantes"][m_nome]
        if _perf_saved:
            st.markdown(f"**{len(_perf_saved)} perfurante(s) registrada(s):**")
            for _pi, _pitem in enumerate(_perf_saved):
                _pc1, _pc2 = st.columns([6, 1])
                with _pc1:
                    _pos_txt = f", {_pitem['posicao_joelho'].lower()} da interlinha do joelho" if _pitem.get("posicao_joelho") else ""
                    _diam_txt = f", {_pitem['diametro_mm']} mm" if _pitem.get("diametro_mm") else ""
                    st.markdown(f"• {_pitem['regiao']}, face {_pitem['face'].lower()} — {_pitem['altura_cm']} cm{_pos_txt}{_diam_txt}")
                with _pc2:
                    if st.button("❌", key=f"rem_perf_{m_nome}_{_pi}_{_rc}"):
                        st.session_state["lista_perfurantes"][m_nome].pop(_pi)
                        st.rerun()

        st.markdown("<sub style='color: #444;'>Adicionar nova perfurante:</sub>", unsafe_allow_html=True)
        cp_1, cp_2, cp_3 = st.columns(3)
        with cp_1:
            _perf_regiao = st.selectbox("Região Anatômica:", ["Coxa", "Perna"], key=f"perf_reg_{m_nome}_{_rc}")
        with cp_2:
            _perf_face = st.selectbox("Face:", ["Medial", "Lateral", "Anterior", "Posterior", "Anterolateral", "Posterointerna"], key=f"perf_face_{m_nome}_{_rc}")
        with cp_3:
            _perf_ref = st.selectbox("Referência:", ["Interlinha do Joelho", "Face Plantar"], key=f"perf_ref_{m_nome}_{_rc}")
        cp_4, cp_5, cp_6 = st.columns(3)
        with cp_4:
            _perf_alt = st.text_input("Altura aferida (cm):", "", key=f"perf_alt_{m_nome}_{_rc}")
        with cp_5:
            _perf_diam = st.text_input("Diâmetro (mm):", "", key=f"perf_diam_{m_nome}_{_rc}")
        with cp_6:
            if _perf_ref == "Interlinha do Joelho":
                _perf_pos_j = st.radio("Posição:", ["Acima", "Abaixo"], horizontal=True, key=f"perf_pos_j_{m_nome}_{_rc}")
            else:
                _perf_pos_j = ""
                st.empty()
        try:
            if float(_perf_diam) > 3.5:
                st.warning(f"⚠️ Diâmetro {_perf_diam} mm > 3,5 mm — critério ESVS 2022 para perfurante patológica")
        except (ValueError, TypeError): pass
        if st.button("💾 Salvar Perfurante", key=f"save_perf_{m_nome}_{_rc}"):
            st.session_state["lista_perfurantes"][m_nome].append({
                "regiao": _perf_regiao, "face": _perf_face, "ref_ponto": _perf_ref,
                "altura_cm": _perf_alt, "diametro_mm": _perf_diam, "posicao_joelho": _perf_pos_j
            })
            st.rerun()

        perfurantes_coletadas = st.session_state["lista_perfurantes"][m_nome]
                        
        # --- 2.4 MAPEAMENTO DE VARICOSIDADES ---
        st.markdown("---")
        st.markdown("#### 2.4 Mapeamento de Varicosidades")

        _varic_saved = st.session_state["lista_varic_reg"][m_nome]
        if _varic_saved:
            st.markdown(f"**{len(_varic_saved)} lesão(ões) registrada(s):**")
            for _vi, _vitem in enumerate(_varic_saved):
                _vc1, _vc2 = st.columns([6, 1])
                with _vc1:
                    st.markdown(f"• **{_vitem['tipo']}** — {_vitem['localizacao']}")
                with _vc2:
                    if st.button("❌", key=f"rem_varic_{m_nome}_{_vi}_{_rc}"):
                        st.session_state["lista_varic_reg"][m_nome].pop(_vi)
                        st.rerun()

        st.markdown("<sub style='color: #444;'>Adicionar nova lesão:</sub>", unsafe_allow_html=True)
        _cv_t1, _cv_t2 = st.columns([2, 3])
        with _cv_t1:
            _varic_tipo = st.selectbox(
                "Tipo:",
                ["Telangiectasias (< 1 mm)", "Microvarizes / Varizes Reticulares (1–3 mm)", "Veias Varicosas Tronculares (> 3 mm)"],
                key=f"varic_tipo_{m_nome}_{_rc}"
            )
        with _cv_t2:
            _varic_loc = st.text_input("Localização:", "", key=f"varic_loc_{m_nome}_{_rc}")
        if st.button("💾 Salvar Lesão", key=f"save_varic_{m_nome}_{_rc}"):
            st.session_state["lista_varic_reg"][m_nome].append({"tipo": _varic_tipo, "localizacao": _varic_loc})
            st.rerun()

        varic_dados = {"lista": _varic_saved}

        st.markdown("---")

        # --- 2.5 VARIZES EXTRASSAFÊNICAS ---
        st.markdown("#### 2.5 Varizes Extrassafênicas")

        _varext_saved = st.session_state["lista_varext_reg"][m_nome]
        if _varext_saved:
            st.markdown(f"**{len(_varext_saved)} registro(s) de varizes extrassafênicas:**")
            for _vei, _veitem in enumerate(_varext_saved):
                _vec1, _vec2 = st.columns([6, 1])
                with _vec1:
                    _ve_lbl = _veitem['origem']
                    if _veitem.get("localizacao"):
                        _ve_lbl += f" — {_veitem['localizacao']}"
                    if _veitem.get("trib_cm"):
                        _pos = f" {_veitem['trib_pos']}" if _veitem.get("trib_pos") else ""
                        _ve_lbl += f" ({_veitem['trib_cm']} cm{_pos} da {_veitem.get('trib_ref','')})"
                    if _veitem.get("ciatica_subtipo"):
                        _ve_lbl += f" ({_veitem['ciatica_subtipo']})"
                    if _veitem.get("pelvico_pontos"):
                        _ve_lbl += f" [{', '.join(_veitem['pelvico_pontos'])}]"
                    st.markdown(f"• {_ve_lbl}")
                with _vec2:
                    if st.button("❌", key=f"rem_varext_{m_nome}_{_vei}_{_rc}"):
                        st.session_state["lista_varext_reg"][m_nome].pop(_vei)
                        st.rerun()

        st.markdown("<sub style='color: #444;'>Adicionar nova entrada:</sub>", unsafe_allow_html=True)
        _varext_origem = st.selectbox(
            "Origem:",
            ["Tributária incompetente", "Refluxo de origem ciática", "Refluxo pélvico"],
            key=f"varext_origem_{m_nome}_{_rc}"
        )
        _varext_loc = ""
        _varext_trib_ref = ""
        _varext_trib_pos = ""
        _varext_trib_cm = ""
        _varext_ciatica = ""
        _varext_pelvico = []
        if _varext_origem == "Tributária incompetente":
            _varext_loc = st.text_input("Localização:", "", key=f"varext_loc_{m_nome}_{_rc}")
            _vt1, _vt2, _vt3 = st.columns(3)
            with _vt1:
                _varext_trib_ref = st.selectbox(
                    "Referência de altura:",
                    ["Junção Safenofemoral", "Interlinha do Joelho", "Face Plantar"],
                    key=f"varext_trib_ref_{m_nome}_{_rc}"
                )
            with _vt2:
                if _varext_trib_ref == "Interlinha do Joelho":
                    _varext_trib_pos = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"varext_trib_pos_{m_nome}_{_rc}")
                else:
                    _varext_trib_pos = ""
                    st.empty()
            with _vt3:
                _varext_trib_cm = st.text_input("Distância (cm):", "", key=f"varext_trib_cm_{m_nome}_{_rc}")
        elif _varext_origem == "Refluxo de origem ciática":
            _varext_ciatica = st.radio(
                "Tipo de refluxo ciático:",
                ["Varizes acompanhando o trajeto do nervo ciático", "Veia ciática persistente"],
                horizontal=True,
                key=f"varext_ciatica_{m_nome}_{_rc}"
            )
        elif _varext_origem == "Refluxo pélvico":
            _varext_pelvico = st.multiselect(
                "Ponto(s) de escape do refluxo pélvico:",
                ["Ponto inguinal", "Ponto perineal", "Ponto obturatório", "Ponto glúteo"],
                key=f"varext_pelvico_{m_nome}_{_rc}"
            )
        if st.button("💾 Salvar Varizes Extrassafênicas", key=f"save_varext_{m_nome}_{_rc}"):
            st.session_state["lista_varext_reg"][m_nome].append({
                "localizacao": _varext_loc,
                "origem": _varext_origem,
                "trib_ref": _varext_trib_ref,
                "trib_pos": _varext_trib_pos,
                "trib_cm": _varext_trib_cm,
                "ciatica_subtipo": _varext_ciatica,
                "pelvico_pontos": _varext_pelvico,
            })
            st.rerun()

        varizes_extrassaf_dados = {"lista": _varext_saved}

        st.markdown("---")

        # 3. ACHADOS ADICIONAIS
        st.markdown("#### 3. Achados Adicionais")
        achados_multiplos = st.multiselect(
            "Patologias de tecidos adjacentes:",
            ["Cisto de Baker na fossa poplítea", "Edema intersticial subcutâneo"],
            default=[],
            key=f"achados_adi_multi_{m_nome}_{_rc}"
        )

        cisto_medidas = {}
        if "Cisto de Baker na fossa poplítea" in achados_multiplos:
            st.markdown("<sub style='color: #444;'>Dimensões do Cisto de Baker:</sub>", unsafe_allow_html=True)
            cc1, cc2, cc3 = st.columns(3)
            with cc1: cisto_medidas["comp"] = st.text_input("Eixo Long (mm):", "", key=f"cb_c_{m_nome}_{_rc}")
            with cc2: cisto_medidas["larg"] = st.text_input("Eixo Transv (mm):", "", key=f"cb_l_{m_nome}_{_rc}")
            with cc3: cisto_medidas["esp"] = st.text_input("Espessura (mm):", "", key=f"cb_e_{m_nome}_{_rc}")

        dados_membros[m_nome] = {
            "svp": svp_res, "vsm_mapeamento": vsm_dados_mapeamento,
            "jsf_mm": jsf_mm, "vsm_prox_coxa": vsm_prox_coxa, "vsm_med_coxa": vsm_med_coxa, "vsm_dist_coxa": vsm_dist_coxa,
            "vsm_prox_perna": vsm_prox_perna, "vsm_med_perna": vsm_med_perna, "vsm_dist_perna": vsm_dist_perna,
            "vsp_dados_input": vsp_dados_input, "jsp_mm": jsp_mm, "vsp_crossa": vsp_crossa, "vsp_med_perna_diam": vsp_med_perna_diam,
            "pos_op_opt": pos_op_opt, "achados_adi_multi": achados_multiplos, "cisto_medidas": cisto_medidas,
            "perfurantes_lista": perfurantes_coletadas,
            "varic_dados": varic_dados,
            "varizes_extrassaf_dados": varizes_extrassaf_dados
        }

st.markdown("---")

# --- HELPERS VSP SEGMENTAR ---
def _fmt_vsp_ref(ref, pos, cm):
    if ref == "Prega Poplítea" and pos:
        return f"{cm} cm {pos} da prega poplítea"
    elif ref == "Prega Poplítea":
        return f"{cm} cm da prega poplítea"
    elif ref == "Face Plantar":
        return f"{cm} cm da face plantar"
    return ""

def _build_seg_vsp_text(vsp_d):
    _vsp_escape    = vsp_d.get("vsp_escape", "")
    _vsp_prox_ref  = vsp_d.get("vsp_prox_ref", "")
    _vsp_prox_pos  = vsp_d.get("vsp_prox_pos", "")
    _vsp_prox_cm   = vsp_d.get("vsp_prox_cm", "")
    _vsp_desague_t = vsp_d.get("vsp_desague_tipo", "Varizes de face posterior")
    _vsp_dist_ref  = vsp_d.get("vsp_dist_ref", "")
    _vsp_dist_pos  = vsp_d.get("vsp_dist_pos", "")
    _vsp_dist_cm   = vsp_d.get("vsp_dist_cm", "")
    escape_txt = f", originado de {_vsp_escape.lower()}" if _vsp_escape else ""
    inicio_txt = _fmt_vsp_ref(_vsp_prox_ref, _vsp_prox_pos, _vsp_prox_cm)
    if _vsp_desague_t == "Continuidade para a região maleolar":
        ext_txt  = f"a partir de {inicio_txt}, com continuidade do refluxo para a região maleolar" if inicio_txt else "com continuidade do refluxo para a região maleolar"
        residual = ""
    else:
        fim_txt = _fmt_vsp_ref(_vsp_dist_ref, _vsp_dist_pos, _vsp_dist_cm)
        if inicio_txt and fim_txt:
            ext_txt = f"de {inicio_txt} até {fim_txt}, com deságue para {_vsp_desague_t.lower()}"
        elif fim_txt:
            ext_txt = f"até {fim_txt}, com deságue para {_vsp_desague_t.lower()}"
        else:
            ext_txt = f"com deságue para {_vsp_desague_t.lower()}"
        residual = " No restante do trajeto a veia safena parva segue competente."
    return escape_txt, ext_txt, residual

# --- FUNÇÃO DE CARTOGRAFIA VENOSA ---
def gerar_cartografia_venosa(m_nome, dados_m, paciente):
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.lines import Line2D
    import numpy as np

    KNEE_H = 38
    JSF_H  = 80
    JSP_H  = 37
    TOTAL_H = 90

    EXTENT_MAP = {
        "toda sua extensão":       0,
        "terço proximal da coxa":  66,
        "terço médio da coxa":     55,
        "terço distal da coxa":    44,
        "altura do joelho":        KNEE_H,
        "terço proximal da perna": 28,
        "terço médio da perna":    18,
        "terço distal da perna":   8,
    }

    def texto_para_altura(txt):
        for chave, h in EXTENT_MAP.items():
            if chave in txt.lower():
                return h
        return 0

    def calc_altura(ref, pos, cm_str):
        try:
            cm = float(str(cm_str).replace(",", ".").strip() or "0")
        except (ValueError, TypeError):
            cm = 0.0
        ref = str(ref)
        pos = str(pos).lower()
        if "Plantar" in ref:
            return cm
        if "JSF" in ref or "Safenofemoral" in ref:
            return JSF_H - cm
        if "Joelho" in ref:
            return (KNEE_H + cm) if "acima" in pos else (KNEE_H - cm)
        return cm

    # ---- silhueta ----
    SIL_Y  = np.array([0,  6,  16,  26,  38,  50,  62,  73,  82])
    SIL_XR = np.array([2.0, 2.3, 2.7, 2.9, 2.3, 3.1, 3.5, 3.3, 2.6])
    SIL_XL = np.array([-2.0,-2.3,-2.6,-2.8,-2.2,-2.7,-3.1,-2.9,-2.3])

    def draw_silhueta(ax):
        xs_f = np.concatenate([SIL_XR, SIL_XL[::-1]])
        ys_f = np.concatenate([SIL_Y, SIL_Y[::-1]])
        ax.fill(xs_f, ys_f, color='#EFD9C6', alpha=0.55, zorder=1)
        ax.plot(SIL_XR, SIL_Y, color='#C49A7A', lw=1.5, zorder=2)
        ax.plot(SIL_XL, SIL_Y, color='#C49A7A', lw=1.5, zorder=2)
        # pé
        ax.fill([-2.0, 2.0, 3.2, -0.5], [0, 0, -3, -3],
                color='#EFD9C6', alpha=0.55, zorder=1)
        ax.plot([-2.0, -0.5, 3.2, 2.0], [0, -3, -3, 0],
                color='#C49A7A', lw=1.5, zorder=2)

    # ---- figura ----
    fig, (ax_med, ax_post) = plt.subplots(1, 2, figsize=(11, 15))
    fig.patch.set_facecolor('#F8F9FA')

    for ax in (ax_med, ax_post):
        ax.set_facecolor('#F8F9FA')
        ax.set_xlim(-7, 7)
        ax.set_ylim(-6, TOTAL_H + 6)
        ax.set_aspect('equal', adjustable='box')
        ax.axis('off')
        draw_silhueta(ax)
        # linha do joelho
        ax.plot([-3.5, 3.5], [KNEE_H, KNEE_H], color='#AAAAAA',
                lw=0.9, ls='--', zorder=3, alpha=0.8)
        ax.text(4.1, KNEE_H, "Joelho", fontsize=7, va='center', color='#888888')
        # régua
        for h in range(0, 85, 10):
            ax.plot([-6.1, -5.7], [h, h], color='#CCCCCC', lw=0.7, zorder=3)
            ax.text(-6.5, h, f"{h}", fontsize=5.5, va='center',
                    ha='right', color='#AAAAAA')
        ax.text(-6.5, -4, "cm", fontsize=5.5, ha='right', color='#AAAAAA')

    # ---- VSM ----
    VSM_X = -0.8
    vsm = dados_m["vsm_mapeamento"]
    vsm_status_geral = vsm.get("status_geral", "")

    if "Ausente" in vsm_status_geral:
        ax_med.plot([VSM_X, VSM_X], [0, JSF_H], color='#BBBBBB',
                    lw=3, ls=':', zorder=5, alpha=0.7)
        ax_med.text(VSM_X + 0.3, 40, "VSM\nAusente", fontsize=8,
                    va='center', color='#AAAAAA', style='italic')
    else:
        jsf_incomp = vsm.get("jsf_status", "") == "Incompetente"
        valvulas   = vsm.get("jsf_valvulas", "")
        det        = vsm.get("jsf_detalhes_input", {})

        regioes_incomp = []

        if jsf_incomp:
            if "Terminal e Pré-terminal" in valvulas:
                y_base = texto_para_altura(det.get("extensao_refluxo", "toda sua extensão"))
                regioes_incomp.append((JSF_H, y_base))
            elif "Apenas a Válvula Pré-terminal" in valvulas:
                try:
                    cm_j = float(str(det.get("cm_ponto_j","10") or "10").replace(",","."))
                except ValueError:
                    cm_j = 10.0
                regioes_incomp.append((JSF_H, JSF_H - cm_j))
            elif "Apenas a Válvula Terminal" in valvulas:
                if "Acessória Anterior" not in det.get("destino_terminal",""):
                    regioes_incomp.append((JSF_H, JSF_H - 4))

        for seg in vsm.get("segmentos_lista", []):
            yt = calc_altura(seg.get("prox_ref",""), seg.get("prox_pos",""), seg.get("prox_cm","0"))
            yb = calc_altura(seg.get("dist_ref",""), seg.get("dist_pos",""), seg.get("dist_cm","0"))
            if seg.get("desague") == "Região maleolar":
                yb = 0
            regioes_incomp.append((max(yt, yb), min(yt, yb)))

        def is_incomp(y):
            return any(yb <= y <= yt for yt, yb in regioes_incomp)

        ys_vsm = np.linspace(0, JSF_H, 400)
        prev_y    = ys_vsm[0]
        prev_incomp = is_incomp(prev_y)
        for cur_y in ys_vsm[1:]:
            cur_incomp = is_incomp(cur_y)
            if cur_incomp != prev_incomp:
                cor = '#C0392B' if prev_incomp else '#1565C0'
                ax_med.plot([VSM_X, VSM_X], [prev_y, cur_y],
                            color=cor, lw=4, solid_capstyle='round', zorder=5)
                prev_y = cur_y
                prev_incomp = cur_incomp
        cor = '#C0392B' if prev_incomp else '#1565C0'
        ax_med.plot([VSM_X, VSM_X], [prev_y, JSF_H],
                    color=cor, lw=4, solid_capstyle='round', zorder=5)

        # setas de refluxo
        for yt, yb in regioes_incomp:
            if yt - yb > 6:
                mid = (yt + yb) / 2
                ax_med.annotate('', xy=(VSM_X + 0.7, mid - 5),
                                xytext=(VSM_X + 0.7, mid + 5),
                                arrowprops=dict(arrowstyle='->', color='#C0392B',
                                                lw=1.5, mutation_scale=13))

        # rótulo VSM
        ax_med.text(VSM_X - 0.35, 45, "VSM", fontsize=8.5, ha='right',
                    va='center', color='#1A1A2E', fontweight='bold', rotation=90)

        # junção JSF
        jsf_cor = '#C0392B' if jsf_incomp else '#1565C0'
        ax_med.scatter([VSM_X], [JSF_H], s=160, color=jsf_cor,
                       marker='D', zorder=10, edgecolors='white', linewidths=0.8)
        ax_med.text(VSM_X + 0.4, JSF_H,
                    f"JSF ({'INC' if jsf_incomp else 'COMP'})",
                    fontsize=7.5, va='center', color=jsf_cor, fontweight='bold')

        # diâmetros VSM
        DIAM_VSM = [
            (JSF_H,        dados_m.get("jsf_mm",""),        8.0),
            (66,           dados_m.get("vsm_prox_coxa",""), 6.0),
            (55,           dados_m.get("vsm_med_coxa",""),  6.0),
            (44,           dados_m.get("vsm_dist_coxa",""), 6.0),
            (KNEE_H - 10,  dados_m.get("vsm_prox_perna",""),6.0),
            (KNEE_H - 20,  dados_m.get("vsm_med_perna",""), 6.0),
            (KNEE_H - 30,  dados_m.get("vsm_dist_perna",""),6.0),
        ]
        for y_d, val, limiar in DIAM_VSM:
            if val and str(val).strip():
                try:
                    cor_d = '#922B21' if float(str(val).replace(",",".")) > limiar else '#2C3E50'
                    ax_med.plot([VSM_X - 0.3, VSM_X + 0.3], [y_d, y_d],
                                color=cor_d, lw=1.3, zorder=6)
                    ax_med.text(VSM_X - 0.45, y_d, f"{val}mm",
                                fontsize=6, ha='right', va='center', color=cor_d)
                except ValueError:
                    pass

    # ---- VSP ----
    VSP_X = 0.0
    vsp_d    = dados_m["vsp_dados_input"]
    vsp_temp = vsp_d.get("template", "")

    if "Ausente" in vsp_temp:
        ax_post.plot([VSP_X, VSP_X], [0, JSP_H], color='#BBBBBB',
                     lw=3, ls=':', zorder=5, alpha=0.7)
        ax_post.text(VSP_X + 0.4, JSP_H / 2, "VSP\nAusente", fontsize=8,
                     va='center', color='#AAAAAA', style='italic')
    else:
        giacomini_cranial    = "extensão cranial" in vsp_temp
        giacomini_ascendente = "ascendente" in vsp_temp
        vsp_escape           = "escape do refluxo para a veia safena parva" in vsp_temp
        vsp_incomp_tipo      = vsp_d.get("vsp_incomp_tipo", "Incompetência completa")
        vsp_incomp_seg       = (vsp_escape and vsp_incomp_tipo == "Incompetência segmentar") or vsp_d.get("tem_seg_vsp", False)
        jsp_incomp           = vsp_escape or giacomini_ascendente
        if giacomini_cranial:
            jsp_cor = '#8E44AD'
        elif jsp_incomp:
            jsp_cor = '#C0392B'
        else:
            jsp_cor = '#1565C0'

        # extensão cranial (JSP anatomicamente ausente — Giacomini como variante)
        if giacomini_cranial:
            ax_post.plot([VSP_X, VSP_X], [JSP_H, JSP_H + 17],
                         color='#8E44AD', lw=2.5, ls='--', zorder=5)
            ax_post.text(VSP_X + 0.9, JSP_H + 8, "Giacomini",
                         fontsize=7, va='center', color='#8E44AD')

        # refluxo ascendente via Giacomini (JSP incompetente, refluxo sobe)
        if giacomini_ascendente:
            ax_post.plot([VSP_X, VSP_X], [JSP_H, JSP_H + 17],
                         color='#C0392B', lw=2.5, ls='--', zorder=5)
            ax_post.annotate('', xy=(VSP_X + 0.7, JSP_H + 14),
                             xytext=(VSP_X + 0.7, JSP_H + 6),
                             arrowprops=dict(arrowstyle='->', color='#C0392B',
                                             lw=1.5, mutation_scale=13))
            ax_post.text(VSP_X + 0.9, JSP_H + 10, "Giacomini\n(↑ ascend.)",
                         fontsize=6.5, va='center', color='#C0392B')

        # tronco da VSP
        if vsp_incomp_seg:
            _dist_ref  = vsp_d.get("vsp_dist_ref", "Prega Poplítea")
            _dist_pos  = vsp_d.get("vsp_dist_pos", "abaixo")
            _desague_t = vsp_d.get("vsp_desague_tipo", "Varizes de face posterior")
            try:
                _dist_cm = float(vsp_d.get("vsp_dist_cm", "") or "10")
            except ValueError:
                _dist_cm = 10.0
            if _desague_t == "Continuidade para a região maleolar":
                y_drain = 2.0
            elif "Poplítea" in _dist_ref:
                y_drain = (JSP_H + _dist_cm) if _dist_pos == "acima" else (JSP_H - _dist_cm)
            else:
                y_drain = _dist_cm
            y_drain = max(2.0, min(float(y_drain), JSP_H - 2))
            ax_post.plot([VSP_X, VSP_X], [y_drain, JSP_H],
                         color='#C0392B', lw=4, solid_capstyle='round', zorder=5)
            ax_post.plot([VSP_X, VSP_X], [0, y_drain],
                         color='#1565C0', lw=4, solid_capstyle='round', zorder=5)
            ax_post.scatter([VSP_X], [y_drain], s=80, color='#C0392B',
                            marker='o', zorder=9, edgecolors='white', linewidths=0.8)
            ax_post.annotate('', xy=(VSP_X + 0.7, y_drain - 4),
                             xytext=(VSP_X + 0.7, y_drain + 4),
                             arrowprops=dict(arrowstyle='->', color='#C0392B',
                                             lw=1.5, mutation_scale=13))
        else:
            vsp_line_cor = '#C0392B' if (vsp_escape and not giacomini_ascendente) else '#1565C0'
            ax_post.plot([VSP_X, VSP_X], [0, JSP_H],
                         color=vsp_line_cor, lw=4, solid_capstyle='round', zorder=5)
            if vsp_escape:
                ax_post.annotate('', xy=(VSP_X + 0.7, 10),
                                 xytext=(VSP_X + 0.7, 22),
                                 arrowprops=dict(arrowstyle='->', color='#C0392B',
                                                 lw=1.5, mutation_scale=13))

        ax_post.scatter([VSP_X], [JSP_H], s=160, color=jsp_cor,
                        marker='D', zorder=10, edgecolors='white', linewidths=0.8)
        if giacomini_cranial:
            jsp_label = "JSP ausente"
        elif jsp_incomp:
            jsp_label = "JSP (INC)"
        else:
            jsp_label = "JSP (COMP)"
        ax_post.text(VSP_X + 0.4, JSP_H,
                     jsp_label, fontsize=7.5, va='center', color=jsp_cor, fontweight='bold')

        ax_post.text(VSP_X - 0.35, 18, "VSP", fontsize=8.5, ha='right',
                     va='center', color='#1A1A2E', fontweight='bold', rotation=90)

        DIAM_VSP = [
            (JSP_H,        dados_m.get("jsp_mm","")          ),
            (JSP_H - 5,    dados_m.get("vsp_crossa","")      ),
            (KNEE_H - 20,  dados_m.get("vsp_med_perna_diam","")),
        ]
        for y_d, val in DIAM_VSP:
            if val and str(val).strip():
                ax_post.plot([VSP_X - 0.3, VSP_X + 0.3], [y_d, y_d],
                             color='#2C3E50', lw=1.3, zorder=6)
                ax_post.text(VSP_X - 0.45, y_d, f"{val}mm",
                             fontsize=6, ha='right', va='center', color='#2C3E50')

    # ---- perfurantes ----
    COR_PF  = '#C0392B'
    X_DEEP  =  0.4   # lado profundo (intrafascial)
    X_SUPER = -3.4   # lado superficial (extrafascial)

    for perf in dados_m.get("perfurantes_lista", []):
        y_pf = calc_altura(
            perf.get("ref_ponto","Face Plantar"),
            perf.get("posicao_joelho","Abaixo").lower(),
            perf.get("altura_cm","0")
        )
        y_pf = max(2.0, min(float(y_pf), JSF_H - 2))

        # Trajeto da perfurante: sistema profundo → VSM → superficial (refluxo)
        ax_med.plot([X_DEEP, X_SUPER], [y_pf, y_pf],
                    color=COR_PF, lw=2, ls='--', zorder=7,
                    dash_capstyle='round')

        # Marcador onde a perfurante cruza (e alimenta) a VSM
        ax_med.scatter([VSM_X], [y_pf], s=75, color=COR_PF, marker='o',
                       zorder=8, edgecolors='white', linewidths=0.8)

        # Seta de escape: sentido profundo → superficial
        ax_med.annotate('', xy=(X_SUPER + 0.15, y_pf),
                        xytext=(X_SUPER + 0.95, y_pf),
                        arrowprops=dict(arrowstyle='->', color=COR_PF,
                                        lw=1.5, mutation_scale=12))

        # Ramos varicosos no ponto de escape superficial
        for dx, dy in [(-0.45, 2.8), (-0.25, 4.5), (-0.55, -2.2), (-0.35, -4.0)]:
            ax_med.plot([X_SUPER, X_SUPER + dx], [y_pf, y_pf + dy],
                        color=COR_PF, lw=1.8, solid_capstyle='round',
                        zorder=7, alpha=0.85)

        # Rótulo acima do ponto profundo
        diam_pf = perf.get("diametro_mm","")
        face_pf = perf.get("face","PF")
        lbl_pf  = face_pf[:3] if face_pf else "PF"
        try:
            if diam_pf:
                lbl_pf += f" {diam_pf}mm"
                if float(str(diam_pf).replace(",",".")) > 3.5:
                    lbl_pf += " ⚠"
        except ValueError:
            pass
        ax_med.text(X_DEEP + 0.15, y_pf + 1.5, lbl_pf, fontsize=6.5,
                    ha='left', va='bottom', color=COR_PF, fontweight='bold')

    # ---- SVP badge ----
    svp = dados_m.get("svp", {})
    if svp.get("status") == "Anormal":
        tipo_svp  = svp.get("tipo","Alteração")
        veias_svp = svp.get("veias",[])
        txt_svp   = f"SVP: {tipo_svp}"
        if veias_svp:
            txt_svp += "\n" + "\n".join(f"• {v}" for v in veias_svp[:3])
        for ax in (ax_med, ax_post):
            ax.text(0, JSF_H + 5.5, txt_svp, fontsize=6.5, ha='center',
                    va='bottom', color='#7B241C', fontweight='bold',
                    bbox=dict(boxstyle='round,pad=0.4', facecolor='#FADBD8',
                              edgecolor='#E74C3C', alpha=0.9, lw=1.5), zorder=15)

    # ---- outros achados ----
    # ---- varicosidades (badge) ----
    vd = dados_m.get("varic_dados", {})
    _varic_lista = vd.get("lista", [])
    if _varic_lista:
        _tipo_abrev = {"Telangiectasias (< 1 mm)": "Telang.", "Microvarizes / Varizes Reticulares (1–3 mm)": "Reticulares", "Veias Varicosas Tronculares (> 3 mm)": "Varicosas"}
        tipos_var = list(dict.fromkeys(_tipo_abrev.get(v["tipo"], v["tipo"][:10]) for v in _varic_lista))
        ax_med.text(-5.5, 6, "Varicosidades:\n" + ", ".join(tipos_var),
                    fontsize=6, ha='left', va='bottom', color='#784212',
                    bbox=dict(boxstyle='round,pad=0.4', facecolor='#FEF9E7',
                              edgecolor='#D4AC0D', alpha=0.9, lw=1))

    # ---- títulos e legenda ----
    ax_med.set_title("Vista Medial\nVeia Safena Magna (VSM)",
                     fontsize=10, fontweight='bold', color='#1A1A2E', pad=12)
    ax_post.set_title("Vista Posterior\nVeia Safena Parva (VSP)",
                      fontsize=10, fontweight='bold', color='#1A1A2E', pad=12)

    legenda = [
        Line2D([0],[0], color='#1565C0', lw=3.5,  label='Veia competente'),
        Line2D([0],[0], color='#C0392B', lw=3.5,  label='Refluxo / incompetência'),
        Line2D([0],[0], color='#BBBBBB', lw=2.5, ls=':', label='Ausente'),
        Line2D([0],[0], color='#8E44AD', lw=2.5, ls='--', label='Veia de Giacomini'),
        Line2D([0],[0], color='#C0392B', lw=2, ls='--',
               marker='o', markersize=6, label='Perfurante incompetente → escape varicoso'),
        mpatches.Patch(color='#1565C0', label='Junção competente (◆)'),
        mpatches.Patch(color='#C0392B', label='Junção incompetente (◆)'),
    ]
    fig.legend(handles=legenda, loc='lower center', ncol=4,
               fontsize=7.5, framealpha=0.95, edgecolor='#CCCCCC',
               bbox_to_anchor=(0.5, 0.005))

    fig.suptitle(
        f"CARTOGRAFIA VENOSA — MEMBRO INFERIOR {m_nome}\nPaciente: {paciente}",
        fontsize=12, fontweight='bold', color='#1A1A2E', y=0.995
    )
    plt.tight_layout(rect=[0, 0.07, 1, 0.975])
    return fig


# --- FUNÇÃO DE MONTAGEM E CONSTRUÇÃO DO DOCUMENTO WORD ---
def construir_laudo_word(membros_lista, dados_m_dict):
    doc = Document()
    doc.styles['Normal'].font.name = fonte_doc
    doc.styles['Normal'].font.size = Pt(tamanho_fonte)
    
    def add_p(text, bold_pre=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, bullet=False, italic=False):
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.alignment = align
        p.paragraph_format.line_spacing = espacamento_linhas
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if bold_pre:
            r_p = p.add_run(bold_pre)
            r_p.bold = True
        r = p.add_run(text)
        if italic:
            r.italic = True

    if nome_clinica.strip():
        p_cl = doc.add_paragraph()
        p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cl = p_cl.add_run(nome_clinica.upper())
        r_cl.bold = True
        r_cl.font.name = fonte_doc
        r_cl.font.size = Pt(tamanho_fonte + 2)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    if formato_exame == "Bilateral (Laudo Único)":
        add_p("DOS MEMBROS INFERIORES", bold_pre='DUPLEX SCAN VENOSO ', space_after=12)
    else:
        add_p(f"DO MEMBRO INFERIOR {membros_lista[0]}", bold_pre='DUPLEX SCAN VENOSO ', space_after=12)

    if nome_paciente.strip():
        add_p(f" {nome_paciente}", bold_pre="Paciente:")
    add_p("TÉCNICA", space_before=12, space_after=6)
    add_p(tecnica_texto, space_after=12)
    
    conclusoes_lista = []
    
    for m_nome in membros_lista:
        dm = dados_m_dict[m_nome]
        add_p("⸻", space_after=12)
        
        add_p(f"MEMBRO INFERIOR {m_nome}", space_after=12)
        lado = "direita" if m_nome == "DIREITO" else "esquerda"

        # 1. SVP
        add_p("SISTEMA VENOSO PROFUNDO", space_after=6)
        if dm["svp"]["status"] == "Normal":
            add_p("As veias femoral comum, femoral, poplítea, tibiais posteriores e fibulares apresentam-se pérvias, compressíveis, com fluxo fásico com a respiração e competentes.")
        else:
            add_p(f"Sistema Venoso Profundo ANORMAL. Detectados sinais de {dm['svp']['tipo']} nas veias: {', '.join(dm['svp'].get('veias', []))}.")

        # 2. SVS - VEIA SAFENA MAGNA (VSM)
        add_p("SISTEMA VENOSO SUPERFICIAL", space_before=12, space_after=6)
        vm = dm["vsm_mapeamento"]
        
        if "Ausente" in vm["status_geral"]:
            add_p(f"Veia safena magna ausente ({vm['status_geral']}).")
        else:
            if "hipoplásico" in vm["status_geral"]:
                _hipo_segs = vm.get("hipoplasico_segs", [])
                _hipo_trib = vm.get("hipoplasico_tributaria", False)
                _segs_txt = ", ".join(s.lower() for s in _hipo_segs) if _hipo_segs else "segmento não especificado"
                _hipo_cal = vm.get("hipoplasico_trib_calibre", "")
                if _hipo_trib and _hipo_cal:
                    _trib_txt = f", com tributária superficial acompanhando o segmento hipoplásico (maior calibre de {_hipo_cal} mm)"
                elif _hipo_trib:
                    _trib_txt = ", com tributária superficial acompanhando o segmento hipoplásico"
                else:
                    _trib_txt = ""
                add_p(f"A veia safena magna apresenta-se pérvia, identificando-se segmento hipoplásico em {_segs_txt}{_trib_txt}.")
                conclusoes_lista.append((m_nome, f"Hipoplasia segmentar da veia safena magna {lado} ({_segs_txt})."))
            if vm["jsf_status"] == "Competente":
                if vm.get("tronco_refluxo") and vm.get("segmentos_lista"):
                    add_p("A junção safenofemoral apresenta-se competente, sem evidências de refluxo valvar patológico.")
                else:
                    add_p("Junção safenofemoral competente. A veia safena magna segue pérvia e competente.")
            else:
                v_padrao = vm["jsf_valvulas"]
                det = vm.get("jsf_detalhes_input", {})
                
                if "Terminal e Pré-terminal" in v_padrao:
                    ref_dist = det.get("ponto_ref_dist", "Interlinha do Joelho")
                    dist_d = det.get("dist_fim", "0")
                    pos_joelho_dist = det.get("posicao_joelho_dist", "")
                    termo_dist = f"a {dist_d} cm {pos_joelho_dist.lower()} da interlinha do joelho" if "Interlinha" in ref_dist and dist_d != "0" else "na altura da interlinha do joelho" if "Interlinha" in ref_dist else f"a {dist_d} cm de distância da {ref_dist}"
                    extensao_txt = det.get("extensao_refluxo", "toda sua extensão")
                    
                    add_p(f"Refluxo originado de incompetência das válvulas pré-terminal e terminal da junção safenofemoral, com escape para a veia safena magna, que segue incompetente até {extensao_txt}. O refluxo apresenta deságue para {det.get('desague_tipo','tributárias epifasciais varicosas')} localizado {termo_dist}.")
                    conclusoes_lista.append((m_nome, f"Insuficiência segmentar do tronco da veia safena magna {lado} por incompetência da junção safenofemoral (até {extensao_txt})."))
                
                elif "Apenas a Válvula Pré-terminal" in v_padrao:
                    add_p(f"Refluxo originado de incompetência da válvula pré-terminal da junção safenofemoral, com escape de refluxo proveniente de tributária para o segmento proximal da veia safena magna que segue incompetente até cerca de {det.get('cm_ponto_j','__')} cm da junção/ponto J, onde ocorre drenagem para {det.get('tipo_drenagem','tributárias epifasciais varicosas')} formando complexos varicosos na face medial da coxa em {det.get('terco_coxa','terço médio')}.")
                    conclusoes_lista.append((m_nome, f"Insuficiência proximal da veia safena magna {lado} por incompetência da junção safenofemoral, associada a varizes calibrosas mediais na coxa."))
                
                elif "Apenas a Válvula Terminal" in v_padrao:
                    dest_term = det.get("destino_terminal", "")
                    if "Acessória Anterior" in dest_term:
                        add_p(f"Refluxo originado de incompetência da válvula terminal da junção safenofemoral com escape para o segmento proximal da veia safena acessória anterior e deságue em veias epifasciais varicosas na face {det.get('vsaa_desague','anterolateral da coxa')}. Há extensão dos trajetos varicosos até a face {det.get('vsaa_extensao','lateral do joelho')}. A veia safena magna segue competente.")
                        conclusoes_lista.append((m_nome, f"Varizes anterolaterais na coxa {lado} originárias da veia safena acessória anterior por insuficiência da válvula terminal da JSF."))
                    else:
                        add_p(f"Refluxo originado de incompetência da válvula terminal da junção safenofemoral fluindo para tributárias não safênicas na {det.get('tributarias_tipo','coxa')}. A veia safena magna segue pérvia e competente com calibre preservado.")
                        conclusoes_lista.append((m_nome, f"Varizes proximais na coxa {lado}, originárias da junção safenofemoral incompetente por falha da válvula terminal."))

            if vm["tronco_refluxo"] and vm.get("segmentos_lista"):
                def _fmt_ref_doc(ref, pos, cm):
                    if ref == "Interlinha do Joelho":
                        return f"{cm} cm {pos} da interlinha do joelho"
                    return f"{cm} cm da {ref}"
                for reg in vm["segmentos_lista"]:
                    origem = reg.get("origem", "")
                    desague = reg.get("desague", "")
                    origem_txt = f" com escape originado de {origem.lower()}" if origem and origem != "Insuficiência valvar isolada" else ""
                    if desague == "Região maleolar":
                        desague_txt = ", estendendo-se até a região maleolar"
                    elif desague:
                        desague_txt = f", com deságue para {desague.lower()}"
                    else:
                        desague_txt = ""
                    inicio = _fmt_ref_doc(reg["prox_ref"], reg["prox_pos"], reg["prox_cm"])
                    tem_fim = reg.get("dist_ref") and desague != "Região maleolar"
                    fim = _fmt_ref_doc(reg["dist_ref"], reg["dist_pos"], reg["dist_cm"]) if tem_fim else None
                    extensao_txt = f"com extensão de {inicio} até {fim}" if fim else f"a partir de {inicio}"
                    if reg["segmentos"]:
                        segs_txt = ", ".join(reg["segmentos"])
                        add_p(f"Identificado(s) segmento(s) incompetente(s) no tronco da veia safena magna ({segs_txt}){origem_txt}, {extensao_txt}{desague_txt}.")
                    else:
                        add_p(f"Insuficiência valvar do tronco da veia safena magna{origem_txt}, {extensao_txt}{desague_txt}.")
                _total_segs_vsm = len(vm["segmentos_lista"])
                if _total_segs_vsm == 1:
                    conclusoes_lista.append((m_nome, f"Insuficiência segmentar de veia safena magna {lado}."))
                elif _total_segs_vsm >= 2:
                    conclusoes_lista.append((m_nome, f"Insuficiência multissegmentar de veia safena magna {lado}."))

            # Competência residual da VSM
            if vm.get("segmentos_lista"):
                refluxo_total = any(reg.get("desague") == "Região maleolar" for reg in vm["segmentos_lista"])
                if not refluxo_total:
                    add_p("No restante do trajeto a veia safena magna segue competente.")
            elif vm["jsf_status"] == "Incompetente":
                _vp  = vm.get("jsf_valvulas", "")
                _det = vm.get("jsf_detalhes_input", {})
                if "Terminal e Pré-terminal" in _vp:
                    if _det.get("extensao_refluxo", "") != "toda sua extensão":
                        add_p("No restante do trajeto a veia safena magna segue competente.")
                elif "Apenas a Válvula Pré-terminal" in _vp:
                    add_p("No restante do trajeto a veia safena magna segue competente.")

        # Impressão das Medidas da VSM
        if "Pérvia" in vm.get("status_geral", ""):
            medidas_vsm = [("Junção safenofemoral:", dm['jsf_mm']), ("Terço proximal da coxa:", dm['vsm_prox_coxa']), ("Terço médio da coxa:", dm['vsm_med_coxa']), ("Terço distal da coxa:", dm['vsm_dist_coxa']), ("Terço proximal da perna:", dm['vsm_prox_perna']), ("Terço médio da perna:", dm['vsm_med_perna']), ("Terço distal da perna:", dm['vsm_dist_perna'])]
            medidas_vsm_ativas = [(lbl, val) for lbl, val in medidas_vsm if str(val).strip()]
            if medidas_vsm_ativas:
                add_p("Medidas da veia safena magna:", space_before=6, space_after=4)
                for lbl, val in medidas_vsm_ativas: add_p(f" {val} mm", bold_pre=lbl, bullet=True)

        # 2.2 VEIA SAFENA PARVA (VSP)
        add_p("Veia Safena Parva:", space_before=8, space_after=4)
        vsp_d = dm["vsp_dados_input"]
        vsp_txt_temp = vsp_d.get("template", "")

        if "Normal" in vsp_txt_temp:
            add_p("A veia safena parva apresenta-se pérvia, com trajeto anatômico habitual, paredes finas, totalmente compressível e competente em todo o seu trajeto.")
        elif "Ausente" in vsp_txt_temp:
            add_p(f"Veia safena parva ausente cirurgicamente ({vsp_txt_temp}).")
        elif "extensão cranial" in vsp_txt_temp:
            add_p("Junção safenopoplítea anatomicamente ausente. Observa-se extensão cranial da veia safena parva cursando no plano fascial posterior da coxa.")
            if vsp_d.get("tem_seg_vsp"):
                _escape_txt, _ext_txt, _residual = _build_seg_vsp_text(vsp_d)
                add_p(f"Identificado segmento com refluxo na veia safena parva{_escape_txt}, {_ext_txt}.{_residual}")
                conclusoes_lista.append((m_nome, f"Insuficiência segmentar da veia safena parva {lado}."))
        elif vsp_txt_temp == "Junção safenopoplítea competente":
            add_p("Junção safenopoplítea competente. A veia safena parva segue pérvia e competente em todo o seu trajeto.")
            if vsp_d.get("tem_seg_vsp"):
                _escape_txt, _ext_txt, _residual = _build_seg_vsp_text(vsp_d)
                add_p(f"Identificado segmento com refluxo na veia safena parva{_escape_txt}, {_ext_txt}.{_residual}")
                conclusoes_lista.append((m_nome, f"Insuficiência segmentar da veia safena parva {lado}."))
        elif "escape do refluxo para a veia safena parva" in vsp_txt_temp:
            incomp_tipo = vsp_d.get("vsp_incomp_tipo", "Incompetência completa")
            if incomp_tipo == "Incompetência completa":
                add_p("Junção safenopoplítea incompetente, com escape do refluxo valvar patológico para a veia safena parva, determinando incompetência troncular completa do vaso.")
                conclusoes_lista.append((m_nome, f"Insuficiência troncular completa da veia safena parva {lado} por incompetência da junção safenopoplítea."))
            else:
                _escape_txt, _ext_txt, _residual = _build_seg_vsp_text(vsp_d)
                add_p(f"Junção safenopoplítea incompetente, com escape do refluxo valvar patológico para a veia safena parva{_escape_txt}, determinando insuficiência segmentar, {_ext_txt}.{_residual}")
                conclusoes_lista.append((m_nome, f"Insuficiência segmentar da veia safena parva {lado} por incompetência da junção safenopoplítea."))
        elif "ascendente" in vsp_txt_temp:
            add_p("Junção safenopoplítea incompetente, com refluxo valvar patológico drenado de forma ascendente através da veia de Giacomini.")
            conclusoes_lista.append((m_nome, f"Insuficiência da junção safenopoplítea {lado} com refluxo ascendente via veia de Giacomini."))

        # Impressão das Medidas da VSP
        if not ("Ausente" in vsp_txt_temp):
            _jsp_ausente = "extensão cranial" in vsp_txt_temp
            medidas_vsp = [] if _jsp_ausente else [("Junção safenopoplítea:", dm['jsp_mm'])]
            medidas_vsp += [
                ("Crossa da safena parva:", dm['vsp_crossa']),
                ("Terço médio da perna:", dm['vsp_med_perna_diam'])
            ]
            medidas_vsp_ativas = [(lbl, val) for lbl, val in medidas_vsp if str(val).strip()]
            if medidas_vsp_ativas:
                add_p("Medidas da veia safena parva:", space_before=4, space_after=4)
                for lbl, val in medidas_vsp_ativas: add_p(f" {val} mm", bold_pre=lbl, bullet=True)

        # 2.3 VEIAS PERFURANTES INCOMPETENTES RELEVANTES
        if dm["perfurantes_lista"]:
            add_p("VEIAS PERFURANTES INCOMPETENTES RELEVANTES", space_before=8, space_after=4)
            for p_dados in dm["perfurantes_lista"]:
                r_reg = p_dados["regiao"]
                r_face = p_dados["face"]
                r_ref = p_dados["ref_ponto"]
                r_alt = p_dados["altura_cm"]
                r_pos = p_dados["posicao_joelho"]
                
                txt_ref_perf = f"localizada {r_pos.lower()} da interlinha do joelho" if "Interlinha" in r_ref else f"a partir da face plantar"
                r_diam = p_dados.get("diametro_mm", "")
                txt_diam = f", com diâmetro de {r_diam} mm" if r_diam else ""
                add_p(f"Identificada veia perfurante incompetente na {r_reg.lower()}, face {r_face.lower()}{txt_diam}, medindo {r_alt} cm de altura {txt_ref_perf}, determinando escape para tributárias varicosas.", bullet=True)
                conclusoes_lista.append((m_nome, f"Insuficiência de veia perfurante na {r_reg.lower()} (face {r_face.lower()})."))

        # 2.4 MAPEAMENTO DE VARICOSIDADES
        vd = dm["varic_dados"]
        _tipo_full = {
            "Telangiectasias (< 1 mm)": "telangiectasias (< 1 mm de diâmetro)",
            "Microvarizes / Varizes Reticulares (1–3 mm)": "microvarizes / varizes reticulares (1 a 3 mm de diâmetro)",
            "Veias Varicosas Tronculares (> 3 mm)": "veias varicosas tronculares (> 3 mm de diâmetro)",
        }
        for _vitem in vd.get("lista", []):
            _t = _tipo_full.get(_vitem["tipo"], _vitem["tipo"].lower())
            _l = _vitem.get("localizacao", "")
            _loc_txt = f" em {_l}" if _l else ""
            add_p(f"Presença de {_t}{_loc_txt}.", space_before=6)

        # 2.5 VARIZES EXTRASSAFÊNICAS
        _PELVICO_TXTS = {
            "Ponto inguinal": (
                "Observa-se fluxo retrógrado originado em topografia de anel inguinal superficial "
                "(Ponto de Escape Inguinal - Ponto I). O referido refluxo estende-se caudalmente "
                "para a raiz da coxa, atuando como fonte nutridora para varizes em região inguinal "
                "anterior e para a face anterior da coxa."
            ),
            "Ponto perineal": (
                "Identificado ponto de escape de origem pélvica em região perineal/vulvar (Ponto P), "
                "secundário ao refluxo de ramos da veia pudenda interna. O fluxo retrógrado calibroso "
                "estende-se para a face medial superior da coxa, alimentando tributárias varicosas superficiais."
            ),
            "Ponto obturatório": (
                "Evidenciado refluxo venoso emergindo pelo forame obturatório (Ponto de Escape "
                "Obturatório - Ponto O), direcionando fluxo retrógrado para a face medial e profunda "
                "do terço superior da coxa. Nota-se conexão deste ponto com feixe de varizes isoladas "
                "e não-safênicas neste compartimento, sem sinais de comunicação direta com o tronco "
                "da veia safena magna."
            ),
            "Ponto glúteo": (
                "Demonstrado ponto de fuga venoso trans-pélvico emergindo através do forame isquiático "
                "(Ponto de Escape Glúteo - Ponto G). O refluxo manifesta-se em região infra-glútea e "
                "direciona-se para a face posterior da coxa, atuando como fonte hemodinâmica para varizes "
                "locais e propagando-se em direção à extensão posterior de coxa."
            ),
        }
        ved = dm.get("varizes_extrassaf_dados", {})
        for _veitem in ved.get("lista", []):
            _ve_loc    = _veitem.get("localizacao", "")
            _ve_origem = _veitem.get("origem", "Tributária incompetente")
            if _ve_origem == "Tributária incompetente":
                _trib_ref = _veitem.get("trib_ref", "")
                _trib_pos = _veitem.get("trib_pos", "")
                _trib_cm  = _veitem.get("trib_cm", "")
                if _trib_cm and _trib_ref:
                    if _trib_ref == "Interlinha do Joelho" and _trib_pos:
                        _altura_txt = f", localizada a {_trib_cm} cm {_trib_pos} da interlinha do joelho"
                    elif _trib_ref == "Junção Safenofemoral":
                        _altura_txt = f", localizada a {_trib_cm} cm da junção safenofemoral"
                    elif _trib_ref == "Face Plantar":
                        _altura_txt = f", localizada a {_trib_cm} cm da face plantar"
                    else:
                        _altura_txt = f", localizada a {_trib_cm} cm da interlinha do joelho"
                else:
                    _altura_txt = ""
                _loc_txt = f" em {_ve_loc}" if _ve_loc else ""
                add_p(f"Identificam-se varizes extrassafênicas{_loc_txt}, com ponto de escape hemodinâmico em tributária incompetente{_altura_txt}.", space_before=6)
                conclusoes_lista.append((m_nome, f"Varizes extrassafênicas{_loc_txt} por tributária incompetente."))
            elif _ve_origem == "Refluxo de origem ciática":
                _subtipo = _veitem.get("ciatica_subtipo", "Varizes acompanhando o trajeto do nervo ciático")
                if _subtipo == "Veia ciática persistente":
                    add_p(f"Identificam-se varizes extrassafênicas em {_ve_loc}, originadas a partir de veia ciática persistente.", space_before=6)
                    conclusoes_lista.append((m_nome, f"Varizes extrassafênicas em {_ve_loc} por veia ciática persistente."))
                else:
                    add_p(f"Identificam-se varizes extrassafênicas em {_ve_loc}, acompanhando o trajeto do nervo ciático, compatíveis com refluxo de origem ciática.", space_before=6)
                    conclusoes_lista.append((m_nome, f"Varizes extrassafênicas em {_ve_loc} acompanhando o trajeto do nervo ciático."))
            else:
                _pontos = _veitem.get("pelvico_pontos", [])
                if _pontos:
                    for _pt in _pontos:
                        _pt_txt = _PELVICO_TXTS.get(_pt)
                        if _pt_txt:
                            add_p(_pt_txt, space_before=6)
                    conclusoes_lista.append((m_nome, f"Varizes extrassafênicas em {_ve_loc} por refluxo pélvico ({', '.join(_pontos)})."))
                else:
                    add_p(f"Identificam-se varizes extrassafênicas em {_ve_loc}, com origem em refluxo de origem pélvica.", space_before=6)
                    conclusoes_lista.append((m_nome, f"Varizes extrassafênicas em {_ve_loc} por refluxo pélvico."))
                add_p(
                    "O padrão de distribuição do refluxo nos membros inferiores sugere fortemente de origem "
                    "pélvica, sustentado pela patência e competência das junções safeno-femorais e "
                    "identificação dos pontos de escape descritos. Sugere-se correlação com propedêutica "
                    "de imagem pélvica (Ultrassom transvaginal com Doppler ou Angiotomografia/Angioressonância "
                    "de pelve).",
                    space_before=4, italic=True
                )

        # Pós-Operatório
        if "6.3" in dm["pos_op_opt"]:
            add_p("Sinais de neovascularização adjacentes à região operada (recidiva pós-operatória).", space_before=8)

        # 5. ACHADOS ADICIONAIS (Cisto Baker / Edema)
        if dm["achados_adi_multi"]:
            add_p("ACHADOS ADICIONAIS / ESTRUTURAS ADJACENTES", space_before=10, space_after=4)
            for achado in dm["achados_adi_multi"]:
                if "Cisto de Baker" in achado:
                    cm = dm.get("cisto_medidas", {})
                    txt_cisto = f"Identificado Cisto de Baker na fossa poplítea medindo {cm.get('comp','__')} x {cm.get('larg','__')} x {cm.get('esp','__')} mm (Eixo Longitudinal x Transversal x Espessura)."
                    add_p(txt_cisto, bullet=True)
                    conclusoes_lista.append((m_nome, "Cisto de Baker na fossa poplítea."))
                if "Edema intersticial" in achado:
                    add_p("Sinais ultrassonográficos compatíveis com edema intersticial / infiltração líquida no tecido subcutâneo.", bullet=True)
                    conclusoes_lista.append((m_nome, "Edema intersticial subcutâneo no membro avaliado."))

    # IMPRESSÃO DIAGNÓSTICA (CONCLUSÃO)
    if quebrar_pagina_diag:
        doc.add_page_break()
    add_p("⸻", space_after=12)
    add_p("IMPRESSÃO DIAGNÓSTICA", space_after=6)
    if not conclusoes_lista:
        add_p("Ausência de trombose venosa profunda.", bullet=True)
        add_p("Veias safenas competentes.", bullet=True)
    else:
        vistas = set()
        conclusoes_unicas = []
        for m_origem, c_txt in conclusoes_lista:
            chave = (m_origem, c_txt)
            if chave not in vistas:
                vistas.add(chave)
                conclusoes_unicas.append(chave)
                
        for m_origem, conclusao_txt in conclusoes_unicas:
            prefixo = f"[{m_origem}] " if formato_exame == "Bilateral (Laudo Único)" else ""
            add_p(f"{prefixo}{conclusao_txt}", bullet=True)

    if incluir_identidade and (nome_medico or crm_medico):
        doc.add_paragraph().paragraph_format.space_before = Pt(25)
        p_assin = doc.add_paragraph()
        p_assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_assin.paragraph_format.line_spacing = espacamento_linhas
        if nome_medico:
            r = p_assin.add_run(f"{nome_medico}\n")
            r.bold = True
        if crm_medico:
            p_assin.add_run(f"CRM-{crm_uf} {crm_medico}\n")
        if rqe_medico.strip():
            p_assin.add_run(f"RQE {rqe_medico}")
    return doc

# --- PROCESSAMENTO DO BOTÃO DE GERAR LAUDO ---
def _exibir_doc(doc, buf, nome_arquivo, label_download):
    st.success("Laudo gerado com sucesso!")
    if modo_saida in ["Somente Visualização", "Visualização + DOCX"]:
        texto_viz = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        st.markdown("## 👁️ Visualização do Laudo")
        st.text_area("Laudo Gerado", value=texto_viz, height=600)
    if modo_saida in ["Somente DOCX", "Visualização + DOCX"]:
        st.download_button(label_download, buf, nome_arquivo,
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)

col_btn1, col_btn2 = st.columns(2)

with col_btn1:
    gerar_laudo_btn = st.button("🚀 Gerar Laudo Venoso Completo", use_container_width=True)

with col_btn2:
    gerar_cart_btn = st.button("🗺️ Gerar Cartografia Venosa", use_container_width=True, type="secondary")

if gerar_laudo_btn:
    if formato_exame == "Bilateral (Laudos Separados)":
        for m_nome in ["DIREITO", "ESQUERDO"]:
            doc_sep = construir_laudo_word([m_nome], dados_membros)
            buf = BytesIO()
            doc_sep.save(buf)
            buf.seek(0)
            _exibir_doc(doc_sep, buf, f"Laudo_Venoso_MMII_{m_nome}.docx",
                        f"📥 Baixar Laudo Membro {m_nome} (.docx)")
    else:
        doc_unico = construir_laudo_word(membros_para_processar, dados_membros)
        buf = BytesIO()
        doc_unico.save(buf)
        buf.seek(0)
        _exibir_doc(doc_unico, buf, "Laudo_Venoso_MMII.docx", "📥 Baixar Laudo Venoso (.docx)")

if gerar_cart_btn:
    import matplotlib.pyplot as plt
    for m_nome in membros_para_processar:
        if m_nome not in dados_membros:
            st.warning(f"Dados do membro {m_nome} não encontrados.")
            continue
        with st.spinner(f"Gerando cartografia do membro {m_nome}..."):
            fig = gerar_cartografia_venosa(m_nome, dados_membros[m_nome], nome_paciente)
            buf_img = BytesIO()
            fig.savefig(buf_img, format='png', dpi=150, bbox_inches='tight',
                        facecolor=fig.get_facecolor())
            buf_img.seek(0)
            plt.close(fig)
        st.markdown(f"### 🗺️ Cartografia Venosa — Membro {m_nome}")
        st.image(buf_img, use_container_width=True)
        buf_img.seek(0)
        st.download_button(
            label=f"📥 Baixar Cartografia {m_nome} (.png)",
            data=buf_img,
            file_name=f"Cartografia_Venosa_MII_{m_nome}.png",
            mime="image/png",
            key=f"dl_cart_{m_nome}_{_rc}",
            use_container_width=True
        )