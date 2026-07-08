# arterial.py
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- CONFIGURACAO DA PAGINA ---
st.title("🫀 Assistente de Laudos: Duplex Scan Arterial de MMII")

# --- SIDEBAR: CONFIGURACOES ---
with st.sidebar:
    st.markdown("## ⚙️ Painel de Controle")
    fonte_doc = st.selectbox("Fonte do Documento:", ["Arial", "Calibri", "Times New Roman"])
    tamanho_fonte = st.slider("Tamanho Texto (pt):", 10, 14, 11)
    espacamento_linhas = st.slider("Espaçamento:", 1.0, 1.5, 1.15)
    modo_saida = st.radio(
        "Modo de saída do laudo:",
        ["Somente DOCX", "Somente Visualização", "Visualização + DOCX"],
        index=2
    )

    st.markdown("---")
    st.markdown("### ✍️ Identidade & Assinatura")
    incluir_identidade = st.toggle("Incluir identidade no laudo", value=False)
    nome_clinica = st.text_input("Cabeçalho / Nome da Clínica:", placeholder="Ex: Instituto de Diagnóstico Vascular")
    nome_medico = st.text_input("Nome do Médico:", "Lucas Santos Guimarães")
    colcrm1, colcrm2 = st.columns([2, 1])
    with colcrm1:
        crm_medico = st.text_input("CRM:", "4061")
    with colcrm2:
        crm_uf = st.selectbox("UF", ["AC","AL","AP","AM","BA","CE","DF","ES","GO","MA","MT","MS","MG","PA","PB","PR","PE","PI","RJ","RN","RS","RO","RR","SC","SP","SE","TO"], index=25)
    rqe_medico = st.text_input("RQE:", "")

# --- IDENTIFICACAO DO PACIENTE ---
nome_paciente = st.text_input("Nome do Paciente:", "")
formato_exame = st.selectbox("Tipo de Exame:", ["Unilateral", "Bilateral (Laudos Separados)", "Bilateral (Laudo Único)"])

if formato_exame == "Unilateral":
    lado_sel = st.selectbox("Selecione o Lado Avaliado:", ["DIREITO", "ESQUERDO"])
    membros_para_processar = [lado_sel]
else:
    membros_para_processar = ["DIREITO", "ESQUERDO"]

# --- LISTA E ORDEM ANATOMICA DAS ARTERIAS ---
ARTERIAS_LISTA = [
    ("femoral comum", "AFC"),
    ("femoral profunda", "AFP"),
    ("femoral superficial", "AFS"),
    ("poplítea", "APOP"),
    ("tibial anterior", "ATA"),
    ("tibial posterior", "ATP"),
    ("fibular", "AFIB")
]

# Dicionário de ordem anatômica para guiar a propagação distal em cascata
ORDEM_ANATOMICA = ["AFC", "AFP", "AFS", "APOP", "ATA", "ATP", "AFIB"]

# Artérias tibiais e fibular não propagam fluxo alterado (não há vasos distais mapeados)
PODE_PROPAGAR = {"AFC", "AFP", "AFS", "APOP"}

PADROES_ONDA_SITIO = [
    "Trifásico",
    "Bifásico",
    "Monofásico de alta resistência",
    "Monofásico de baixa resistência",
    "Monofásico sem diástole",
    "Monofásico tardus parvus",
]

# Mesma lista de padrões: garante que qualquer valor propagado por cascata
# hemodinâmica seja sempre uma opção válida no seletor "Padrão de Onda" das artérias distais
PADROES_POS_ESTENOTICO = PADROES_ONDA_SITIO

# Descrição padronizada do padrão de fluxo para o texto do laudo (todas as artérias seguem o mesmo modelo)
DESCRICOES_FLUXO = {
    "Trifásico": "anterógrado, hiperresistente, com onda espectral multifásica",
    "Bifásico": "anterógrado, de resistência intermediária, com onda espectral bifásica",
    "Monofásico de alta resistência": "anterógrado, hiperresistente, com onda espectral monofásica",
    "Monofásico de baixa resistência": "anterógrado, de baixa resistência, com onda espectral monofásica",
    "Monofásico sem diástole": "anterógrado, com onda espectral monofásica, sem componente diastólico",
    "Monofásico tardus parvus": "anterógrado, com onda espectral monofásica de morfologia tardus parvus",
}

def descrever_fluxo(padrao_onda):
    return DESCRICOES_FLUXO.get(padrao_onda, DESCRICOES_FLUXO["Trifásico"])

# Ordem crescente de gravidade hemodinâmica, usada para apurar a lesão mais
# alterada dentro de uma série de estenoses consecutivas (em tandem)
ORDEM_GRAVIDADE_FLUXO = {
    "Trifásico": 0,
    "Bifásico": 1,
    "Monofásico de alta resistência": 2,
    "Monofásico de baixa resistência": 3,
    "Monofásico sem diástole": 4,
    "Monofásico tardus parvus": 4,
}

def pior_padrao_fluxo(padroes):
    if not padroes:
        return "Trifásico"
    return max(padroes, key=lambda p: ORDEM_GRAVIDADE_FLUXO.get(p, 0))

def formatar_segmentos(segmentos):
    if not segmentos:
        return "segmento avaliado"
    nomes = [s.lower() for s in segmentos]
    if len(nomes) == 1:
        return nomes[0]
    return f"{', '.join(nomes[:-1])} e {nomes[-1]}"

SEGMENTOS_ARTERIA_PADRAO = ["Segmento Proximal", "Segmento Médio", "Segmento Distal", "Todo o trajeto"]
SEGMENTOS_ARTERIA_POPLITEA = [
    "Segmento Supra-articular",
    "Segmento Médio (Nível da Interlinha Articular)",
    "Segmento Distal (Infra-articular)",
    "Todo o trajeto"
]

def obter_segmentos_estenose(art_id):
    if art_id == "APOP":
        return SEGMENTOS_ARTERIA_POPLITEA
    return SEGMENTOS_ARTERIA_PADRAO

TERCOS_OCLUSAO_PADRAO = ["Terço Proximal", "Terço Médio", "Terço Distal"]
TERCOS_OCLUSAO_AFS = ["Terço Proximal", "Terço Médio", "Terço Distal (Canal dos Adutores)"]
SEGMENTOS_OCLUSAO_POPLITEA = [
    "Segmento Supra-articular",
    "Segmento Médio (Nível da Interlinha Articular)",
    "Segmento Distal (Infra-articular)"
]

def obter_segmentos_oclusao(art_id):
    if art_id == "AFS":
        return TERCOS_OCLUSAO_AFS
    if art_id == "APOP":
        return SEGMENTOS_OCLUSAO_POPLITEA
    return TERCOS_OCLUSAO_PADRAO

dados_membros = {}

# --- INTERFACE DE ENTRADA DE DADOS ---
if len(membros_para_processar) > 1:
    abas = st.tabs(["🔴 Membro Inferior Direito (MID)", "🔵 Membro Inferior Esquerdo (MIE)"])
else:
    abas = [st.container()]

for idx, m_nome in enumerate(membros_para_processar):
    with abas[idx]:
        st.markdown(f"### 📋 Mapeamento Arterial - Membro {m_nome}")

        # Primeiro passo: Coleta e armazenamento inicial de inputs de forma limpa
        dados_brutos_membro = {}

        # Estrutura auxiliar para rastrear se alguma artéria proximal disparou propagação em cascata
        propagacoes_ativas = {}

        for art_nome, art_id in ARTERIAS_LISTA:
            st.markdown(f"**📍 Artéria {art_nome.title()}**")

            c1, c2, c3, c4 = st.columns([1.5, 1.5, 1.8, 2.2])

            with c1:
                status = st.selectbox(
                    "Status da Artéria:",
                    ["Normal", "Estenose Focal", "Estenoses Consecutivas", "Ocluído"],
                    key=f"status_{art_id}_{m_nome}"
                )

                # Input de Ateromatose Difusa (Parietal)
                tem_ateromatose = st.checkbox("Ateromatose difusa?", value=False, key=f"ateroma_{art_id}_{m_nome}")
                desc_ateromatose = ""
                if tem_ateromatose:
                    desc_ateromatose = st.radio(
                        "Padrão da ateromatose:",
                        ["calcificações parietais multissegmentares", "placas de ateroma calcificadas multissegmentares"],
                        key=f"tipo_ateroma_{art_id}_{m_nome}"
                    )

            is_estenose_focal = status == "Estenose Focal"
            is_estenose_consecutiva = status == "Estenoses Consecutivas"
            is_ocluido = status == "Ocluído"

            pvs_pre = 0.0
            pvs_max = 0.0
            pvs_distal = 0.0
            razao_v = 0.0

            with c2:
                if is_estenose_focal:
                    pvs_pre = st.number_input("PVS Pré-Estenótico (cm/s):", min_value=1.0, max_value=600.0, value=60.0, step=5.0, key=f"pvs_pre_{art_id}_{m_nome}")
                    pvs_max = st.number_input("PVS Maior Estreitamento (cm/s):", min_value=1.0, max_value=600.0, value=150.0, step=5.0, key=f"pvs_max_{art_id}_{m_nome}")
                    pvs_distal = st.number_input("PVS Distal à Estenose (cm/s):", min_value=0.0, max_value=600.0, value=40.0, step=5.0, key=f"pvs_distal_{art_id}_{m_nome}")
                elif is_estenose_consecutiva:
                    st.caption("Defina cada lesão da série na coluna ao lado →")
                elif not is_ocluido:
                    val_sugerido = 90.0 if "femoral" in art_nome else 60.0
                    pvs_max = st.number_input("PVS (cm/s):", min_value=0.0, max_value=600.0, value=val_sugerido, key=f"pvs_max_{art_id}_{m_nome}")
                else:
                    st.markdown("<p style='color:red; margin-top:20px; font-weight:bold;'>Oclusão Luminal</p>", unsafe_allow_html=True)

            with c3:
                seg_afetado = ""
                onda_pos = ""
                origem_reenchimento = ""
                direcao_fluxo = ""
                onda_reenchimento = ""
                reenchimento_colat = False
                segs_ocluidos = []
                seg_revasc = ""
                lesoes_consecutivas = []

                if is_estenose_focal:
                    seg_afetado = st.selectbox(
                        "Localização da Estenose:",
                        obter_segmentos_estenose(art_id),
                        key=f"seg_{art_id}_{m_nome}"
                    )

                    onda_pos = st.selectbox(
                        "Fluxo Distal (Pós-Estenótico):",
                        PADROES_POS_ESTENOTICO,
                        key=f"ondapos_{art_id}_{m_nome}"
                    )
                elif is_estenose_consecutiva:
                    lesoes_key = f"lesoes_consec_{art_id}_{m_nome}"
                    if lesoes_key not in st.session_state:
                        st.session_state[lesoes_key] = []
                    if st.session_state[lesoes_key]:
                        for li, lesao in enumerate(st.session_state[lesoes_key]):
                            lc1, lc2 = st.columns([5, 1])
                            with lc1:
                                st.caption(f"`{li+1:02d}` {lesao['segmento']} — fluxo distal {lesao['fluxo'].lower()}")
                            with lc2:
                                if st.button("❌", key=f"rem_lesao_{art_id}_{m_nome}_{li}"):
                                    st.session_state[lesoes_key].pop(li)
                                    st.rerun()
                    novo_seg_lesao = st.selectbox(
                        "Segmento da Lesão:",
                        obter_segmentos_estenose(art_id),
                        key=f"novo_seg_lesao_{art_id}_{m_nome}"
                    )
                    novo_fluxo_lesao = st.selectbox(
                        "Padrão de Fluxo Distal:",
                        PADROES_ONDA_SITIO,
                        key=f"novo_fluxo_lesao_{art_id}_{m_nome}"
                    )
                    if st.button("💾 Adicionar Lesão", key=f"add_lesao_{art_id}_{m_nome}"):
                        st.session_state[lesoes_key].append({"segmento": novo_seg_lesao, "fluxo": novo_fluxo_lesao})
                        st.rerun()
                    lesoes_consecutivas = st.session_state[lesoes_key]
                elif is_ocluido:
                    segmentos_oclusao_opcoes = obter_segmentos_oclusao(art_id)
                    segs_ocluidos = st.multiselect(
                        "Segmento(s) Ocluído(s):",
                        segmentos_oclusao_opcoes,
                        default=segmentos_oclusao_opcoes,
                        key=f"segs_ocl_{art_id}_{m_nome}"
                    )
                    reenchimento_colat = st.checkbox("Há reenchimento distal?", value=True, key=f"colat_{art_id}_{m_nome}")
                    if reenchimento_colat:
                        seg_revasc = st.selectbox(
                            "Segmento da Revascularização:",
                            segmentos_oclusao_opcoes,
                            key=f"seg_revasc_{art_id}_{m_nome}"
                        )
                        origem_reenchimento = st.selectbox(
                            "Origem do Reenchimento:",
                            ["por artéria colateral", "por fluxo retrógrado de ramo arterial"],
                            key=f"orig_reench_{art_id}_{m_nome}"
                        )
                        direcao_fluxo = st.selectbox(
                            "Direção do Fluxo:",
                            ["anterógrada", "retrógrada"],
                            key=f"dir_fluxo_{art_id}_{m_nome}"
                        )
                        onda_reenchimento = st.selectbox(
                            "Padrão de Onda no Reenchimento:",
                            PADROES_ONDA_SITIO,
                            key=f"onda_reench_{art_id}_{m_nome}"
                        )
                else:
                    # Se alguma artéria proximal propagou fluxo alterado em cascata, força
                    # o valor deste seletor antes de renderizá-lo para refletir automaticamente
                    idx_atual_onda = ORDEM_ANATOMICA.index(art_id)
                    onda_forcada = None
                    for vaso_prop_id, padrao_forcado in propagacoes_ativas.items():
                        if ORDEM_ANATOMICA.index(vaso_prop_id) < idx_atual_onda:
                            onda_forcada = padrao_forcado
                    if onda_forcada:
                        st.session_state[f"onda_{art_id}_{m_nome}"] = onda_forcada
                    onda_sitio = st.selectbox("Padrão de Onda:", PADROES_ONDA_SITIO, key=f"onda_{art_id}_{m_nome}")

            with c4:
                comp_placa = ""
                grau_estenose_estimado = ""
                propagar_fluxo_distal = False

                if is_estenose_focal:
                    if pvs_pre > 0:
                        razao_v = pvs_max / pvs_pre

                    if razao_v < 2.0:
                        grau_estenose_estimado = "Estenose leve (< 50%)"
                    elif 2.0 <= razao_v <= 4.0:
                        grau_estenose_estimado = "Estenose moderada (50-70%)"
                    else:
                        grau_estenose_estimado = "Estenose severa (> 70%)"

                    st.markdown(f"""
                    <div style='background-color:#fff3cd; padding:6px; border-radius:5px; border-left:4px solid #ffc107; font-size:13px;'>
                        <b>Razão:</b> {razao_v:.2f} | <b>{grau_estenose_estimado}</b>
                    </div>
                    """, unsafe_allow_html=True)

                    # Tipo histológico/composição da placa causadora
                    comp_placa = st.radio(
                        "Composição da Placa:",
                        ["Placa calcificada", "Placa não calcificada"],
                        key=f"comp_placa_{art_id}_{m_nome}"
                    )

                    # Caixa para cascata hemodinâmica total se a estenose for de moderada a severa (Razão >= 2)
                    if razao_v >= 2.0 and art_id in PODE_PROPAGAR:
                        propagar_fluxo_distal = st.checkbox(
                            "🛒 Propagar fluxo alterado para todas as artérias distais?",
                            value=False,
                            key=f"prop_distal_{art_id}_{m_nome}"
                        )
                        if propagar_fluxo_distal:
                            propagacoes_ativas[art_id] = onda_pos

                elif is_estenose_consecutiva:
                    if lesoes_consecutivas:
                        pior_padrao = pior_padrao_fluxo([l["fluxo"] for l in lesoes_consecutivas])
                        if art_id in PODE_PROPAGAR and pior_padrao != "Trifásico":
                            propagar_fluxo_distal = st.checkbox(
                                "🛒 Propagar fluxo alterado para todas as artérias distais?",
                                value=False,
                                key=f"prop_distal_consec_{art_id}_{m_nome}"
                            )
                            if propagar_fluxo_distal:
                                propagacoes_ativas[art_id] = pior_padrao
                    else:
                        st.info("Adicione ao menos uma lesão da série.")

                elif is_ocluido:
                    if reenchimento_colat:
                        st.info(f"Reenchimento ativo {origem_reenchimento}.")
                        if art_id in PODE_PROPAGAR and onda_reenchimento != "Trifásico":
                            propagar_fluxo_distal = st.checkbox(
                                "🛒 Propagar fluxo alterado para todas as artérias distais?",
                                value=False,
                                key=f"prop_distal_reench_{art_id}_{m_nome}"
                            )
                            if propagar_fluxo_distal:
                                propagacoes_ativas[art_id] = onda_reenchimento
                    else:
                        st.warning("Ausência completa de fluxo.")
                else:
                    st.write("")

                # Alocando dados coletados na memória temporária do loop
                dados_brutos_membro[art_id] = {
                    "nome": art_nome,
                    "status": status,
                    "tem_ateromatose": tem_ateromatose,
                    "desc_ateromatose": desc_ateromatose,
                    "pvs_pre": pvs_pre,
                    "pvs_max": pvs_max,
                    "pvs_distal": pvs_distal,
                    "razao_v": razao_v,
                    "seg_afetado": seg_afetado,
                    "grau_estenose": grau_estenose_estimado,
                    "comp_placa": comp_placa,
                    "onda_sitio": onda_sitio if status == "Normal" else "",
                    "onda_pos": onda_pos,
                    "reenchimento_colat": reenchimento_colat,
                    "origem_reenchimento": origem_reenchimento,
                    "direcao_fluxo": direcao_fluxo,
                    "onda_reenchimento": onda_reenchimento,
                    "propagar_fluxo_distal": propagar_fluxo_distal,
                    "segs_ocluidos": segs_ocluidos,
                    "seg_revasc": seg_revasc,
                    "lesoes_consecutivas": lesoes_consecutivas
                }
            st.markdown("<hr style='margin: 8px 0px; border-top: 1px dashed #ddd;' />", unsafe_allow_html=True)

        # SEGUNDO PASSO: Processamento pós-interatividade para aplicar a cascata hemodinâmica distal
        dados_finais_membro = {}
        for art_id in ORDEM_ANATOMICA:
            info_vaso = dados_brutos_membro[art_id]

            # Varre se alguma artéria acima na árvore anatômica exigiu propagação forçada
            idx_atual = ORDEM_ANATOMICA.index(art_id)
            for vaso_prop_id, padrao_forcado in propagacoes_ativas.items():
                idx_prop = ORDEM_ANATOMICA.index(vaso_prop_id)
                # Se o vaso atual está abaixo (distal) do vaso causador da cascata e está como "Normal"
                if idx_atual > idx_prop and info_vaso["status"] == "Normal":
                    info_vaso["onda_sitio"] = padrao_forcado  # Altera dinamicamente o fluxo livre
                    info_vaso["detalhes_cascata"] = f"Apresenta padrão de fluxo modificado secundário à lesão proximal na artéria {dados_brutos_membro[vaso_prop_id]['nome']}."

            dados_finais_membro[art_id] = info_vaso

        dados_membros[m_nome] = dados_finais_membro

st.markdown("---")

# --- FUNCAO DE MONTAGEM DO DOCUMENTO WORD ---
def construir_laudo_arterial_word(membros_lista, dados_m_dict):
    doc = Document()
    doc.styles['Normal'].font.name = fonte_doc
    doc.styles['Normal'].font.size = Pt(tamanho_fonte)

    def add_p(text, bold_pre=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, bullet=False):
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.alignment = align
        p.paragraph_format.line_spacing = espacamento_linhas
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if bold_pre:
            r_p = p.add_run(bold_pre)
            r_p.bold = True
        p.add_run(text)

    if nome_clinica.strip():
        p_cl = doc.add_paragraph()
        p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cl = p_cl.add_run(nome_clinica.upper())
        r_cl.bold = True
        r_cl.font.name = fonte_doc
        r_cl.font.size = Pt(tamanho_fonte + 2)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    sufixo_exame = "DOS MEMBROS INFERIORES" if formato_exame == "Bilateral (Laudo Único)" else f"DO MEMBRO INFERIOR {membros_lista[0]}"
    titulo_exame = f"DUPLEX SCAN ARTERIAL {sufixo_exame}"
    add_p("", bold_pre=titulo_exame, space_after=12)
    if nome_paciente.strip():
        add_p(f" {nome_paciente}", bold_pre="Paciente:")

    add_p("TÉCNICA", space_before=12, space_after=6)
    add_p("Exame realizado com mapeamento duplex colorido e análise espectral das velocidades sistólicas, utilizando transdutor linear de alta frequência. Paciente avaliado em decúbito dorsal horizontal.", space_after=12)

    conclusoes_lista = []

    for m_nome in membros_lista:
        add_p("⸻", space_after=12)
        if formato_exame == "Bilateral (Laudo Único)":
            add_p(f"RELATÓRIO TÉCNICO – MEMBRO INFERIOR {m_nome}", space_after=12)
        else:
            add_p("RELATÓRIO TÉCNICO", space_after=12)

        m_dados = dados_m_dict[m_nome]

        for art_id, info in m_dados.items():
            txt_art = f"Artéria {info['nome']}, "

            # Injeção de texto base sobre a parede (Ateromatose difusa) se ativa
            prefixo_parede = ""
            if info["tem_ateromatose"]:
                prefixo_parede = f"com evidências de ateromatose difusa caracterizada por {info['desc_ateromatose']}, "

            if info["status"] == "Normal":
                txt_art += f"{prefixo_parede}pérvia, com trajeto regular, apresentando fluxo {descrever_fluxo(info['onda_sitio'])}."
                tem_cascata = "detalhes_cascata" in info
                if tem_cascata:
                    txt_art += f" {info['detalhes_cascata']}"
                if not tem_cascata and info["onda_sitio"] != "Trifásico":
                    conclusoes_lista.append((m_nome, f"Padrão de fluxo {info['onda_sitio'].lower()} na artéria {info['nome']}, sem lesão estenosante focal identificada."))

            elif info["status"] == "Ocluído":
                segs_ocl_texto = formatar_segmentos(info["segs_ocluidos"])
                txt_art += f"{prefixo_parede}ocluída no {segs_ocl_texto}, com completa ausência de sinal ao mapeamento Doppler colorido e espectral."
                if info["reenchimento_colat"]:
                    seg_revasc_texto = info["seg_revasc"].lower()
                    txt_art += (f" Nota-se reenchimento no {seg_revasc_texto}, alimentado {info['origem_reenchimento']}, "
                                f"apresentando direção de fluxo {info['direcao_fluxo']} e padrão de onda {info['onda_reenchimento'].lower()}.")
                    conclusoes_lista.append((m_nome, f"Oclusão do {segs_ocl_texto} da artéria {info['nome']} com reenchimento no {seg_revasc_texto}, alimentado {info['origem_reenchimento']} ({info['direcao_fluxo']}/{info['onda_reenchimento']})."))
                else:
                    txt_art += " Não há sinais de reenchimento arterial distal evidente por circulação colateral significativa."
                    conclusoes_lista.append((m_nome, f"Oclusão do {segs_ocl_texto} da artéria {info['nome']} sem sinais de reenchimento distal."))

            elif info["status"] == "Estenose Focal":
                loc_texto = info['seg_afetado'].lower()
                txt_art += (f"{prefixo_parede}pérvia, identificando-se lesão estenosante focal causada por uma {info['comp_placa'].lower()} localizada no {loc_texto}, caracterizando {info['grau_estenose'].lower()}. "
                            f"Ao Doppler espectral, registra-se PVS pré-estenótica de {info['pvs_pre']:.1f} cm/s, PVS no ponto de maior estreitamento de {info['pvs_max']:.1f} cm/s "
                            f"(Razão de Velocidades: {info['razao_v']:.2f}) e PVS distal de {info['pvs_distal']:.1f} cm/s.")

                if info["onda_pos"]:
                    txt_art += f" O padrão de fluxo no segmento imediatamente pós-estenótico exibe morfologia {info['onda_pos'].lower()}."

                if info["grau_estenose"] != "Estenose leve (< 50%)":
                    sufixo_concl = " com repercussão em cascata nos vasos distais" if info["propagar_fluxo_distal"] else ""
                    conclusoes_lista.append((m_nome, f"{info['grau_estenose']} por {info['comp_placa'].lower()} na artéria {info['nome']} ({loc_texto}) - Fluxo pós-lesão: {info['onda_pos'].lower()}{sufixo_concl}."))

            else:
                # Estenoses Consecutivas (em tandem)
                lesoes = info["lesoes_consecutivas"]
                if not lesoes:
                    txt_art += f"{prefixo_parede}pérvia, sem lesões estenosantes cadastradas na série."
                else:
                    partes_lesoes = [f"{l['segmento'].lower()} (fluxo distal {l['fluxo'].lower()})" for l in lesoes]
                    if len(partes_lesoes) == 1:
                        desc_lesoes = partes_lesoes[0]
                    else:
                        desc_lesoes = ", ".join(partes_lesoes[:-1]) + " e " + partes_lesoes[-1]
                    txt_art += f"{prefixo_parede}pérvia, identificando-se estenoses consecutivas (em tandem) no {desc_lesoes}."

                    pior_padrao = pior_padrao_fluxo([l["fluxo"] for l in lesoes])
                    sufixo_concl = " com repercussão em cascata nos vasos distais" if info["propagar_fluxo_distal"] else ""
                    segs_texto = formatar_segmentos([l["segmento"] for l in lesoes])
                    conclusoes_lista.append((m_nome, f"Estenoses consecutivas na artéria {info['nome']} ({segs_texto}) - Fluxo pós-lesão mais alterado: {pior_padrao.lower()}{sufixo_concl}."))

            add_p(txt_art, space_after=6)

    # IMPRESSÃO DIAGNÓSTICA (CONCLUSÃO)
    add_p("⸻", space_before=12, space_after=6)
    add_p("IMPRESSÃO DIAGNÓSTICA", space_after=6)

    # Processa se há relato de ateromatose geral no exame para enriquecer a conclusão
    tem_qualquer_ateroma = any(dados_m_dict[m][art_id]["tem_ateromatose"] for m in membros_lista for _, art_id in ARTERIAS_LISTA)

    if not conclusoes_lista:
        msg_normal = ("Sistema arterial fêmoro-poplíteo-tibial livre de espessamentos parietais significativos, "
                      "tortuosidades e dilatações. Padrão e velocidades de fluxo compatíveis com ausência de "
                      "obstruções, lesões estenosantes significativas e comunicações artério-venosas.")
        if tem_qualquer_ateroma:
            msg_normal += " Sinais de ateromatose difusa parietal sem repercussão luminal estrutural."
        add_p(msg_normal, bullet=True)
    else:
        if tem_qualquer_ateroma:
            add_p("Sinais macroscópicos de ateromatose arterial difusa de padrão parietal nos segmentos examinados.", bullet=True)
        for m_origem, conclusao_txt in conclusoes_lista:
            prefixo = f"[{m_origem}] " if formato_exame == "Bilateral (Laudo Único)" else ""
            add_p(f"{prefixo}{conclusao_txt}", bullet=True)

    if incluir_identidade and (nome_medico or crm_medico):
        doc.add_paragraph().paragraph_format.space_before = Pt(25)
        p_assin = doc.add_paragraph()
        p_assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_assin.paragraph_format.line_spacing = espacamento_linhas
        if nome_medico:
            r = p_assin.add_run(f"{nome_medico}\n")
            r.bold = True
        if crm_medico:
            p_assin.add_run(f"CRM-{crm_uf} {crm_medico}\n")
        if rqe_medico.strip():
            p_assin.add_run(f"RQE {rqe_medico}")

    return doc

# --- PROCESSAMENTO DO BOTAO DE EMISSAO ---
def _exibir_doc(doc, buf, nome_arquivo, label_download):
    st.success("Laudo Arterial estruturado e atualizado com sucesso!")
    if modo_saida in ["Somente Visualização", "Visualização + DOCX"]:
        texto_viz = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        st.markdown("## 👁️ Visualização do Laudo")
        st.text_area("Laudo Gerado", value=texto_viz, height=600)
    if modo_saida in ["Somente DOCX", "Visualização + DOCX"]:
        st.download_button(
            label_download,
            buf,
            nome_arquivo,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

if st.button("🚀 Gerar Laudo Arterial Completo", use_container_width=True):
    if formato_exame == "Bilateral (Laudos Separados)":
        for m_nome in ["DIREITO", "ESQUERDO"]:
            doc_sep = construir_laudo_arterial_word([m_nome], dados_membros)
            buf = BytesIO()
            doc_sep.save(buf)
            buf.seek(0)
            _exibir_doc(doc_sep, buf, f"Laudo_Doppler_Arterial_MMII_{m_nome}.docx",
                        f"📥 Baixar Laudo Membro {m_nome} (.docx)")
    else:
        doc_gerado = construir_laudo_arterial_word(membros_para_processar, dados_membros)
        buf = BytesIO()
        doc_gerado.save(buf)
        buf.seek(0)
        _exibir_doc(doc_gerado, buf, "Laudo_Doppler_Arterial_MMII.docx", "📥 Baixar Laudo Arterial (.docx)")
