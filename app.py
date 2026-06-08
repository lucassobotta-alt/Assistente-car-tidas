import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Inicialização segura do estado da sessão
if 'lista_placas' not in st.session_state:
    st.session_state.lista_placas = []
if 'lesoes_incipientes' not in st.session_state:
    st.session_state.lesoes_incipientes = []
if 'calcificacoes_isoladas' not in st.session_state:
    st.session_state.calcificacoes_isoladas = []
if 'lesoes_nao_ateromatosas' not in st.session_state:
    st.session_state.lesoes_nao_ateromatosas = []

# --- FUNÇÃO AUXILIAR PARA ESTIMAR PLAQUE-RADS ---
def retirar_prefixo_numerico(opcao_texto):
    if ". " in opcao_texto:
        return opcao_texto.split(". ", 1)[1]
    return opcao_texto

def estimar_plaque_rads(opcao_texto):
    if opcao_texto.startswith("1."):
        return "Plaque-RADS 1 (Placa puramente calcificada)"
    elif opcao_texto.startswith("2."):
        return "Plaque-RADS 2 (Placa fibrocalcificada ou fibrosa estável)"
    elif opcao_texto.startswith("3."):
        return "Plaque-RADS 3 (Placa predominantemente fibrosa)"
    elif opcao_texto.startswith("4."):
        return "Plaque-RADS 4 (Placa com núcleo lipídico/necrótico delimitado)"
    elif opcao_texto.startswith("5."):
        return "Plaque-RADS 5 (Placa predominantemente anecogênica/vulnerável)"
    elif opcao_texto.startswith("6."):
        return "Plaque-RADS 5 (Placa uniformemente anecogênica de alta vulnerabilidade)"
    return None

# --- CLASSIFICADORES HEMODINÂMICOS ADAPTATIVOS (SBC 2023 vs NASCET) ---
def obter_texto_hemo_continuo(estado, vps_aci, vcc, tem_placa=False, diretriz="Diretriz SBC 2023"):
    if estado == "Oclusão":
        return "Oclusão", "determinando oclusão completa do vaso, caracterizada por ausência total de fluxo ao estudo Doppler pulsado e mapeamento a cores."
    elif estado == "Suboclusão":
        return "Suboclusão", "determinando suboclusão do vaso, caracterizada por estreitamento luminal severo com padrão de fluxo filiforme ('trickle flow') ao estudo Doppler."
    
    relacao = round(vps_aci / max(vcc, 1), 2)
    
    if diretriz == "Diretriz SBC 2023":
        if vps_aci < 140:
            if tem_placa:
                return "Estenose < 50%", f"determinando estenose leve (<50% pelos critérios da Diretriz SBC 2023), caracterizada por velocidade de pico sistólico na artéria carótida interna de {vps_aci} cm/s."
            else:
                return "Normal", "pérvia, com fluxo bifásico anterógrado de baixa resistência, caracterizado por diástole sustentada e velocidades dentro da normalidade, compatível com irrigação de leito encefálico de baixa impedância. Não há sinais de estenose ou turbulência."
        
        if vps_aci > 400 or relacao > 5.0:
            return "Estenose > 90%", f"determinando estenose acentuada (>90% pelos critérios da Diretriz SBC 2023), caracterizada por acentuada elevação das velocidades de fluxo com VPS na artéria carótida interna de {vps_aci} cm/s e relação artéria carótida interna / artéria carótida comum de {relacao}."
        elif 230 < vps_aci <= 400 or relacao > 4.0:
            return "Estenose de 70-89%", f"determinando estenose hemodinamicamente significativa (70-89% pelos critérios da Diretriz SBC 2023), caracterizada por VPS na artéria carótida interna de {vps_aci} cm/s e relação artéria carótida interna / artéria carótida comum de {relacao}."
        elif 3.2 <= relacao <= 4.0:
            return "Estenose de 60-69%", f"determinando estenose moderada (60-69% pelos critérios da Diretriz SBC 2023), caracterizada por relação artéria carótida interna / artéria carótida comum de {relacao} e VPS na artéria carótida interna de {vps_aci} cm/s."
        else:
            return "Estenose de 50-59%", f"determinando estenose moderada (50-59% pelos critérios da Diretriz SBC 2023), caracterizada por VPS na artéria carótida interna de {vps_aci} cm/s e relação artéria carótida interna / artéria carótida comum de {relacao}."
            
    else:  # Critérios Clássicos do NASCET
        if vps_aci < 125:
            if tem_placa:
                return "Estenose < 50%", f"determinando estenose leve (<50% pelos critérios do Consenso NASCET), com VPS na artéria carótida interna de {vps_aci} cm/s."
            else:
                return "Normal", "pérvia, apresentando padrão de velocidades normais ao estudo Doppler, sem critérios para estenose hemodinâmica pelo Consenso NASCET."
        
        if vps_aci >= 230 or relacao >= 4.0:
            return "Estenose ≥ 70%", f"determinando estenose severa (≥70% pelos critérios do Consenso NASCET), caracterizada por VPS na artéria carótida interna de {vps_aci} cm/s e relação ACI/ACC de {relacao}."
        else:
            return "Estenose de 50-69%", f"determinando estenose moderada (50-69% pelos critérios do Consenso NASCET), caracterizada por VPS na artéria carótida interna de {vps_aci} cm/s e relação ACI/ACC de {relacao}."

def avaliar_vertebral(espectro, vps_vert):
    if espectro == "Normal (Fluxo Anterógrado)":
        if vps_vert >= 100:
            return "Estenose de Vertebral (>50%)", f"Artéria vertebral apresentando fluxo anterógrado com acentuada elevação focal de velocidades (VPS de {vps_vert} cm/s) e turbulência local, compatível com estenose segmentar superior a 50%."
        else:
            return "Normal", "Artéria vertebral pérvia, com fluxo anterógrado de baixa resistência e diástole sustentada, compatível com adequada perfusão vertebrobasilar."
    elif espectro == "Hipoplasia": 
        return "Hipoplasia de Vertebral", f"Artéria vertebral apresentando fluxo anterógrado de baixa resistência, porém exibindo calibre reduzido e velocidades proporcionalmente baixas (VPS de {vps_vert} cm/s), compatível com variante anatomofuncional (hipoplasia)."
    elif espectro == "Roubo Latente": 
        return "Sinal de Roubo Latente da Subclávia", "Artéria vertebral apresentando fluxo anterógrado, porém com morfologia de onda alterada devido a uma desaceleração mesosistólica abrupta, sugerindo alteração hemodinâmica inicial por estenose da artéria subclávia proximal ipsilateral."
    elif espectro == "Roubo Parcial (Fluxo Alternante)": 
        return "Sinal de Roubo Parcial da Subclávia", "Artéria vertebral apresentando padrão de fluxo alternante, caracterizado por vetor sistólico retrógrado e vetor diastólico anterógrado, indicando inversão parcial do fluxo por estenose acentuada da artéria subclávia proximal ipsilateral."
    elif espectro == "Roubo Total (Fluxo Retrógrado)": 
        return "Sinal de Roubo Total da Subclávia", "Artéria vertebral apresentando inversão completa e contínua do seu vetor de fluxo, confirmando o fenômeno de roubo de subclávia secundário a oclusão da artéria subclávia proximal ipsilateral."
    return "Alterada", f"Artéria vertebral com alterações inespecíficas do padrão de fluxo. VPS: {vps_vert} cm/s."

# --- CONFIGURAÇÃO DA INTERFACE STREAMLIT ---
st.set_page_config(page_title="Laudo Neurovascular Avançado", layout="wide")
st.title("⚕️ Assistente de Laudos Vascular Completo")

# ==========================================
#       PAINEL DE CONTROLE LATERAL
# ==========================================
with st.sidebar:
    st.markdown("## ⚙️ Painel de Controle Avançado")
    
    # 1. Parâmetros Científicos / Diretrizes
    st.markdown("### 📚 Critérios Científicos")
    diretriz_selecionada = st.radio(
        "Diretriz Hemodinâmica de Referência:",
        ["Diretriz SBC 2023", "Consenso Clássico NASCET"],
        help="Altera os pontos de corte de velocidade e a redação descritiva das estenoses."
    )
    ano_dislipidemia = st.selectbox("Ano da Diretriz de Dislipidemia:", ["2025", "2023"], index=0)
    
    st.markdown("---")
    
    # 2. Configurações de Formatação Visual do Documento
    st.markdown("### 📝 Formatação Externa (.docx)")
    fonte_doc = st.selectbox("Família da Fonte:", ["Arial", "Calibri", "Times New Roman"], index=0)
    tamanho_fonte = st.slider("Tamanho do Texto (pt):", 10, 14, 11, step=1)
    espacamento_linhas = st.slider("Espaçamento entre Linhas:", 1.0, 1.5, 1.15, step=0.05)
    quebrar_pagina_diag = st.toggle("Separar Impressão Diagnóstica em Nova Página", value=False)
    
    st.markdown("---")
    
    # 3. Identidade Visual e Assinatura Automatizada
    st.markdown("### ✍️ Identidade & Assinatura")
    nome_clinica = st.text_input("Cabeçalho / Nome da Clínica:", placeholder="Ex: Instituto de Diagnóstico por Imagem")
    nome_medico = st.text_input("Nome do Médico:", "Dr. Ekhator")
    crm_medico = st.text_input("CRM / RQE:")
    
    st.markdown("---")
    
    if st.button("🔄 Resetar Todos os Parâmetros", use_container_width=True, type="secondary"):
        st.session_state.lista_placas = []
        st.session_state.lesoes_incipientes = []
        st.session_state.calcificacoes_isoladas = []
        st.session_state.lesoes_nao_ateromatosas = []
        st.toast("🔄 Todos os dados clínicos foram limpos!")
        st.rerun()

# Opções técnicas fixas
opcoes_tecnicas = {
    "1. Exame sem limitações técnicas": "Exame realizado em decúbito dorsal, utilizando transdutor linear de alta frequência, com avaliação bidimensional, mapeamento de fluxo a cores e Doppler pulsado, sem limitações técnicas.",
    "2. Exame com limitação por condições anatômicas desfavoráveis": "Exame realizado em decúbito dorsal, utilizando transdutor linear de alta frequência, com avaliação bidimensional, mapeamento de fluxo a cores e Doppler pulsado. Devido a condições anatômicas desfavoráveis para insonação dos vasos cervicais, foi necessária avaliação complementar com transdutor convexo, o que pode reduzir a sensibilidade para identificação de placas ateroscleróticas de pequenas dimensões.",
    "3. Exame realizado à beira do leito (UTI)": "Exame realizado à beira do leito em unidade de terapia intensiva, utilizando transdutor linear de alta frequência, com limitações técnicas inerentes às condições do exame.",
    "4. Exame à beira do leito (UTI) com curativos cervicais": "Exame realizado à beira do leito em unidade de terapia intensiva, utilizando transdutor linear de alta frequência, com limitações técnicas inerentes às condições do exame e à presença de curativos cervicais sobre acessos jugulares."
}

col_id1, col_id2 = st.columns([2, 2])
with col_id1:
    nome = st.text_input("Nome do Paciente", "Paciente Exemplo")
with col_id2:
    opcao_selecionada = st.selectbox("Condições Técnicas do Exame:", list(opcoes_tecnicas.keys()))
    texto_tecnica_final = opcoes_tecnicas[opcao_selecionada]

st.markdown("---")
st.markdown("### 📊 Parâmetros Hemodinâmicos")
col_hemo_dir, col_hemo_esq = st.columns(2)

with col_hemo_dir:
    st.header("LADO DIREITO")
    cmi_dir = st.number_input("CMI Artéria Carótida Comum Direita (mm)", min_value=0.0, max_value=5.0, value=0.4, step=0.1)
    estado_aci_dir = st.selectbox("Estado da Artéria Carótida Interna Direita", ["Pérvia (Calcular por Velocidade)", "Suboclusão", "Oclusão"])
    vps_aci_dir = st.number_input("VPS Artéria Carótida Interna Direita (cm/s)", min_value=0.0, value=90.0, step=5.0)
    vcc_dir = st.number_input("VPS Artéria Carótida Comum Direita (cm/s)", min_value=1.0, value=60.0, step=5.0)
    ace_dir = st.selectbox("Artéria Carótida Externa Direita", ["Com padrão espectral de alta resistência, compatível com perfusão de leitos musculares extracranianos.", "Alterada / Estenose hemodinâmica"])
    espectro_vert_dir = st.selectbox("Espectro Artéria Vertebral Direita", ["Normal (Fluxo Anterógrado)", "Hipoplasia", "Roubo Latente", "Roubo Parcial (Fluxo Alternante)", "Roubo Total (Fluxo Retrógrado)"])
    vps_vert_dir = st.number_input("VPS Artéria Vertebral Direita (cm/s)", min_value=0.0, value=30.0, step=5.0)

with col_hemo_esq:
    st.header("LADO ESQUERDO")
    cmi_esq = st.number_input("CMI Artéria Carótida Comum Esquerda (mm)", min_value=0.0, max_value=5.0, value=0.4, step=0.1)
    estado_aci_esq = st.selectbox("Estado da Artéria Carótida Interna Esquerda", ["Pérvia (Calcular por Velocidade)", "Suboclusão", "Oclusão"])
    vps_aci_esq = st.number_input("VPS Artéria Carótida Interna Esquerda (cm/s)", min_value=0.0, value=100.0, step=5.0)
    vcc_esq = st.number_input("VPS Artéria Carótida Comum Esquerda (cm/s)", min_value=1.0, value=60.0, step=5.0)
    ace_esq = st.selectbox("Artéria Carótida Externa Esquerda", ["Com padrão espectral de alta resistência, compatível com perfusão de leitos musculares extracranianos.", "Alterada / Estenose hemodinâmica"])
    espectro_vert_esq = st.selectbox("Espectro Artéria Vertebral Esquerda", ["Normal (Fluxo Anterógrado)", "Hipoplasia", "Roubo Latente", "Roubo Parcial (Fluxo Alternante)", "Roubo Total (Fluxo Retrógrado)"])
    vps_vert_esq = st.number_input("VPS Artéria Vertebral Esquerda (cm/s)", min_value=0.0, value=30.0, step=5.0)

st.markdown("---")

with st.expander("🔄 2. Lesões Não Ateromatosas (Tortuosidades e Vasculite)"):
    col_na1, col_na2, col_na3 = st.columns(3)
    with col_na1:
        vaso_na = st.selectbox("Vaso com Lesão Não Ateromatosa:", ["Artéria carótida interna", "Artéria carótida comum", "Bulbo carotídeo", "Artéria vertebral"])
        lado_na = st.selectbox("Lado do Achado:", ["Direito", "Esquerdo", "Bilateral"])
    with col_na2:
        tipo_na = st.selectbox("Selecione a Lesão Não Ateromatosa:", [
            "Tortuosidade / Alongamento (Alongamento simples)",
            "Tortuosidade / Alongamento (Kinking - Angulação aguda < 90°)",
            "Tortuosidade / Alongamento (Coiling - Alça completa em espiral)",
            "Vasculite / Arterite (Espessamento parietal concêntrico e homogêneo - Sinal do Halo)",
            "Vasculite / Arterite (Espessamento difuso irregular de padrão inflamatório)"
        ])
    with col_na3:
        mensuracao_na = st.number_input("Mensuração / Espessura parietal (mm) se aplicável:", min_value=0.0, max_value=10.0, value=2.1, step=0.1)
        hemo_na = st.toggle("Gera alteração hemodinâmica ou aceleração focal?", value=False, key="hemo_na")

    if st.button("💾 Registrar Lesão Não Ateromatosa"):
        lados_add = ["Direito", "Esquerdo"] if lado_na == "Bilateral" else [lado_na]
        for ld in lados_add:
            item_na = {"vaso": vaso_na, "lado": ld, "tipo": tipo_na, "medida": mensuracao_na, "hemo": hemo_na}
            if item_na not in st.session_state.lesoes_nao_ateromatosas:
                st.session_state.lesoes_nao_ateromatosas.append(item_na)
        st.toast("✅ Lesão não ateromatosa registrada com sucesso!")

    if st.session_state.lesoes_nao_ateromatosas:
        for idx, na in enumerate(st.session_state.lesoes_nao_ateromatosas):
            st.write(f"• `{idx+1:02d}` **{na['vaso']} {na['lado']}**: {na['tipo']} ({na['medida']} mm) — Repercussão: {na['hemo']}")
        if st.button("❌ Limpar Lista Não Ateromatosa"):
            st.session_state.lesoes_nao_ateromatosas = []
            st.rerun()

st.markdown("---")

with st.expander("🌱 3. Lesões Estruturais Incipientes (Alterações Precoces ≤ 1.5 mm)"):
    tem_incipiente = st.checkbox("Incluir achados de Lesão Incipiente?")
    if tem_incipiente:
        col_inc1, col_inc2, col_inc3 = st.columns(3)
        with col_inc1:
            vaso_inc = st.selectbox("Vaso (Incipiente):", [
                "Artéria carótida comum direita", "Bulbo carotídeo direito", "Artéria carótida interna direita", 
                "Artéria carótida comum esquerda", "Bulbo carotídeo esquerdo", "Artéria carótida interna esquerda"
            ], key="vaso_inc")
        with col_inc2:
            local_inc = st.selectbox("Localização / Segmento (Incipiente):", ["terço proximal", "terço médio", "terço distal", "segmento total"], key="local_inc")
        with col_inc3:
            espessura_inc = st.number_input("Espessura Máxima (mm) [Max: 1.5mm]:", min_value=0.5, max_value=1.5, value=1.2, step=0.1)

        if st.button("💾 Registrar Lesão Incipiente"):
            nova_incipiente = {"vaso": vaso_inc, "localizacao": local_inc, "espessura": espessura_inc}
            if nova_incipiente not in st.session_state.lesoes_incipientes:
                st.session_state.lesoes_incipientes.append(nova_incipiente)
            st.toast("✅ Lesão estrutural incipiente registrada!")

        if st.session_state.lesoes_incipientes:
            for idx, inc in enumerate(st.session_state.lesoes_incipientes):
                st.write(f"• `{idx+1:02d}` **{inc['vaso']}** ({inc['localizacao']}) — Espessura precoce: {inc['espessura']} mm.")
            if st.button("❌ Limpar Lista Incipientes"):
                st.session_state.lesoes_incipientes = []
                st.rerun()

st.markdown("---")

with st.expander("🔎 4. Mapeamento de Placas Ateroscleróticas (Consolidadas ≥ 2.0 mm)"):
    tem_placa_ateromatosa = st.checkbox("Incluir achados de placas ateromatosas?")
    if tem_placa_ateromatosa:
        c_vaso, c_loc = st.columns(2)
        with c_vaso:
            vaso_selecionado = st.selectbox("Selecione a Artéria:", [
                "Artéria carótida comum direita", "Bulbo carotídeo direito", "Artéria carótida interna direita", 
                "Artéria carótida comum esquerda", "Bulbo carotídeo esquerdo", "Artéria carótida interna esquerda"
            ], key="vaso_placa")
        with c_loc:
            localizacao_selecionada = st.selectbox("Localização / Segmento:", ["terço proximal", "terço médio", "terço distal", "segmento total/bifurcação"], key="local_placa")
            
        culpada_hemo = st.toggle("Esta lesão determina estenose significativa?", value=False, key="culpada_placa")
            
        c1, c2, c3 = st.columns(3)
        with c1:
            composicao = st.selectbox("Composição / Tipo da Placa (Plaque-RADS):", [
                "1. Placa calcificada",
                "2. Placa uniformemente ecogênica (fibrosa estável)",
                "3. Placa predominantemente ecogênica (fibrosa)",
                "4. Placa com componente anecogênico delimitado por halo hiperecogênico",
                "5. Placa predominantemente anecogênica",
                "6. Placa uniformemente anecogênica sem cápsula fibrosa definida"
            ])
            usar_plaque_rads = st.toggle("Deseja estimar e incluir a classificação Plaque-RADS?", value=False)
            espessura = st.number_input("Espessura Máxima (mm):", min_value=2.0, max_value=20.0, value=2.5, step=0.1)
        with c2:
            superficie = st.selectbox("Superfície:", ["Regular", "Irregular", "Ulcerada"])
            profundidade_ulcura = st.text_input("Nicho Ulceroso:", placeholder="Ex: medindo 1.2 mm") if superficie == "Ulcerada" else ""
        with c3:
            st.markdown("**Achados Adicionais:**")
            tem_trombo = st.checkbox("Trombo adjacente")
            tem_fluxo = st.checkbox("Fluxo intraplaca")
            tem_calc_intra = st.checkbox("Focos de calcificação de permeio")

        if st.button("💾 Registrar Placa Ateromatosa"):
            achados_texto_lista = []
            if tem_trombo: achados_texto_lista.append("trombo adjacente")
            if tem_fluxo: achados_texto_lista.append("fluxo intraplaca")
            if tem_calc_intra: achados_texto_lista.append("calcificação intraplaca")
            
            pr_estimado = estimar_plaque_rads(composicao) if usar_plaque_rads else None
            composicao_limpa = retirar_prefixo_numerico(composicao)
                
            nova_placa = {
                "vaso": vaso_selecionado, "localizacao": localizacao_selecionada, "culpada_hemo": culpada_hemo,
                "composicao_texto": composicao_limpa, "superficie_texto": superficie, "profundidade_ulcura": profundidade_ulcura,
                "espessura": espessura, "achados_adicionais": achados_texto_lista,
                "plaque_rads": pr_estimado  
            }
            if nova_placa not in st.session_state.lista_placas:
                st.session_state.lista_placas.append(nova_placa)
            st.toast("✅ Placa registrada!")

        if st.session_state.lista_placas:
            for idx, p in enumerate(st.session_state.lista_placas):
                pr_tag = f" | {p['plaque_rads']}" if p['plaque_rads'] else ""
                st.write(f"`Item {idx+1:02d}` 🔸 **{p['vaso']}** ({p['localizacao']}) — {p['composicao_texto']} | {p['espessura']} mm de espessura ({p['superficie_texto'].lower()}){pr_tag}.")
            if st.button("❌ Limpar Lista de Placas"):
                st.session_state.lista_placas = []
                st.rerun()

st.markdown("---")

with st.expander("🪨 5. Calcificações Parietais Isoladas (Sem Formação de Placa)"):
    tem_calc_parietal = st.checkbox("Incluir registro de Calcificações Parietais Isoladas?")
    if tem_calc_parietal:
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            lado_calc = st.selectbox("Lado da Calcificação:", ["Direito", "Esquerdo", "Bilateral"])
        with col_c2:
            topografia_calc = st.selectbox("Localização Anatômica:", ["bulbo carotídeo", "artéria carótida comum", "artéria carótida interna"])
        
        if st.button("💾 Adicionar Calcificação Isolada"):
            lados_calc_add = ["Direito", "Esquerdo"] if lado_calc == "Bilateral" else [lado_calc]
            for ld in lados_calc_add:
                item_calc = {"lado": ld, "topografia": topografia_calc}
                if item_calc not in st.session_state.calcificacoes_isoladas:
                    st.session_state.calcificacoes_isoladas.append(item_calc)
            st.toast("✅ Calcificação registrada!")

        if st.session_state.calcificacoes_isoladas:
            for idx, c in enumerate(st.session_state.calcificacoes_isoladas):
                st.write(f"• `{idx+1:02d}` Calcificação no Lado **{c['lado']}** — {c['topografia'].title()}")
            if st.button("❌ Limpar Lista de Calcificações"):
                st.session_state.calcificacoes_isoladas = []
                st.rerun()

st.markdown("---")
gerar_laudo = st.button("🚀 Gerar Laudo Clínico Completo", use_container_width=True)

# --- MOTOR DE GERAÇÃO TEXTUAL ADAPTATIVO & FORMATADOR ---
if gerar_laudo:
    tem_placa_aci_dir = any("interna direita" in p['vaso'].lower() for p in st.session_state.lista_placas)
    tem_placa_aci_esq = any("interna esquerda" in p['vaso'].lower() for p in st.session_state.lista_placas)

    status_aci_dir_limpo, sufixo_hemo_aci_dir = obter_texto_hemo_continuo(estado_aci_dir, vps_aci_dir, vcc_dir, tem_placa_aci_dir, diretriz_selecionada)
    status_aci_esq_limpo, sufixo_hemo_aci_esq = obter_texto_hemo_continuo(estado_aci_esq, vps_aci_esq, vcc_esq, tem_placa_aci_esq, diretriz_selecionada)
    
    _, texto_vert_dir = avaliar_vertebral(espectro_vert_dir, vps_vert_dir)
    _, texto_vert_esq = avaliar_vertebral(espectro_vert_esq, vps_vert_esq)

    doc = Document()
    
    # Configuração Dinâmica da Fonte Padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = fonte_doc
    font.size = Pt(tamanho_fonte)
    
    # Funções Auxiliares de Injeção de Texto Adaptativo
    def adicionar_titulo(texto):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(texto.upper())
        run.bold = True
        run.font.name = fonte_doc
        run.font.size = Pt(tamanho_fonte + 1)

    def adicionar_subtitulo(texto):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(texto)
        run.bold = True
        run.font.name = fonte_doc
        run.font.size = Pt(tamanho_fonte)

    def adicionar_texto_esquerda(texto, bold_prefix=None, force_page_break=False):
        p = doc.add_paragraph()
        if force_page_break:
            p.insert_paragraph_before().add_run().add_break()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = espacamento_linhas
        
        if bold_prefix:
            run_pre = p.add_run(bold_prefix)
            run_pre.bold = True
            run_pre.font.name = fonte_doc
            run_pre.font.size = Pt(tamanho_fonte)
            
        run = p.add_run(texto)
        run.font.name = fonte_doc
        run.font.size = Pt(tamanho_fonte)

    # Injeção Opcional do Cabeçalho da Clínica
    if nome_clinica:
        p_clinica = doc.add_paragraph()
        p_clinica.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_clinica = p_clinica.add_run(nome_clinica.upper())
        r_clinica.bold = True
        r_clinica.font.size = Pt(tamanho_fonte + 2)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    adicionar_titulo('DUPLEX SCAN DAS ARTÉRIAS CARÓTIDAS E VERTEBRAIS')
    adicionar_texto_esquerda(f"Paciente: {nome}")
    adicionar_texto_esquerda(texto_tecnica_final, bold_prefix="Técnica: ")
    adicionar_subtitulo('RELATÓRIO')
    
    # --- LADO DIREITO ---
    adicionar_subtitulo('LADO DIREITO')
    
    txt_comum_dir = f"Artéria carótida comum direita pérvia, com diâmetro e trajeto conservados, apresentando fluxo bifásico anterógrado de baixa resistência. Espessura do complexo médio-intimal: {cmi_dir} mm."
    for inc in [x for x in st.session_state.lesoes_incipientes if "carótida comum direita" in x['vaso'].lower()]:
        txt_comum_dir += f" Identifica-se alteração estrutural incipiente precoce (espessamento focal igual ou inferior a 1.5 mm) no {inc['localizacao']}, medindo {inc['espessura']} mm de espessura."
    for p in [x for x in st.session_state.lista_placas if "carótida comum direita" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        txt_comum_dir += f" Identifica-se na parede uma placa de ateroma {p['composicao_texto'].lower()}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Direito" and "comum" in c['topografia']:
            txt_comum_dir += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."
    adicionar_texto_esquerda(txt_comum_dir)

    txt_bulbo_dir = "Bulbo carotídeo direito pérvio, com diâmetro e trajeto conservados."
    tem_achado_bulbo_dir = False
    for na in [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Direito" and "bulbo" in x['vaso'].lower()]:
        suf_h = "com alteração hemodinâmica local" if na['hemo'] else "sem repercussão hemodinâmica"
        txt_bulbo_dir += f" Identifica-se alteração não ateromatosa do tipo {na['tipo'].lower()} medindo {na['medida']} mm, {suf_h}."
        tem_achado_bulbo_dir = True
    for inc in [x for x in st.session_state.lesoes_incipientes if "bulbo carotídeo direito" in x['vaso'].lower()]:
        txt_bulbo_dir += f" Identifica-se lesão incipiente no {inc['localizacao']}, medindo {inc['espessura']} mm de espessura."
        tem_achado_bulbo_dir = True
    for p in [x for x in st.session_state.lista_placas if "bulbo carotídeo direito" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        txt_bulbo_dir += f" Apresentando na parede uma placa de ateroma {p['composicao_texto'].lower()}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
        tem_achado_bulbo_dir = True
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Direito" and "bulbo" in c['topografia']:
            txt_bulbo_dir += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."
            tem_achado_bulbo_dir = True
    if not tem_achado_bulbo_dir:
        txt_bulbo_dir += " Sem evidências de placas ou alterações estruturais."
    adicionar_texto_esquerda(txt_bulbo_dir)

    placas_aci_dir = [p for p in st.session_state.lista_placas if "interna direita" in p['vaso'].lower()]
    if placas_aci_dir:
        p = placas_aci_dir[0]
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        txt_aci_dir = f"Artéria carótida interna direita pérvia, apresentando na parede uma placa de ateroma {p['composicao_texto'].lower()}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, {sufixo_hemo_aci_dir}"
    else:
        txt_aci_dir = f"Artéria carótida interna direita pérvia, {sufixo_hemo_aci_dir}"
    adicionar_texto_esquerda(txt_aci_dir)

    adicionar_texto_esquerda(f"Artéria carótida externa direita {ace_dir.lower()}")
    adicionar_texto_esquerda(texto_vert_dir)

    # --- LADO ESQUERDO ---
    adicionar_subtitulo('LADO ESQUERDO')
    
    txt_comum_esq = f"Artéria carótida comum esquerda pérvia, com diâmetro e trajeto conservados, apresentando fluxo bifásico anterógrado de baixa resistência. Espessura do complexo médio-intimal: {cmi_esq} mm."
    for inc in [x for x in st.session_state.lesoes_incipientes if "carótida comum esquerda" in x['vaso'].lower()]:
        txt_comum_esq += f" Identifica-se alteração estrutural incipiente precoce (espessamento focal igual ou inferior a 1.5 mm) no {inc['localizacao']}, medindo {inc['espessura']} mm de espessura."
    for p in [x for x in st.session_state.lista_placas if "carótida comum esquerda" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        txt_comum_esq += f" Identifica-se na parede uma placa de ateroma {p['composicao_texto'].lower()}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Esquerdo" and "comum" in c['topografia']:
            txt_comum_esq += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."
    adicionar_texto_esquerda(txt_comum_esq)

    txt_bulbo_esq = "Bulbo carotídeo esquerdo pérvio, com diâmetro e trajeto conservados."
    tem_achado_bulbo_esq = False
    for na in [x for x in st.session_state.lesoes_nao_ateromatosas if x['lado'] == "Esquerdo" and "bulbo" in x['vaso'].lower()]:
        suf_h = "com alteração hemodinâmica local" if na['hemo'] else "sem repercussão hemodinâmica"
        txt_bulbo_esq += f" Identifica-se alteração não ateromatosa do tipo {na['tipo'].lower()} medindo {na['medida']} mm, {suf_h}."
        tem_achado_bulbo_esq = True
    for inc in [x for x in st.session_state.lesoes_incipientes if "bulbo carotídeo esquerdo" in x['vaso'].lower()]:
        txt_bulbo_esq += f" Identifica-se lesão incipiente no {inc['localizacao']}, medindo {inc['espessura']} mm de espessura."
        tem_achado_bulbo_esq = True
    for p in [x for x in st.session_state.lista_placas if "bulbo carotídeo esquerdo" in x['vaso'].lower()]:
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        txt_bulbo_esq += f" Apresentando na parede uma placa de ateroma {p['composicao_texto'].lower()}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}."
        tem_achado_bulbo_esq = True
    for c in st.session_state.calcificacoes_isoladas:
        if c['lado'] == "Esquerdo" and "bulbo" in c['topografia']:
            txt_bulbo_esq += " Identificam-se calcificações parietais isoladas sem repercussão hemodinâmica."
            tem_achado_bulbo_esq = True
    if not tem_achado_bulbo_esq:
        txt_bulbo_esq += " Sem evidências de placas ou alterações estruturais."
    adicionar_texto_esquerda(txt_bulbo_esq)

    placas_aci_esq = [p for p in st.session_state.lista_placas if "interna esquerda" in p['vaso'].lower()]
    if placas_aci_esq:
        p = placas_aci_esq[0]
        suffix_pr = f" ({p['plaque_rads']})" if p['plaque_rads'] else ""
        txt_aci_esq = f"Artéria carótida interna esquerda pérvia, apresentando na parede uma placa de ateroma {p['composicao_texto'].lower()}, medindo {p['espessura']} mm de espessura máxima, com superfície {p['superficie_texto'].lower()}{suffix_pr}, {sufixo_hemo_aci_esq}"
    else:
        txt_aci_esq = f"Artéria carótida interna esquerda pérvia, {sufixo_hemo_aci_esq}"
    adicionar_texto_esquerda(txt_aci_esq)

    adicionar_texto_esquerda(f"Artéria carótida externa esquerda {ace_esq.lower()}")
    adicionar_texto_esquerda(texto_vert_esq)

    # --- IMPRESSÃO DIAGNÓSTICA (Com quebra de página opcional) ---
    if quebrar_pagina_diag:
        doc.add_page_break()
        adicionar_subtitulo('IMPRESSÃO DIAGNÓSTICA')
    else:
        adicionar_subtitulo('IMPRESSÃO DIAGNÓSTICA')
        
    tem_achado = False
    cmi_alterado = (cmi_dir > 0.9 or cmi_esq > 0.9)
    tem_placa = len(st.session_state.lista_placas) > 0
    maior_que_plaque_rads_2 = any(p['espessura'] > 2.0 for p in st.session_state.lista_placas)

    if cmi_dir > 0.9 and cmi_esq > 0.9:
        adicionar_texto_esquerda("– Espessamento do complexo médio-intimal bilateralmente.")
        tem_achado = True
    elif cmi_dir > 0.9:
        adicionar_texto_esquerda("– Espessamento do complexo médio-intimal à direita.")
        tem_achado = True
    elif cmi_esq > 0.9:
        adicionar_texto_esquerda("– Espessamento do complexo médio-intimal à esquerda.")
        tem_achado = True

    for na in st.session_state.lesoes_nao_ateromatosas:
        tem_achado = True
        if "tortuosidade" in na['tipo'].lower():
            suf_msg = "hemodinamicamente significativa." if na['hemo'] else "sem repercussão hemodinâmica."
            adicionar_texto_esquerda(f"– Tortuosidade de trajeto na {na['vaso'].lower()} {na['lado'].lower()} {suf_msg}")
        elif "vasculite" in na['tipo'].lower() or "arterite" in na['tipo'].lower():
            adicionar_texto_esquerda(f"– Alterações sugestivas de processo inflamatório (vasculite) na {na['vaso'].lower()} {na['lado'].lower()}.")

    for inc in st.session_state.lesoes_incipientes:
        adicionar_texto_esquerda(f"– Alteração estrutural inicial precoce (lesão incipiente) na {inc['vaso'].lower()}.")
        tem_achado = True

    for p in st.session_state.lista_placas:
        tem_achado = True
        v_nome = p['vaso'].lower()
        if "interna" in v_nome:
            status_hemo, _ = obter_texto_hemo_continuo(
                estado_aci_dir if "direita" in v_nome else estado_aci_esq, 
                vps_aci_dir if "direita" in v_nome else vps_aci_esq, 
                vcc_dir if "direita" in v_nome else vcc_esq, 
                True, 
                diretriz_selecionada
            )
            vps_val = vps_aci_dir if "direita" in v_nome else vps_aci_esq
            vcc_val = vcc_dir if "direita" in v_nome else vcc_esq
            rel_val = round(vps_val / vcc_val, 2)
            justificativa_hemo = f", caracterizada por VPS de {vps_val} cm/s e relação ACI/ACC de {rel_val}" if "Normal" not in status_hemo else ""
        else:
            status_hemo = "Estenose < 50%" if not p['culpada_hemo'] else "Estenose de 50-59%"
            justificativa_hemo = ""

        suffix_diag_pr = f" [{p['plaque_rads']}]" if p['plaque_rads'] else ""

        if "Oclusão" in status_hemo:
            adicionar_texto_esquerda(f"– Placa de ateroma determinando oclusão completa de {v_nome}{suffix_diag_pr}.")
        elif "Suboclusão" in status_hemo:
            adicionar_texto_esquerda(f"– Placa de ateroma determinando suboclusão em {v_nome}{suffix_diag_pr}.")
        elif "Normal" in status_hemo or "< 50%" in status_hemo:
            adicionar_texto_esquerda(f"– Placa de ateroma discreta na {v_nome}{suffix_diag_pr}{justificativa_hemo}.")
        else:
            adicionar_texto_esquerda(f"– Placa de ateroma determinando estenose de {status_hemo.replace('Estenose de ', '')} na {v_nome}{justificativa_hemo}{suffix_diag_pr}.")

    has_bulbo_dir = any(c['lado'] == "Direito" and c['topografia'] == "bulbo carotídeo" for c in st.session_state.calcificacoes_isoladas)
    has_bulbo_esq = any(c['lado'] == "Esquerdo" and c['topografia'] == "bulbo carotídeo" for c in st.session_state.calcificacoes_isoladas)

    if has_bulbo_dir and has_bulbo_esq:
        adicionar_texto_esquerda("– Calcificações parietais isoladas sem repercussão nos bulbos carotídeos bilateralmente.")
        tem_achado = True
    else:
        for c in st.session_state.calcificacoes_isoladas:
            adicionar_texto_esquerda(f"– Calcificações parietais isoladas sem repercussão no {c['topografia']} {c['lado'].lower()}.")
            tem_achado = True

    if not tem_achado:
        adicionar_texto_esquerda("– Artérias carótidas e vertebrais pérvias, com trajetos e padrões de fluxo normais, dentro dos limites da normalidade.")

    # --- OBSERVAÇÕES DINÂMICAS COM PARÂMETROS ---
    obs_ativas = []
    if cmi_alterado:
        obs_ativas.append(
            "\"O espessamento do complexo médio-intimal carotídeo é considerado marcador de aterosclerose subclínica "
            "e associa-se a aumento do risco de eventos cardiovasculares, devendo sua interpretação ser integrada ao "
            "contexto clínico e aos demais fatores de risco do paciente.\" Referências: Mannheim Carotid Intima-Media "
            "Thickness Consensus (2004–2006); ESC/EAS Guidelines for the Management of Dyslipidaemias (2021)."
        )
    if tem_placa:
        obs_ativas.append(
            f"\"A presença de placa aterosclerótica carotídea, independentemente do grau de estenose, caracteriza "
            f"aterosclerose subclínica e constitui fator agravante de risco cardiovascular, devendo ser considerada "
            f"na estratificação global do risco cardiovascular, conforme a Diretriz Brasileira de Dislipidemias e "
            f"Prevenção da Aterosclerose – {ano_dislipidemia}.\""
        )
    if maior_que_plaque_rads_2 or any(p['plaque_rads'] is not None for p in st.session_state.lista_placas):
        obs_ativas.append(
            "\"A classificação Plaque-RADS padroniza a caracterização ultrassonográfica das placas carotídeas, "
            "incorporando aspectos morfológicos relacionados à vulnerabilidade da placa e fornecendo informação "
            "complementar ao grau de estenose na estratificação do risco de eventos cerebrovasculares.\" Referências: "
            "Plaque-RADS™ Consensus Statement (2023); recomendações da Society of Radiologists in Ultrasound para "
            "avaliação ultrassonográfica da doença carotídea."
        )

    total_obs = len(obs_ativas)
    if total_obs == 1:
        adicionar_texto_esquerda(obs_ativas[0], bold_prefix="– Observação: ")
    elif total_obs > 1:
        for i, texto_obs in enumerate(obs_ativas):
            adicionar_texto_esquerda(texto_obs, bold_prefix=f"– Observação {i+1}: ")

    # Bloco Dinâmico de Assinatura
    if nome_medico or crm_medico:
        doc.add_paragraph().paragraph_format.space_before = Pt(36)
        p_assinatura = doc.add_paragraph()
        p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if nome_medico:
            run_n = p_assinatura.add_run(f"{nome_medico}\n")
            run_n.bold = True
        if crm_medico:
            p_assinatura.add_run(crm_medico)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("Laudo integrado gerado com sucesso!")
    st.download_button(
        label="📥 Baixar Laudo Formatado (.docx)", 
        data=buffer, 
        file_name="Laudo_Vascular_Completo.docx", 
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
