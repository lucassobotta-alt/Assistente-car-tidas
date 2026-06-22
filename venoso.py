# venoso.py
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Laudo de Duplex Scan Venoso MMII", layout="wide")
st.title("🌀 Assistente de Laudos: Duplex Scan Venoso de MMII")

# --- SIDEBAR: CONFIGURAÇÕES E DADOS DO MÉDICO ---
with st.sidebar:
    st.markdown("## ⚙️ Painel de Controle Avançado")

    st.markdown("### 📚 Referências Científicas")
    st.markdown(
        "📄 <a href='https://www.ejves.com/article/S1078-5884(21)00979-5/fulltext' target='_blank'>"
        "ESVS 2022 Clinical Practice Guidelines – Chronic Venous Disease</a><br>"
        "📄 <a href='https://www.minervamedica.it/en/journals/international-angiology/article.php?cod=R34Y2023N01A0045' target='_blank'>"
        "The First Latin American Consensus on Superficial and Perforating Venous Mapping</a><br>"
        "📄 <a href='https://cbr.org.br/wp-content/uploads/2020/05/Consenso-para-a-Sociedade-Bras.-de-Angiologia-e-Cirurgia-Vascular_2020.pdf' target='_blank'>"
        "Consenso Brasileiro de Angiologia e Cirurgia Vascular – SBACV 2020</a>",
        unsafe_allow_html=True
    )

    st.markdown("---")
    st.markdown("### 📝 Formatação Externa (.docx)")
    fonte_doc = st.selectbox("Família da Fonte:", ["Arial", "Calibri", "Times New Roman"])
    tamanho_fonte = st.slider("Tamanho do Texto (pt):", 10, 14, 11)
    espacamento_linhas = st.slider("Espaçamento entre Linhas:", 1.0, 1.5, 1.15, step=0.05)
    quebrar_pagina_diag = st.toggle("Separar Impressão Diagnóstica em Nova Página", value=False)
    modo_saida = st.radio(
        "Modo de saída do laudo:",
        ["Somente DOCX", "Somente Visualização", "Visualização + DOCX"],
        index=2
    )

    st.markdown("---")
    st.markdown("### ✍️ Identidade & Assinatura")
    nome_clinica = st.text_input("Cabeçalho / Nome da Clínica:", placeholder="Ex: Instituto de Diagnóstico Vascular")
    nome_medico = st.text_input("Nome do Médico:", "Lucas Santos Guimarães")
    colcrm1, colcrm2 = st.columns([2, 1])
    with colcrm1:
        crm_medico = st.text_input("CRM:", "4061")
    with colcrm2:
        crm_uf = st.selectbox("UF", ["AC","AL","AP","AM","BA","CE","DF","ES","GO","MA","MT","MS","MG","PA","PB","PR","PE","PI","RJ","RN","RS","RO","RR","SC","SP","SE","TO"], index=25)
    rqe_medico = st.text_input("RQE:", "")

# --- IDENTIFICAÇÃO DO PACIENTE ---
nome_paciente = st.text_input("Nome do Paciente:", "Paciente Exemplo Venoso")
formato_exame = st.selectbox("Tipo de Exame:", ["Unilateral", "Bilateral (Laudos Separados)", "Bilateral (Laudo Único)"])

membros_para_processar = [st.selectbox("Selecione o Lado Avaliado:", ["DIREITO", "ESQUERDO"])] if formato_exame == "Unilateral" else ["DIREITO", "ESQUERDO"]

st.markdown("---")

# Inicialização do estado de sessão para estruturas dinâmicas
if "lista_perfurantes" not in st.session_state:
    st.session_state["lista_perfurantes"] = {}
if "segmentos_vsm_reg" not in st.session_state:
    st.session_state["segmentos_vsm_reg"] = {}
if "vsp_reg" not in st.session_state:
    st.session_state["vsp_reg"] = {}

dados_membros = {}

if len(membros_para_processar) > 1:
    abas = st.tabs(["🔴 Membro Inferior Direito (MID)", "🔵 Membro Inferior Esquerdo (MIE)"])
else:
    abas = [st.container()]

for idx, m_nome in enumerate(membros_para_processar):
    if m_nome not in st.session_state["lista_perfurantes"]:
        st.session_state["lista_perfurantes"][m_nome] = []
    if m_nome not in st.session_state["segmentos_vsm_reg"]:
        st.session_state["segmentos_vsm_reg"][m_nome] = []
    if m_nome not in st.session_state["vsp_reg"]:
        st.session_state["vsp_reg"][m_nome] = []

    with abas[idx]:
        st.markdown(f"### 📋 Parâmetros Clínicos - Membro {m_nome}")
        
        # 1. SISTEMA VENOSO PROFUNDO (SVP)
        st.markdown("#### 1. Sistema Venoso Profundo (SVP)")
        svp_status = st.radio(f"Status do SVP ({m_nome}):", ["Normal", "Anormal"], horizontal=True, key=f"svp_stat_{m_nome}")
        svp_res = {"status": svp_status}
        if svp_status == "Anormal":
            c_svp1, c_svp2 = st.columns(2)
            with c_svp1: svp_res["tipo"] = st.selectbox("Tipo de Alteração:", ["Refluxo", "Trombose Venosa Profunda (TVP)"], key=f"svp_tipo_{m_nome}")
            with c_svp2: svp_res["veias"] = st.multiselect("Veias Acometidas:", ["Veia Femoral Comum (VFC)", "Veia Femoral (VF)", "Veia Femoral Profunda (VFP)", "Veia Poplítea (V POP)", "Veias Gastrocnêmias", "Veias Soleares", "Veias Tibiais Posteriores (VTP)", "Veias Fibulares"], key=f"svp_veias_{m_nome}")
        
        st.markdown("---")
        
        # 2. SISTEMA VENOSO SUPERFICIAL (SVS) - VEIA SAFENA MAGNA (VSM)
        st.markdown("#### 2. Sistema Venoso Superficial (SVS)")
        st.markdown("##### 2.1 Veia Safena Magna (VSM)")
        
        vsm_status_geral = st.selectbox("Status Geral da VSM:", ["Pérvia / Presente", "Ausente (Safenectomia Total)", "Ausente (Safenectomia Segmentar)"], key=f"vsm_status_geral_{m_nome}")
        vsm_dados_mapeamento = {"status_geral": vsm_status_geral}
        
        if vsm_status_geral == "Pérvia / Presente":
            st.markdown("**Avaliação da Junção Safenofemoral (JSF):**")
            jsf_status = st.radio("Status da JSF:", ["Competente", "Incompetente"], horizontal=True, key=f"jsf_status_{m_nome}")
            vsm_dados_mapeamento["jsf_status"] = jsf_status
            
            jsf_detalhes_input = {}
            
            if jsf_status == "Incompetente":
                jsf_valvulas = st.selectbox(
                    "Padrão de Incompetência Valvar da JSF:", 
                    ["Válvulas Terminal e Pré-terminal incompetentes", "Apenas a Válvula Pré-terminal incompetente", "Apenas a Válvula Terminal incompetente"], 
                    key=f"jsf_valvulas_{m_nome}"
                )
                vsm_dados_mapeamento["jsf_valvulas"] = jsf_valvulas
                
                st.markdown("<div style='background-color: #f0f7ff; padding: 12px; border-left: 4px solid #0056b3; border-radius: 4px; margin-top: 10px; margin-bottom: 10px;'><strong>📐 Extensão e Deságue do Refluxo Juncional</strong></div>", unsafe_allow_html=True)
                
                if "Terminal e Pré-terminal" in jsf_valvulas:
                    c1, c2, c3 = st.columns(3)
                    with c1: jsf_detalhes_input["extensao_refluxo"] = st.selectbox("O refluxo estende-se até:", ["toda sua extensão", "o terço proximal da coxa", "o terço médio da coxa", "o terço distal da coxa", "a altura do joelho", "o terço proximal da perna", "o terço médio da perna", "o terço distal da perna"], key=f"jsf_ext_{m_nome}")
                    with c2: jsf_detalhes_input["desague_tipo"] = st.selectbox("Drenagem / Deságue para:", ["tributárias epifasciais varicosas", "veia perfurante incompetente", "veia comunicante", "veias colaterais de face medial"], key=f"jsf_des_tipo_{m_nome}")
                    with c3: jsf_detalhes_input["ponto_ref_dist"] = st.selectbox("Ponto de Referência do Deságue:", ["Interlinha do Joelho", "Junção Safenofemoral (JSF)", "Face Plantar"], key=f"jsf_ref_dist_{m_nome}")
                    
                    c4, c5 = st.columns(2)
                    with c4:
                        if jsf_detalhes_input["ponto_ref_dist"] == "Interlinha do Joelho":
                            jsf_detalhes_input["posicao_joelho_dist"] = st.radio("Posição em relação ao joelho:", ["acima", "abaixo"], horizontal=True, key=f"jsf_pos_joelho_{m_nome}")
                        else: jsf_detalhes_input["posicao_joelho_dist"] = ""
                    with c5: jsf_detalhes_input["dist_fim"] = st.text_input("Distância do ponto de referência (cm):", "0" if jsf_detalhes_input["ponto_ref_dist"]=="Interlinha do Joelho" else "15", key=f"jsf_dist_fim_{m_nome}")

                elif "Apenas a Válvula Pré-terminal" in jsf_valvulas:
                    c1, c2, c3 = st.columns(3)
                    with c1: jsf_detalhes_input["cm_ponto_j"] = st.text_input("Incompetente até cerca de (cm) da JSF/Ponto J:", "10", key=f"jsf_det_j_{m_nome}")
                    with c2: jsf_detalhes_input["tipo_drenagem"] = st.selectbox("Drenagem para:", ["tributárias epifasciais varicosas", "veia perfurante incompetente", "veia comunicante"], key=f"jsf_det_dren_{m_nome}")
                    with c3: jsf_detalhes_input["terco_coxa"] = st.selectbox("Em qual terço anatômico:", ["terço proximal da coxa", "terço médio da coxa", "terço distal da coxa"], key=f"jsf_det_terco_{m_nome}")
                
                elif "Apenas a Válvula Terminal" in jsf_valvulas:
                    jsf_detalhes_input["destino_terminal"] = st.selectbox(
                        "Destino do refluxo da Válvula Terminal:",
                        ["Fluxo para tributárias não safênicas (Mantém VSM competente)", "Escape para a Veia Safena Acessória Anterior (VSAA)"],
                        key=f"jsf_dest_term_{m_nome}"
                    )
                    if "Acessória Anterior" in jsf_detalhes_input["destino_terminal"]:
                        c1, c2 = st.columns(2)
                        with c1: jsf_detalhes_input["vsaa_desague"] = st.text_input("Deságue da VSAA em veias epifasciais da face:", "anterolateral da coxa", key=f"jsf_vsaa_des_{m_nome}")
                        with c2: jsf_detalhes_input["vsaa_extensao"] = st.text_input("Extensão dos trajetos varicosos até a face:", "lateral do joelho", key=f"jsf_vsaa_ext_{m_nome}")
                    else:
                        jsf_detalhes_input["tributarias_tipo"] = st.text_input("Fluindo para tributárias não safênicas na:", "coxa", key=f"jsf_det_term_trib_{m_nome}")

                vsm_dados_mapeamento["jsf_detalhes_input"] = jsf_detalhes_input

            incluir_segmento = st.toggle("Incluir Segmento Incompetente", key=f"incluir_seg_{m_nome}")

            seg_sel = []
            seg_origem = "Insuficiência valvar isolada"
            seg_desague = "Tributária epifascial varicosa"
            seg_prox_ref = "Junção Safenofemoral (JSF)"
            seg_prox_pos = ""
            seg_prox_cm = "0"
            seg_dist_ref = "Interlinha do Joelho"
            seg_dist_pos = "abaixo"
            seg_dist_cm = "15"

            if incluir_segmento:
                # Linha 1: Origem
                seg_origem = st.selectbox(
                    "Origem do refluxo (escape proximal):",
                    ["Insuficiência valvar isolada", "Tributárias pélvicas", "Varizes ganglionares",
                     "Tributária epifascial incompetente", "Veia de Giacomini incompetente",
                     "Veia perfurante incompetente"],
                    key=f"seg_origem_{m_nome}"
                )

                # Linha 2: Início — Referência e Distância
                _refs = ["Junção Safenofemoral (JSF)", "Interlinha do Joelho", "Face Plantar"]
                c_prox1, c_prox2, c_prox3 = st.columns(3)
                with c_prox1:
                    seg_prox_ref = st.selectbox("Início — Referência:", _refs, key=f"seg_prox_ref_{m_nome}")
                with c_prox2:
                    if seg_prox_ref == "Interlinha do Joelho":
                        seg_prox_pos = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"seg_prox_pos_{m_nome}")
                    else:
                        seg_prox_pos = ""
                        st.empty()
                with c_prox3:
                    seg_prox_cm = st.text_input("Distância (cm):", "0", key=f"seg_prox_cm_{m_nome}")

                # Linha 3: Ponto distal
                seg_desague = st.selectbox(
                    "Ponto distal (deságue do refluxo):",
                    ["Tributária epifascial varicosa", "Tributária varicosa troncular",
                     "Veia perfurante de drenagem", "Região maleolar"],
                    key=f"seg_desague_{m_nome}"
                )

                # Linha 4: Fim — Referência e Distância (oculto se região maleolar)
                if seg_desague != "Região maleolar":
                    c_dist1, c_dist2, c_dist3 = st.columns(3)
                    with c_dist1:
                        seg_dist_ref = st.selectbox("Fim — Referência:", _refs, key=f"seg_dist_ref_{m_nome}")
                    with c_dist2:
                        if seg_dist_ref == "Interlinha do Joelho":
                            seg_dist_pos = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"seg_dist_pos_{m_nome}")
                        else:
                            seg_dist_pos = ""
                            st.empty()
                    with c_dist3:
                        seg_dist_cm = st.text_input("Distância (cm):", "15", key=f"seg_dist_cm_{m_nome}")
                else:
                    seg_dist_ref = ""
                    seg_dist_pos = ""
                    seg_dist_cm = ""

                if st.button("💾 Registrar Segmento", key=f"reg_seg_vsm_{m_nome}"):
                    st.session_state["segmentos_vsm_reg"][m_nome].append({
                        "segmentos": seg_sel,
                        "origem": seg_origem, "desague": seg_desague,
                        "prox_ref": seg_prox_ref, "prox_pos": seg_prox_pos, "prox_cm": seg_prox_cm,
                        "dist_ref": seg_dist_ref, "dist_pos": seg_dist_pos, "dist_cm": seg_dist_cm,
                    })
                    st.rerun()

            if st.session_state["segmentos_vsm_reg"][m_nome]:
                for i, reg in enumerate(st.session_state["segmentos_vsm_reg"][m_nome]):
                    def _fmt_ref(ref, pos, cm):
                        if ref == "Interlinha do Joelho":
                            return f"{cm} cm {pos} da interlinha do joelho"
                        return f"{cm} cm da {ref}"
                    segs_label = ", ".join(reg["segmentos"]) if reg["segmentos"] else "insuficiência valvar"
                    desc = (f"`{i+1:02d}` **{segs_label}** | origem: {reg.get('origem','')} | deságue: {reg.get('desague','')} — "
                            f"de {_fmt_ref(reg['prox_ref'], reg['prox_pos'], reg['prox_cm'])} "
                            f"até {_fmt_ref(reg['dist_ref'], reg['dist_pos'], reg['dist_cm'])}")
                    st.markdown(desc)
                if st.button("❌ Limpar Segmentos Registrados", key=f"clear_seg_vsm_{m_nome}"):
                    st.session_state["segmentos_vsm_reg"][m_nome] = []
                    st.rerun()

            vsm_dados_mapeamento["tronco_refluxo"] = bool(st.session_state["segmentos_vsm_reg"][m_nome])
            vsm_dados_mapeamento["segmentos_lista"] = st.session_state["segmentos_vsm_reg"][m_nome]

        # Mensurações da Veia Safena Magna
        if "Pérvia" in vsm_status_geral:
            st.markdown("**Mensurações da Veia Safena Magna (Diâmetros em mm):**")
            cm1, cm2, cm3, cm4 = st.columns(4)
            with cm1: jsf_mm = st.text_input("Junção safenofemoral (mm):", "4.5", key=f"jsf_mm_{m_nome}")
            with cm2: vsm_prox_coxa = st.text_input("Terço proximal da coxa (mm):", "3.8", key=f"prox_c_{m_nome}")
            with cm3: vsm_med_coxa = st.text_input("Terço médio da coxa (mm):", "3.5", key=f"med_c_{m_nome}")
            with cm4: vsm_dist_coxa = st.text_input("Terço distal da coxa (mm):", "3.2", key=f"dist_c_{m_nome}")
            cm5, cm6, cm7 = st.columns(3)
            with cm5: vsm_prox_perna = st.text_input("Terço proximal da perna (mm):", "3.0", key=f"prox_p_{m_nome}")
            with cm6: vsm_med_perna = st.text_input("Terço médio da perna (mm):", "2.8", key=f"med_p_{m_nome}")
            with cm7: vsm_dist_perna = st.text_input("Terço distal da perna (mm):", "2.5", key=f"dist_p_{m_nome}")
            # Alertas de diâmetro da VSM (ESVS 2022)
            _alertas_vsm = []
            try:
                if float(jsf_mm) > 8.0:
                    _alertas_vsm.append(f"JSF: {jsf_mm} mm > 8,0 mm — dilatação acentuada da junção safenofemoral")
            except ValueError: pass
            for _lbl, _val in [("Terço proximal da coxa", vsm_prox_coxa), ("Terço médio da coxa", vsm_med_coxa),
                                ("Terço distal da coxa", vsm_dist_coxa), ("Terço proximal da perna", vsm_prox_perna),
                                ("Terço médio da perna", vsm_med_perna), ("Terço distal da perna", vsm_dist_perna)]:
                try:
                    if float(_val) > 6.0:
                        _alertas_vsm.append(f"{_lbl}: {_val} mm > 6,0 mm — tronco dilatado (limiar ESVS 2022 para indicação terapêutica)")
                except ValueError: pass
            if _alertas_vsm:
                st.warning("⚠️ **Diâmetros fora do limiar de referência (ESVS 2022):**\n\n" + "\n\n".join(f"• {a}" for a in _alertas_vsm))
        else: jsf_mm = vsm_prox_coxa = vsm_med_coxa = vsm_dist_coxa = vsm_prox_perna = vsm_med_perna = vsm_dist_perna = ""

        st.markdown("---")

        # 2.2 VEIA SAFENA PARVA (VSP)
        st.markdown("##### 2.2 Veia Safena Parva (VSP)")
        vsp_template = st.selectbox(
            "Padrão / Alterações na Veia Safena Parva (VSP):",
            [
                "Normal (Pérvia, de trajeto anatômico habitual e competente)",
                "Ausente (Safenectomia Total)",
                "Ausente (Safenectomia Segmentar)",
                "Junção safenopoplítea ausente, com extensão cranial da veia safena parva (Veia de Giacomini)",
                "JSP Incompetente -> refluxo na VSP em toda extensão com dilatação/tortuosidade"
            ],
            key=f"vsp_temp_choise_{m_nome}"
        )
        vsp_dados_form = {"template": vsp_template}

        if "JSP Incompetente" in vsp_template:
            st.markdown("<sub style='color: #444;'>Detalhamento do Deságue da VSP:</sub>", unsafe_allow_html=True)
            vsp_dados_form["desague_tipo"] = st.selectbox("Drenagem do refluxo da VSP para:", ["tributárias epifasciais varicosas", "veia perfurante incompetente", "veia comunicante", "malha reticular maleolar"], key=f"vsp_des_tipo_{m_nome}")
            _refs_vsp = ["Interlinha do Joelho", "Face Plantar"]
            cvsp_1, cvsp_2, cvsp_3 = st.columns(3)
            with cvsp_1:
                vsp_dados_form["desague_ref"] = st.selectbox("Fim do refluxo — Referência:", _refs_vsp, key=f"vsp_des_ref_{m_nome}")
            with cvsp_2:
                if vsp_dados_form["desague_ref"] == "Interlinha do Joelho":
                    vsp_dados_form["desague_pos"] = st.radio("Posição:", ["acima", "abaixo"], horizontal=True, key=f"vsp_des_pos_{m_nome}")
                else:
                    vsp_dados_form["desague_pos"] = ""
                    st.empty()
            with cvsp_3:
                vsp_dados_form["desague_cm"] = st.text_input("Distância (cm):", "10", key=f"vsp_des_cm_{m_nome}")

        if "Ausente" not in vsp_template:
            st.markdown("**Mensurações da Veia Safena Parva (Diâmetros em mm):**")
            cp1, cp2, cp3 = st.columns(3)
            label_jsp_dinamico = "Extensão cranial (mm):" if "extensão cranial" in vsp_template else "Junção safenopoplítea (mm):"
            with cp1: vsp_dados_form["jsp_mm"] = st.text_input(label_jsp_dinamico, "4.2", key=f"jsp_mm_{m_nome}")
            with cp2: vsp_dados_form["vsp_crossa"] = st.text_input("Crossa da safena parva (mm):", "3.8", key=f"crossa_{m_nome}")
            with cp3: vsp_dados_form["vsp_med_perna_diam"] = st.text_input("Terço médio da perna (mm):", "3.0", key=f"med_per_{m_nome}")

        if st.button("💾 Registrar Achado da VSP", key=f"reg_vsp_{m_nome}"):
            st.session_state["vsp_reg"][m_nome] = vsp_dados_form
            st.rerun()

        vsp_dados_input = st.session_state["vsp_reg"][m_nome] if st.session_state["vsp_reg"][m_nome] else vsp_dados_form

        if st.session_state["vsp_reg"][m_nome]:
            reg_vsp = st.session_state["vsp_reg"][m_nome]
            lbl_vsp = reg_vsp["template"].split("(")[0].strip() if "(" in reg_vsp["template"] else reg_vsp["template"][:60]
            st.write(f"✅ **VSP registrada:** {lbl_vsp}")
            if st.button("❌ Limpar Registro VSP", key=f"clear_vsp_{m_nome}"):
                st.session_state["vsp_reg"][m_nome] = []
                st.rerun()

        jsp_mm = vsp_dados_input.get("jsp_mm", "")
        vsp_crossa = vsp_dados_input.get("vsp_crossa", "")
        vsp_med_perna_diam = vsp_dados_input.get("vsp_med_perna_diam", "")

        st.markdown("---")

        # --- 2.3 VEIAS PERFURANTES INCOMPETENTES ISOLADAS ---
        st.markdown("#### 2.3 Veias Perfurantes Incompetentes (Mapeamento Separado)")
        possui_perfurantes = st.checkbox("Identifica veias perfurantes incompetentes neste membro?", value=False, key=f"has_perf_{m_nome}")
        perfurantes_coletadas = []
        
        if possui_perfurantes:
            st.markdown("<div style='background-color: #fff9e6; padding: 10px; border-left: 4px solid #ffcc00; border-radius: 4px; margin-bottom: 15px;'><strong>📍 Localização Dinâmica de Perfurantes Incompetentes</strong></div>", unsafe_allow_html=True)
            qtd_perf = len(st.session_state["lista_perfurantes"][m_nome])
            
            for p_idx in range(qtd_perf):
                st.markdown(f"**Veia Perfurante Incompetente #{p_idx + 1}**")
                perf_dados = {}
                cp_1, cp_2, cp_3, cp_4, cp_5 = st.columns(5)
                with cp_1: perf_dados["regiao"] = st.selectbox("Região Anatômica:", ["Coxa", "Perna"], key=f"perf_reg_{m_nome}_{p_idx}")
                with cp_2: perf_dados["face"] = st.selectbox("Face Medida:", ["Medial", "Lateral", "Anterior", "Posterior", "Anterolateral", "Posterointerna"], key=f"perf_face_{m_nome}_{p_idx}")
                with cp_3: perf_dados["ref_ponto"] = st.selectbox("Referência de Medida:", ["Interlinha do Joelho", "Face Plantar"], key=f"perf_ref_{m_nome}_{p_idx}")
                with cp_4: perf_dados["altura_cm"] = st.text_input("Altura aferida (cm):", "12", key=f"perf_alt_{m_nome}_{p_idx}")
                with cp_5: perf_dados["diametro_mm"] = st.text_input("Diâmetro (mm):", "3.5", key=f"perf_diam_{m_nome}_{p_idx}")
                try:
                    if float(perf_dados["diametro_mm"]) > 3.5:
                        st.warning(f"⚠️ Perfurante #{p_idx + 1}: diâmetro {perf_dados['diametro_mm']} mm > 3,5 mm — critério ESVS 2022 para perfurante patológica")
                except ValueError: pass
                
                if perf_dados["ref_ponto"] == "Interlinha do Joelho":
                    perf_dados["posicao_joelho"] = st.radio("Posição do plano do joelho:", ["Acima", "Abaixo"], horizontal=True, key=f"perf_pos_j_{m_nome}_{p_idx}")
                else: perf_dados["posicao_joelho"] = ""
                
                perfurantes_coletadas.append(perf_dados)
                st.markdown("---")
                
            c_pbtn1, c_pbtn2, _ = st.columns([1.5, 1.5, 3])
            with c_pbtn1:
                if st.button("➕ Adicionar Nova Perfurante", key=f"add_perf_btn_{m_nome}"):
                    st.session_state["lista_perfurantes"][m_nome].append(1)
                    st.rerun()
            with c_pbtn2:
                if len(st.session_state["lista_perfurantes"][m_nome]) > 0:
                    if st.button("❌ Remover Última Perfurante", key=f"rem_perf_btn_{m_nome}"):
                        st.session_state["lista_perfurantes"][m_nome].pop()
                        st.rerun()
                        
        # --- 2.4 MAPA DE VARICOSIDADES ---
        st.markdown("---")
        st.markdown("#### 2.4 Mapeamento de Varicosidades / Malhas Reticulares")
        possui_varicosidades = st.checkbox("Descrever presença de Telangiectasias, Microvarizes ou Reticulares?", key=f"has_varic_{m_nome}")
        varic_dados = {"possui": possui_varicosidades}
        
        if possui_varicosidades:
            cv_1, cv_2, cv_3 = st.columns(3)
            with cv_1: varic_dados["telangiectasias"] = st.checkbox("Telangiectasias (< 1 mm)", key=f"var_tel_{m_nome}")
            with cv_2: varic_dados["micro_reticulares"] = st.checkbox("Microvarizes / Varizes Reticulares (1 a 3 mm)", key=f"var_mic_{m_nome}")
            with cv_3: varic_dados["veias_varicosas"] = st.checkbox("Veias Varicosas Tronculares (> 3 mm)", key=f"var_tronc_{m_nome}")
            varic_dados["localizacao"] = st.text_input("Localização predominante das lesões superficiais:", "em faces lateral da coxa e posterior da perna", key=f"var_loc_{m_nome}")

        st.markdown("---")

        # 3. MÓDULOS ADICIONAIS & VARIÁVEIS EXTRAS
        st.markdown("#### 3. Módulos Adicionais")
        c_add1, c_add2 = st.columns(2)
        with c_add1:
            giacomini_opt = st.selectbox("Veia de Giacomini Isolada:", ["Não se aplica / Normal", "3.1 Refluxo ostial drenado de forma ascendente", "3.2 Refluxo ostial transferindo para VSM"], key=f"giacomini_{m_nome}")
            varizes_pelvicas_opt = st.selectbox("Varizes Pélvicas (Pontos de Escape):", ["Ausentes", "5.1 Plexo venoso ciático (região infraglútea)", "5.2 Escape inguinal/perineal"], key=f"pelvicas_{m_nome}")
            pos_op_opt = st.selectbox("Pós-Operatório / Recidiva:", ["Não se aplica", "6.3 Sinais de neovascularização adjacentes"], key=f"pos_op_{m_nome}")
        
        with c_add2:
            st.markdown("**Achados Extras / Patologias de Tecidos Adjacentes:**")
            achados_multiplos = st.multiselect(
                "Selecione os achados adicionais observados:",
                ["Cisto de Baker na fossa poplítea", "Edema intersticial subcutâneo"],
                default=[],
                key=f"achados_adi_multi_{m_nome}"
            )
            
            cisto_medidas = {}
            if "Cisto de Baker na fossa poplítea" in achados_multiplos:
                st.markdown("<sub style='color: #444;'>Dimensões do Cisto de Baker:</sub>", unsafe_allow_html=True)
                cc1, cc2, cc3 = st.columns(3)
                with cc1: cisto_medidas["comp"] = st.text_input("Eixo Long (mm):", "35", key=f"cb_c_{m_nome}")
                with cc2: cisto_medidas["larg"] = st.text_input("Eixo Transv (mm):", "18", key=f"cb_l_{m_nome}")
                with cc3: cisto_medidas["esp"] = st.text_input("Espessura (mm):", "12", key=f"cb_e_{m_nome}")

        dados_membros[m_nome] = {
            "svp": svp_res, "vsm_mapeamento": vsm_dados_mapeamento,
            "jsf_mm": jsf_mm, "vsm_prox_coxa": vsm_prox_coxa, "vsm_med_coxa": vsm_med_coxa, "vsm_dist_coxa": vsm_dist_coxa,
            "vsm_prox_perna": vsm_prox_perna, "vsm_med_perna": vsm_med_perna, "vsm_dist_perna": vsm_dist_perna,
            "vsp_dados_input": vsp_dados_input, "jsp_mm": jsp_mm, "vsp_crossa": vsp_crossa, "vsp_med_perna_diam": vsp_med_perna_diam,
            "giacomini_opt": giacomini_opt, "varizes_pelvicas_opt": varizes_pelvicas_opt, 
            "pos_op_opt": pos_op_opt, "achados_adi_multi": achados_multiplos, "cisto_medidas": cisto_medidas,
            "perfurantes_lista": perfurantes_coletadas if possui_perfurantes else [],
            "varic_dados": varic_dados
        }

st.markdown("---")

# --- FUNÇÃO DE MONTAGEM E CONSTRUÇÃO DO DOCUMENTO WORD ---
def construir_laudo_word(membros_lista, dados_m_dict):
    doc = Document()
    doc.styles['Normal'].font.name = fonte_doc
    doc.styles['Normal'].font.size = Pt(tamanho_fonte)
    
    def add_p(text, bold_pre=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4, bullet=False):
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.alignment = align
        p.paragraph_format.line_spacing = espacamento_linhas
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if bold_pre:
            r_p = p.add_run(bold_pre)
            r_p.bold = True
        p.add_run(text)

    if nome_clinica.strip():
        p_cl = doc.add_paragraph()
        p_cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cl = p_cl.add_run(nome_clinica.upper())
        r_cl.bold = True
        r_cl.font.name = fonte_doc
        r_cl.font.size = Pt(tamanho_fonte + 2)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    if formato_exame == "Bilateral (Laudo Único)":
        add_p("DOS MEMBROS INFERIORES", bold_pre='DUPLEX SCAN VENOSO ', space_after=12)
    else:
        add_p(f"DO MEMBRO INFERIOR {membros_lista[0]}", bold_pre='DUPLEX SCAN VENOSO ', space_after=12)

    add_p(f" {nome_paciente}", bold_pre="Paciente:")
    add_p("TÉCNICA", space_before=12, space_after=6)
    add_p("Exame realizado com transdutor linear de alta frequência...", space_after=12)
    
    conclusoes_lista = []
    
    for m_nome in membros_lista:
        dm = dados_m_dict[m_nome]
        add_p("⸻", space_after=12)
        
        if formato_exame == "Bilateral (Laudo Único)":
            add_p(f"RELATÓRIO TÉCNICO – MEMBRO INFERIOR {m_nome}", space_after=12)
        else: add_p("RELATÓRIO", space_after=12)
            
        # 1. SVP
        add_p("SISTEMA VENOSO PROFUNDO", space_after=6)
        if dm["svp"]["status"] == "Normal":
            add_p("As veias femoral comum, femoral, poplítea, tibiais posteriores e fibulares apresentam-se pérvias, compressíveis, com fluxo fásico com a respiração e competentes...")
        else:
            add_p(f"Sistema Venoso Profundo ANORMAL. Detectados sinais de {dm['svp']['tipo']} nas veias: {', '.join(dm['svp'].get('veias', []))}.")

        # 2. SVS - VEIA SAFENA MAGNA (VSM)
        add_p("SISTEMA VENOSO SUPERFICIAL", space_before=12, space_after=6)
        vm = dm["vsm_mapeamento"]
        
        if "Ausente" in vm["status_geral"]:
            add_p(f"Veia safena magna ausente ({vm['status_geral']}).")
        else:
            if vm["jsf_status"] == "Competente":
                add_p("A junção safenofemoral apresenta-se competente, sem evidências de refluxo valvar patológico.")
            else:
                v_padrao = vm["jsf_valvulas"]
                det = vm.get("jsf_detalhes_input", {})
                
                if "Terminal e Pré-terminal" in v_padrao:
                    ref_dist = det.get("ponto_ref_dist", "Interlinha do Joelho")
                    dist_d = det.get("dist_fim", "0")
                    pos_joelho_dist = det.get("posicao_joelho_dist", "")
                    termo_dist = f"a {dist_d} cm {pos_joelho_dist.lower()} da interlinha do joelho" if "Interlinha" in ref_dist and dist_d != "0" else "na altura da interlinha do joelho" if "Interlinha" in ref_dist else f"a {dist_d} cm de distância da {ref_dist}"
                    extensao_txt = det.get("extensao_refluxo", "toda sua extensão")
                    
                    add_p(f"Refluxo originado de incompetência das válvulas pré-terminal e terminal da junção safenofemoral, com escape para a veia safena magna, que segue incompetente até {extensao_txt}. O refluxo apresenta deságue para {det.get('desague_tipo','tributárias epifasciais varicosas')} localizado {termo_dist}.")
                    conclusoes_lista.append((m_nome, f"Insuficiência segmentar do tronco da veia safena magna por incompetência da junção safenofemoral (até {extensao_txt})."))
                
                elif "Apenas a Válvula Pré-terminal" in v_padrao:
                    add_p(f"Refluxo originado de incompetência da válvula pré-terminal da junção safenofemoral, com escape de refluxo proveniente de tributária para o segmento proximal da veia safena magna que segue incompetente até cerca de {det.get('cm_ponto_j','__')} cm da junção/ponto J, onde ocorre drenagem para {det.get('tipo_drenagem','tributárias epifasciais varicosas')} formando complexos varicosos na face medial da coxa em {det.get('terco_coxa','terço médio')}.")
                    conclusoes_lista.append((m_nome, "Insuficiência proximal da veia safena magna por incompetência da junção safenofemoral, associada a varizes calibrosas mediais na coxa."))
                
                elif "Apenas a Válvula Terminal" in v_padrao:
                    dest_term = det.get("destino_terminal", "")
                    if "Acessória Anterior" in dest_term:
                        add_p(f"Refluxo originado de incompetência da válvula terminal da junção safenofemoral com escape para o segmento proximal da veia safena acessória anterior e deságue em veias epifasciais varicosas na face {det.get('vsaa_desague','anterolateral da coxa')}. Há extensão dos trajetos varicosos até a face {det.get('vsaa_extensao','lateral do joelho')}. A veia safena magna segue competente.")
                        conclusoes_lista.append((m_nome, "Varizes anterolaterais na coxa originárias da veia safena acessória anterior por insuficiência da válvula terminal da JSF."))
                    else:
                        add_p(f"Refluxo originado de incompetência da válvula terminal da junção safenofemoral fluindo para tributárias não safênicas na {det.get('tributarias_tipo','coxa')}. A veia safena magna segue pérvia e competente com calibre preservado.")
                        conclusoes_lista.append((m_nome, "Varizes proximais na coxa, originárias da junção safenofemoral incompetente por falha da válvula terminal."))

            if vm["tronco_refluxo"] and vm.get("segmentos_lista"):
                def _fmt_ref_doc(ref, pos, cm):
                    if ref == "Interlinha do Joelho":
                        return f"{cm} cm {pos} da interlinha do joelho"
                    return f"{cm} cm da {ref}"
                for reg in vm["segmentos_lista"]:
                    origem = reg.get("origem", "")
                    desague = reg.get("desague", "")
                    origem_txt = f" com escape originado de {origem.lower()}" if origem and origem != "Insuficiência valvar isolada" else ""
                    if desague == "Região maleolar":
                        desague_txt = ", estendendo-se até a região maleolar"
                    elif desague:
                        desague_txt = f", com deságue para {desague.lower()}"
                    else:
                        desague_txt = ""
                    inicio = _fmt_ref_doc(reg["prox_ref"], reg["prox_pos"], reg["prox_cm"])
                    tem_fim = reg.get("dist_ref") and desague != "Região maleolar"
                    fim = _fmt_ref_doc(reg["dist_ref"], reg["dist_pos"], reg["dist_cm"]) if tem_fim else None
                    extensao_txt = f"com extensão de {inicio} até {fim}" if fim else f"a partir de {inicio}"
                    if reg["segmentos"]:
                        segs_txt = ", ".join(reg["segmentos"])
                        add_p(f"Identificado(s) segmento(s) incompetente(s) no tronco da veia safena magna ({segs_txt}){origem_txt}, {extensao_txt}{desague_txt}.")
                        for terco_m in reg["segmentos"]:
                            conclusoes_lista.append((m_nome, f"Insuficiência segmentar do tronco da VSM ({terco_m})."))
                    else:
                        add_p(f"Insuficiência valvar do tronco da veia safena magna{origem_txt}, {extensao_txt}{desague_txt}.")
                        conclusoes_lista.append((m_nome, "Insuficiência valvar do tronco da VSM sem segmento incompetente definido."))

        # Impressão das Medidas da VSM
        if "Pérvia" in vm.get("status_geral", ""):
            medidas_vsm = [("Junção safenofemoral:", dm['jsf_mm']), ("Terço proximal da coxa:", dm['vsm_prox_coxa']), ("Terço médio da coxa:", dm['vsm_med_coxa']), ("Terço distal da coxa:", dm['vsm_dist_coxa']), ("Terço proximal da perna:", dm['vsm_prox_perna']), ("Terço médio da perna:", dm['vsm_med_perna']), ("Terço distal da perna:", dm['vsm_dist_perna'])]
            medidas_vsm_ativas = [(lbl, val) for lbl, val in medidas_vsm if str(val).strip()]
            if medidas_vsm_ativas:
                add_p("Medidas da veia safena magna:", space_before=6, space_after=4)
                for lbl, val in medidas_vsm_ativas: add_p(f" {val} mm", bold_pre=lbl, bullet=True)

        # 2.2 VEIA SAFENA PARVA (VSP) - TEXTO ADICIONADO E CORRIGIDO
        add_p("Veia Safena Parva:", space_before=8, space_after=4)
        vsp_d = dm["vsp_dados_input"]
        vsp_txt_temp = vsp_d.get("template", "")
        
        if "Normal" in vsp_txt_temp:
            add_p("A veia safena parva apresenta-se pérvia, com trajeto anatômico habitual, paredes finas, totalmente compressível e competente em todo o seu trajeto.")
        elif "Ausente" in vsp_txt_temp:
            add_p(f"Veia safena parva ausente cirurgicamente ({vsp_txt_temp}).")
        elif "Junção safenopoplítea ausente" in vsp_txt_temp:
            add_p("Junção safenopoplítea anatomicamente ausente. Observa-se extensão cranial da veia safena parva (Veia de Giacomini) cursando no plano fascial posterior da coxa.")
            conclusoes_lista.append((m_nome, "Extensão cranial anatômica da veia safena parva."))
        elif "JSP Incompetente" in vsp_txt_temp:
            des_vsp = vsp_d.get("desague_tipo", "tributárias epifasciais")
            ref_vsp = vsp_d.get("desague_ref", "")
            pos_vsp = vsp_d.get("desague_pos", "")
            cm_vsp = vsp_d.get("desague_cm", "")
            if ref_vsp == "Interlinha do Joelho" and pos_vsp:
                loc_txt = f"a {cm_vsp} cm {pos_vsp} da interlinha do joelho"
            elif ref_vsp == "Face Plantar" and cm_vsp:
                loc_txt = f"a {cm_vsp} cm da face plantar"
            else:
                loc_txt = vsp_d.get("localizacao_desague", "terço médio da perna")
            add_p(f"Junção safenopoplítea (JSP) incompetente, gerando refluxo valvar patológico na veia safena parva em toda a sua extensão, associado a dilatação e tortuosidade do tronco. O refluxo apresenta deságue/drenagem terminal para {des_vsp}, {loc_txt}.")
            conclusoes_lista.append((m_nome, "Insuficiência troncular da veia safena parva por incompetência da junção safenopoplítea."))

        # Impressão das Medidas da VSP
        if not ("Ausente" in vsp_txt_temp):
            medidas_vsp = [
                ("Junção/Extensão Safenopoplítea:", dm['jsp_mm']),
                ("Crossa da safena parva:", dm['vsp_crossa']),
                ("Terço médio da perna:", dm['vsp_med_perna_diam'])
            ]
            medidas_vsp_ativas = [(lbl, val) for lbl, val in medidas_vsp if str(val).strip()]
            if medidas_vsp_ativas:
                add_p("Medidas da veia safena parva:", space_before=4, space_after=4)
                for lbl, val in medidas_vsp_ativas: add_p(f" {val} mm", bold_pre=lbl, bullet=True)

        # 2.3 VEIAS PERFURANTES INCOMPETENTES RELEVANTES
        if dm["perfurantes_lista"]:
            add_p("VEIAS PERFURANTES INCOMPETENTES RELEVANTES", space_before=8, space_after=4)
            for p_dados in dm["perfurantes_lista"]:
                r_reg = p_dados["regiao"]
                r_face = p_dados["face"]
                r_ref = p_dados["ref_ponto"]
                r_alt = p_dados["altura_cm"]
                r_pos = p_dados["posicao_joelho"]
                
                txt_ref_perf = f"localizada {r_pos.lower()} da interlinha do joelho" if "Interlinha" in r_ref else f"a partir da face plantar"
                r_diam = p_dados.get("diametro_mm", "")
                txt_diam = f", com diâmetro de {r_diam} mm" if r_diam else ""
                add_p(f"Identificada veia perfurante incompetente na {r_reg.lower()}, face {r_face.lower()}{txt_diam}, medindo {r_alt} cm de altura {txt_ref_perf}.", bullet=True)
                conclusoes_lista.append((m_nome, f"Insuficiência de veia perfurante na {r_reg.lower()} (face {r_face.lower()})."))

        # 2.4 MAPA DE VARICOSIDADES
        vd = dm["varic_dados"]
        if vd.get("possui", False):
            v_tipos = []
            if vd.get("telangiectasias"): v_tipos.append("telangiectasias (< 1 mm de diâmetro)")
            if vd.get("micro_reticulares"): v_tipos.append("microvarizes / varizes reticulares (1 a 3 mm de diâmetro)")
            if vd.get("veias_varicosas"): v_tipos.append("veias varicosas tronculares (> 3 mm de diâmetro)")
            
            if v_tipos:
                txt_varicosidades = "Presença de lesões vasculares superficiais do tipo: " + ", ".join(v_tipos) + f", localizadas predominantemente {vd.get('localizacao', '')}."
                add_p(txt_varicosidades, space_before=6)

        # 3. MÓDULOS ADICIONAIS EXTRA (Giacomini Isolado, Pélvicas)
        if dm["giacomini_opt"] != "Não se aplica / Normal" or dm["varizes_pelvicas_opt"] != "Ausentes":
            add_p("OUTROS ACHADOS FLUXOMÉTRICOS / PONTOS DE ESCAPE", space_before=8)
            if "3.1" in dm["giacomini_opt"]:
                add_p("• Veia de Giacomini com refluxo ostial de sentido ascendente.", bullet=True)
            elif "3.2" in dm["giacomini_opt"]:
                add_p("• Veia de Giacomini isolada transferindo refluxo diretamente para o tronco da VSM.", bullet=True)
            if "5.1" in dm["varizes_pelvicas_opt"]:
                add_p("• Sinais de escape hemodinâmico de origem pélvica via plexo venoso ciático (região infraglútea).", bullet=True)
            elif "5.2" in dm["varizes_pelvicas_opt"]:
                add_p("• Sinais de escape hemodinâmico de origem pélvica via região inguinal/perineal.", bullet=True)
            if "6.3" in dm["pos_op_opt"]:
                add_p("• Sinais de neovascularização adjacentes à região operada (recidiva pós-operatória).", bullet=True)

        # 5. ACHADOS ADICIONAIS (Cisto Baker / Edema)
        if dm["achados_adi_multi"]:
            add_p("ACHADOS ADICIONAIS / ESTRUTURAS ADJACENTES", space_before=10, space_after=4)
            for achado in dm["achados_adi_multi"]:
                if "Cisto de Baker" in achado:
                    cm = dm.get("cisto_medidas", {})
                    txt_cisto = f"Identificado Cisto de Baker na fossa poplítea medindo {cm.get('comp','__')} x {cm.get('larg','__')} x {cm.get('esp','__')} mm (Eixo Longitudinal x Transversal x Espessura)."
                    add_p(txt_cisto, bullet=True)
                    conclusoes_lista.append((m_nome, "Cisto de Baker na fossa poplítea."))
                if "Edema intersticial" in achado:
                    add_p("Sinais ultrassonográficos compatíveis com edema intersticial / infiltração líquida no tecido subcutâneo.", bullet=True)
                    conclusoes_lista.append((m_nome, "Edema intersticial subcutâneo no membro avaliado."))

    # IMPRESSÃO DIAGNÓSTICA (CONCLUSÃO)
    if quebrar_pagina_diag:
        doc.add_page_break()
    add_p("⸻", space_after=12)
    add_p("IMPRESSÃO DIAGNÓSTICA", space_after=6)
    if not conclusoes_lista:
        add_p("Veias safenas magna e parva competentes, sem evidências de refluxo hemodinamicamente significativo.", bullet=True)
    else:
        vistas = set()
        conclusoes_unicas = []
        for m_origem, c_txt in conclusoes_lista:
            chave = (m_origem, c_txt)
            if chave not in vistas:
                vistas.add(chave)
                conclusoes_unicas.append(chave)
                
        for m_origem, conclusao_txt in conclusoes_unicas:
            prefixo = f"[{m_origem}] " if formato_exame == "Bilateral (Laudo Único)" else ""
            add_p(f"{prefixo}{conclusao_txt}", bullet=True)

    if nome_medico or crm_medico:
        doc.add_paragraph().paragraph_format.space_before = Pt(25)
        p_assin = doc.add_paragraph()
        p_assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_assin.paragraph_format.line_spacing = espacamento_linhas
        if nome_medico:
            r = p_assin.add_run(f"{nome_medico}\n")
            r.bold = True
        if crm_medico:
            p_assin.add_run(f"CRM-{crm_uf} {crm_medico}\n")
        if rqe_medico.strip():
            p_assin.add_run(f"RQE {rqe_medico}")
    return doc

# --- PROCESSAMENTO DO BOTÃO DE GERAR LAUDO ---
def _exibir_doc(doc, buf, nome_arquivo, label_download):
    st.success("Laudo gerado com sucesso!")
    if modo_saida in ["Somente Visualização", "Visualização + DOCX"]:
        texto_viz = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        st.markdown("## 👁️ Visualização do Laudo")
        st.text_area("Laudo Gerado", value=texto_viz, height=600)
    if modo_saida in ["Somente DOCX", "Visualização + DOCX"]:
        st.download_button(label_download, buf, nome_arquivo,
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)

if st.button("🚀 Gerar Laudo Venoso Completo", use_container_width=True):
    if formato_exame == "Bilateral (Laudos Separados)":
        for m_nome in ["DIREITO", "ESQUERDO"]:
            doc_sep = construir_laudo_word([m_nome], dados_membros)
            buf = BytesIO()
            doc_sep.save(buf)
            buf.seek(0)
            _exibir_doc(doc_sep, buf, f"Laudo_Venoso_MMII_{m_nome}.docx",
                        f"📥 Baixar Laudo Membro {m_nome} (.docx)")
    else:
        doc_unico = construir_laudo_word(membros_para_processar, dados_membros)
        buf = BytesIO()
        doc_unico.save(buf)
        buf.seek(0)
        _exibir_doc(doc_unico, buf, "Laudo_Venoso_MMII.docx", "📥 Baixar Laudo Venoso (.docx)")